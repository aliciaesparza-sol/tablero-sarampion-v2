import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

docx_path = r"C:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\CONASABI\EVIDENCIAS CONASABI_27FEBRERO2026\Resumen_Vacunacion_Simplificado.docx"
pdf_path = r"C:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\CONASABI\EVIDENCIAS CONASABI_27FEBRERO2026\Resumen_Vacunacion_Simplificado.pdf"

# Generate DOCX
try:
    doc = Document()
    doc.add_heading('Resumen de Vacunación - Corte: 27 de febrero de 2026', 0)

    doc.add_heading('Dosis Aplicadas por Grupos Prioritarios y Rubros Clave (2026)', level=1)
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Grupo Prioritario / Rubro'
    hdr_cells[1].text = 'Dosis Acumuladas (Corte 20 feb)'
    hdr_cells[2].text = 'Dosis Acumuladas (Corte 27 feb)'
    hdr_cells[3].text = 'Incremento (Nuevas Dosis)'
    hdr_cells[4].text = 'Incremento (%)'

    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    data = [
        ('Vacunación Institucional (10 a 49 años)', '', '', '', ''),
        ('Personal de Salud', '2,625', '2,780', '+155', '+5.9%'),
        ('Personal Educativo', '606', '645', '+39', '+6.4%'),
        ('Jornaleros Agrícolas', '301', '320', '+19', '+6.3%'),
        ('Total Institucional', '3,532', '3,745', '+213', '+6.0%'),
        ('Grupos Etarios Prioritarios', '', '', '', ''),
        ('20–29 años', '7,455', '7,866', '+411', '+5.5%'),
        ('30–39 años', '7,061', '7,587', '+526', '+7.4%'),
        ('Total Etario Prioritario', '14,516', '15,453', '+937', '+6.5%'),
        ('Vacunación General (Infantil)', '', '', '', ''),
        ('SRP 6–11 meses (Dosis Cero)', '7,992', '8,764', '+772', '+9.6%'),
        ('SRP 18 meses (2ª Dosis)', '15,495', '16,382', '+887', '+5.7%'),
        ('Total Vacunación Infantil', '23,487', '25,146', '+1,659', '+7.1%'),
        ('TOTAL GLOBAL (Todas las estrategias)', '41,535', '44,344', '+2,809', '+6.8%')
    ]

    for row in data:
        row_cells = table.add_row().cells
        for i in range(5):
             row_cells[i].text = row[i]
             # If it's the Total Global row, make it bold
             if row[0].startswith('TOTAL GLOBAL') or row[0].startswith('Total'):
                 for paragraph in row_cells[i].paragraphs:
                     for run in paragraph.runs:
                         run.bold = True
             
    doc.add_heading('Nuevas Directrices de Vacunación (Actualización Vigente)', level=1)
    p = doc.add_paragraph()
    runner = p.add_run('Regla de oro: Todas aquellas personas que ya cuenten con 2 dosis documentadas en su cartilla NO requieren vacunarse nuevamente.')
    runner.bold = True

    doc.add_heading('1. Población Infantil (6 meses a 12 años)', level=2)
    doc.add_paragraph('6 a 11 meses: Aplicar Dosis Cero.', style='List Bullet')
    doc.add_paragraph('12 meses: Aplicar 1ª dosis.', style='List Bullet')
    doc.add_paragraph('18 meses: Aplicar 2ª dosis.', style='List Bullet')
    doc.add_paragraph('Población Rezagada (1 a 12 años): Iniciar o completar el esquema conforme a su edad.', style='List Bullet')

    doc.add_heading('2. Población General (13 a 49 años)', level=2)
    doc.add_paragraph('Sin esquema completo: Aplicar una dosis de SR o SRP (dependiendo del escenario epidemiológico vigente).', style='List Bullet')

    doc.add_heading('3. Sectores Institucionales', level=2)
    doc.add_paragraph('Personal de Salud: Aplicar una dosis adicional de SR (independientemente de su esquema previo documentado).', style='List Bullet')
    doc.add_paragraph('Personal Docente y Jornaleros Agrícolas (hasta 49 años): Iniciar o completar esquema en caso de estar incompleto.', style='List Bullet')

    doc.save(docx_path)
    print(f"Created DOCX at: {docx_path}")
except Exception as e:
    print(f"DOCX error: {e}")

try:
    from docx2pdf import convert
    convert(docx_path, pdf_path)
    print(f"Created PDF at: {pdf_path}")
except Exception as e:
    print(f"PDF creation failed: {e}")
