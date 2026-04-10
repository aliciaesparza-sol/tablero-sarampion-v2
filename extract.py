import sys
from PyPDF2 import PdfReader
from docx import Document

def extract_pdf(path, out_file):
    try:
        reader = PdfReader(path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n\n"
        with open(out_file, 'w', encoding='utf-8') as f:
            f.write(text)
    except Exception as e:
        print(f"Error reading PDF: {e}")

def extract_docx(path, out_file):
    try:
        doc = Document(path)
        text = ""
        for p in doc.paragraphs:
            text += p.text + "\n"
        
        # also get tables
        for table in doc.tables:
            text += "\n--- Table ---\n"
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.replace('\n', ' '))
                text += " | ".join(row_text) + "\n"
            text += "-------------\n"
            
        with open(out_file, 'w', encoding='utf-8') as f:
            f.write(text)
    except Exception as e:
        print(f"Error reading DOCX: {e}")

extract_pdf(r"c:\Users\aicil\OneDrive\Escritorio\PVU\SEMANA NACIONAL DE VACUNACION\SEMANA NACIONAL DE VACUNACIÓN 2026\56_CIR CENSIA_SemanaNacionaldeVacunación_2026.pdf", "pdf_text.txt")
extract_docx(r"c:\Users\aicil\OneDrive\Escritorio\PVU\SEMANA NACIONAL DE VACUNACION\SEMANA NACIONAL DE VACUNACIÓN 2026\FICHA TECNICA CAMAPAÑA 25 ABRIL AL 2 DE MAYO.docx", "ficha_text.txt")
extract_docx(r"c:\Users\aicil\OneDrive\Escritorio\hoja MEMBRETADA GIGANTE2026_carta.docx", "template_text.txt")
