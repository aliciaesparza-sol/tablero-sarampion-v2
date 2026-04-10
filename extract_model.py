import docx
import os

def read_full_docx(file_path, output_txt):
    if not os.path.exists(file_path):
        return
    doc = docx.Document(file_path)
    with open(output_txt, "w", encoding="utf-8") as f:
        f.write(f"--- CONTENT OF {os.path.basename(file_path)} ---\n\n")
        f.write("PARAGRAPHS:\n")
        for para in doc.paragraphs:
            f.write(para.text + "\n")
        
        f.write("\nTABLES:\n")
        for i, table in enumerate(doc.tables):
            f.write(f"\nTable {i}:\n")
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                f.write(" | ".join(row_text) + "\n")

model_path = r"c:\Users\aicil\OneDrive\Escritorio\PVU\CAMPAÑA DE VACUNACION MASIVA INTERINSTITUCIONAL DICIEMBRE 2025\MACROCENTROS UNIVERSITARIOS 25 MARZO 2026\FICHA TECNICA MACROCENTROS DE VACUNACIÓN UNIVERSITARIOS 25 MARZO 2026.docx"
read_full_docx(model_path, "model_content.txt")
