from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ---- Page margins ----
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3)
section.right_margin = Cm(2.5)

def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, color=(31,56,100)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.bold = True
    return p

def add_para(doc, text, bold=False, italic=False, color=None, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if bold_prefix:
        run1 = p.add_run(bold_prefix + " ")
        set_font(run1, bold=True, color=(31,56,100))
    run2 = p.add_run(text)
    set_font(run2)
    p.paragraph_format.space_after = Pt(3)

def add_divider(doc):
    p = doc.add_paragraph("─" * 80)
    p.runs[0].font.color.rgb = RGBColor(189,215,238)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

def add_highlight_box(doc, label, text, bg=(240,245,255)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run(f"{label}  ")
    set_font(r1, bold=True, color=(31,56,100), size=11)
    r2 = p.add_run(text)
    set_font(r2, size=11, color=(50,50,50))
    p.paragraph_format.space_after = Pt(4)
    return p

# ==================== DOCUMENT ====================

# COVER
add_para(doc, "SECRETARÍA DE SALUD — DURANGO", bold=True, size=13, color=(31,56,100), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "CAMPAÑA DE VACUNACIÓN CONTRA SARAMPIÓN 2026", bold=True, size=15, color=(31,56,100), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Periodo: 18 de febrero al 14 de marzo de 2026", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para(doc, "INFORME DE ANÁLISIS: EXISTENCIAS, CONSUMO Y DOSIS APLICADAS DE VACUNAS SR Y SRP POR JURISDICCIÓN SANITARIA", bold=True, size=14, color=(31,56,100), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para(doc, "Elaborado por: Programa de Vacunación Universal — Dirección General de Salud Pública", size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Fecha: 14 de marzo de 2026", size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ==== SECCIÓN 1: EXPLICACIÓN SIMPLE ====
add_heading(doc, "1.  ¿De qué se trata este informe? (En palabras simples)", level=1)

add_para(doc, "Imagina que tienes una bodega con manzanas y cada día las repartes a cuatro tiendas (cada tienda es una jurisdicción). Al final del día, cuentas cuántas manzanas quedaron en cada tienda. Con eso puedes saber cuántas manzanas se \"usaron\" ese día (aunque no las hayas visto vender directamente).", italic=True, color=(80,80,80))
doc.add_paragraph()

add_para(doc, "En este informe, las manzanas son las vacunas (SR y SRP) y las tiendas son las 4 jurisdicciones sanitarias del estado de Durango. Hacemos tres preguntas clave:")
add_bullet(doc, "¿Cuántas vacunas había en cada jurisdicción cada día? → Existencias", bold_prefix="📦")
add_bullet(doc, "¿Cuántas vacunas se usaron (desaparecieron del inventario) cada día? → Consumo teórico", bold_prefix="📉")
add_bullet(doc, "¿Cuántas dosis de vacuna se registraron como aplicadas en el sistema oficial? → Dosis aplicadas", bold_prefix="💉")

doc.add_paragraph()
add_para(doc, "Si el inventario baja en 500 vacunas pero el sistema solo registra 200 dosis aplicadas, hay una diferencia de 300. Esas 300 dosis o no se registraron, o salieron del almacén por otra razón. Eso es lo que analizamos aquí.", bold=False, size=11)

add_divider(doc)

# ==== SECCIÓN 2: ¿QUÉ SON LAS VACUNAS SR Y SRP? ====
add_heading(doc, "2.  ¿Qué es la vacuna SR y cuál es la SRP?", level=1)

add_bullet(doc, "Protege contra: Sarampión y Rubéola. Se aplica como refuerzo en adultos y campaña intensiva.", bold_prefix="SR (Doble Viral) →")
add_bullet(doc, "Protege contra: Sarampión, Rubéola y Parotitis (paperas). Se aplica principalmente en niños.", bold_prefix="SRP (Triple Viral) →")
doc.add_paragraph()
add_para(doc, "En esta campaña 2026, ambas vacunas se distribuyeron a los 4 jurisdicciones del estado y los puntos de vacunación bajo su responsabilidad.")

add_divider(doc)

# ==== SECCIÓN 3: Las 4 Jurisdicciones ====
add_heading(doc, "3.  Las 4 Jurisdicciones Sanitarias", level=1)
add_bullet(doc, "Durango (Zona urbana, mayor concentración de población)", bold_prefix="Jurisdicción No. 1 →")
add_bullet(doc, "Santiago Papasquiaro (Zona serrana norte)", bold_prefix="Jurisdicción No. 2 →")
add_bullet(doc, "Rodeo (Zona sur / cañones)", bold_prefix="Jurisdicción No. 3 →")
add_bullet(doc, "Mezquital (Zona indígena / barrancas)", bold_prefix="Jurisdicción No. 4 →")

add_divider(doc)

# ==== SECCIÓN 4: HALLAZGOS ====
add_heading(doc, "4.  ¿Qué encontramos? — Hallazgos por Jurisdicción", level=1)

add_heading(doc, "4.1  Jurisdicción No. 1 — Durango (la tienda más grande)", level=2, color=(31,56,100))
add_para(doc, "Esta es la jurisdicción con más actividad. Al inicio del periodo (18 feb) tenía solo 350 dosis SR en los puntos de vacunación. Al final del periodo (14 mar) tenía 12,090 dosis. Eso significa que recibió muchas vacunas durante el mes.")
doc.add_paragraph()
add_para(doc, "En la vacuna SRP pasó lo contrario: empezó con 5,127 dosis y terminó con solo 415. Consumió casi toda su existencia SRP — eso es bueno, es señal de que se aplicaron.")
doc.add_paragraph()
add_para(doc, "⚠️  Problema detectado: En varios días (03 al 11 de marzo principalmente) el stock SR bajó mucho, pero el sistema CENSIA tiene muy pocas dosis registradas como aplicadas en esas fechas. El inventario dice que se usaron las vacunas, pero el sistema no lo refleja. Esto podría ser subreporte.", size=11, color=(192,0,0))

add_heading(doc, "4.2  Jurisdicción No. 2 — Santiago Papasquiaro (la tienda serrana)", level=2, color=(31,56,100))
add_para(doc, "Esta jurisdicción arrancó sin nada de vacuna SR (0 dosis) y fue recibiendo dosis durante el mes hasta llegar a 5,190 al final. Eso indica que el abastecimiento llegó tarde o fue gradual.")
doc.add_paragraph()
add_para(doc, "La vacuna SRP inició en 2,014 dosis y bajó a 989. Usó aproximadamente 1,025 dosis SRP durante el periodo — bien consumidas.")
doc.add_paragraph()
add_para(doc, "⚠️  Problema detectado: Los días 10, 11, 12 y 13 de marzo presentan diferencias mayores a 300 dosis de SR entre lo que bajó del inventario y lo que se registró en CENSIA. Posiblemente hay aplicaciones sin capturar.", size=11, color=(192,0,0))

add_heading(doc, "4.3  Jurisdicción No. 3 — Rodeo (la tienda pequeña)", level=2, color=(31,56,100))
add_para(doc, "Es la jurisdicción con menor volumen. Nunca tuvo más de 390 dosis SR en sus puntos — territorio pequeño, pocos puntos de vacunación.")
doc.add_paragraph()
add_para(doc, "La vacuna SRP bajó de 2,211 a 1,300 dosis durante el periodo (usó ~911 dosis).")
doc.add_paragraph()
add_para(doc, "⚠️  Problema detectado: El sistema CENSIA tiene muy pocos registros de dosis aplicadas para esta jurisdicción, a pesar de que el inventario sí disminuyó. Necesita refuerzo en la captura diaria del registro.", size=11, color=(192,0,0))

add_heading(doc, "4.4  Jurisdicción No. 4 — Mezquital (la tienda más lejana)", level=2, color=(31,56,100))
add_para(doc, "Esta jurisdicción tuvo entre 140 y 1,072 dosis SR en distintos momentos del mes, con fluctuaciones que indican entregas intermedias.")
doc.add_paragraph()
add_para(doc, "Consumo acumulado aproximado: 712 dosis SR y 235 dosis SRP.")
doc.add_paragraph()
add_para(doc, "⚠️  Problema detectado: Los días 02/03, 09/03 y del 04 al 06 de marzo presentan variaciones de stock sin registro correspondiente en CENSIA. Desfase medio de ~168 dosis/día en SR.", size=11, color=(192,0,0))

add_divider(doc)

# ==== SECCIÓN 5: CUADRO RESUMEN ====
add_heading(doc, "5.  Cuadro Resumen del Periodo (18 Feb — 14 Mar 2026)", level=1)

table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, h in enumerate(['Jurisdicción', 'Stock SR Final', 'Consumo SR Neto', 'Stock SRP Final', 'Consumo SRP Neto']):
    hdr_cells[i].text = h
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

rows_data = [
    ('Jur. 1 — Durango',         '12,090 dosis', '+11,740 (reabastecida)', '415 dosis', '4,712 consumidas' ),
    ('Jur. 2 — Sto. Papasquiaro','5,190 dosis',  '+5,190 (recibió stock)', '989 dosis', '1,025 consumidas' ),
    ('Jur. 3 — Rodeo',           '60 dosis',     '160 neto',               '1,300 dosis','911 consumidas'  ),
    ('Jur. 4 — Mezquital',       '360 dosis',    '712 consumidas',         '813 dosis', '235 consumidas'   ),
]
for rd in rows_data:
    row = table.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        row[i].paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

add_divider(doc)

# ==== SECCIÓN 6: QUÉ SIGNIFICA EL DESFASE ====
add_heading(doc, "6.  ¿Qué significa cuando hay una diferencia (desfase)?", level=1)

add_para(doc, "En términos simples: si el inventario dice que salieron 500 manzanas de la tienda, pero la cajera solo registró 200 ventas, hay 300 manzanas \"sin explicación\". Lo mismo pasa con las vacunas.")
doc.add_paragraph()

add_para(doc, "Posibles causas del desfase:")
add_bullet(doc, "El vacunador aplicó dosis pero no las capturó en el sistema ese mismo día")
add_bullet(doc, "Se trasladaron dosis de un punto a otro sin registrarse como salida")
add_bullet(doc, "Hubo bajas (caducidad, ruptura de cadena de frío) que reducen el inventario sin contar como dosis aplicadas")
add_bullet(doc, "Error en el conteo de existencias reportado")

doc.add_paragraph()
add_para(doc, "🔍 En este periodo, la mayor parte de los desfases apuntan a subreporte en CENSIA: las vacunas sí se aplicaron, pero no se registraron a tiempo o correctamente en el sistema.", bold=True, color=(31,56,100))

add_divider(doc)

# ==== SECCIÓN 7: RECOMENDACIONES ====
add_heading(doc, "7.  Recomendaciones", level=1)

add_bullet(doc, "Revisar y actualizar el registro de dosis en CENSIA para los días con desfases detectados (especialmente 03–13 de marzo).", bold_prefix="1.")
add_bullet(doc, "Implementar una conciliación diaria: al cerrar el turno, el responsable jurisdiccional debe cruzar el conteo físico de vacunas con el registro del sistema.", bold_prefix="2.")
add_bullet(doc, "Fortalecer la capacitación en la Jurisdicción No. 3 (Rodeo) para mejorar la captura en sistema.", bold_prefix="3.")
add_bullet(doc, "Verificar, para las fechas donde el stock subió sin reporte de entrega, si hubo traslados de lote no registrados.", bold_prefix="4.")
add_bullet(doc, "Continuar el monitoreo diario de existencias para garantizar abasto suficiente hasta el cierre de la campaña.", bold_prefix="5.")

add_divider(doc)

# ==== SECCIÓN 8: CONCLUSIÓN ====
add_heading(doc, "8.  Conclusión", level=1)
add_para(doc, "El estado de Durango mantuvo un abasto de vacunas SR y SRP distribuido a sus 4 jurisdicciones durante todo el periodo analizado. La Jurisdicción No. 1 (Durango) concentra el mayor movimiento de biológicos, mientras que las jurisdicciones serranas operan con volúmenes más pequeños pero estables.")
doc.add_paragraph()
add_para(doc, "La principal área de mejora detectada es el registro oportuno y completo de dosis aplicadas en el sistema CENSIA. El inventario físico y el registro digital no coinciden en múltiples fechas y jurisdicciones. Corregir esto es esencial para garantizar reportes precisos al nivel federal y para la toma de decisiones oportunas.")
doc.add_paragraph()
add_para(doc, "Con base en las existencias actuales y el ritmo de consumo observado, se recomienda verificar el abasto proyectado para el resto de la campaña, especialmente en las jurisdicciones No. 3 y No. 4 que presentan los menores stocks finales.", bold=False)

out_path = r"c:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\INFORME_ANALISIS_VACUNAS_2026.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
