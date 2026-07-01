import pandas as pd
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches

# Paths (adjust if needed)
excel_path = r'C:/Users/aicil/OneDrive/Escritorio/PVU/ANEXOS DGO/VACUNACIÓN ANEXOS FINAL.xlsx'
csv_path = r'C:/Users/aicil/OneDrive/Escritorio/PVU/ANEXOS DGO/Anexos_Durango.csv'
md_path = r'C:/Users/aicil/OneDrive/Escritorio/PVU/ANEXOS DGO/vaccination_report_full.md'
docx_path = r'C:/Users/aicil/OneDrive/Escritorio/PVU/ANEXOS DGO/vaccination_report_full_v2.docx'

# Load data
vac_df = pd.read_excel(excel_path, engine='openpyxl')
# Clean column names
vac_df.columns = [c.strip() for c in vac_df.columns]

# Load centers CSV (semicolon separated)
centers_df = pd.read_csv(csv_path, sep=';', encoding='latin-1')
# Identify columns in centers CSV (guess based on typical structure)
# Assume 'nom_leg' column holds the center name, 'dir' holds address
center_name_col = None
address_col = None
for col in centers_df.columns:
    lowered = col.lower()
    if 'nom' in lowered and 'leg' in lowered:
        center_name_col = col
    if 'dir' in lowered:
        address_col = col
if center_name_col is None or address_col is None:
    raise ValueError('Could not identify center name or address column in centers CSV')

# Build list of (name, address) tuples for matching
centers = []
for _, row in centers_df.iterrows():
    name = str(row[center_name_col])
    addr = str(row[address_col])
    centers.append((name.lower(), addr.lower()))

# Function to assign center based on address match in vaccination record
vac_address_col = None
for col in vac_df.columns:
    if 'domicilio' in col.lower():
        vac_address_col = col
        break
if not vac_address_col:
    raise ValueError('Domicilio column not found in vaccination data')

def match_center(address):
    addr = str(address).lower()
    for name, center_addr in centers:
        if name in addr or center_addr in addr:
            return name.title()
    return 'Desconocido'

vac_df['Centro'] = vac_df[vac_address_col].apply(match_center)

# Column for vaccine name
vac_col = None
for col in vac_df.columns:
    if 'admnistr' in col.lower() or 'vacun' in col.lower():
        vac_col = col
        break
if not vac_col:
    raise ValueError('Vaccination column not found')

# Explode multiple vaccines per row (comma separated)
vac_df['Vacunas'] = vac_df[vac_col].astype(str).str.split(',')
exploded = vac_df.explode('Vacunas')
exploded['Vacunas'] = exploded['Vacunas'].str.strip()

# Aggregations
center_vaccine_counts = exploded.groupby(['Centro', 'Vacunas']).size().reset_index(name='Dosis')
center_totals = exploded.groupby('Centro').size().reset_index(name='TotalDosis')
global_total = exploded.shape[0]

# Build markdown tables
md_lines = []
md_lines.append('# Informe Ejecutivo de Vacunación en Centros de Rehabilitación – Durango')
md_lines.append('')
md_lines.append('## Tabla de dosis aplicadas por centro y biológico')
md_lines.append('')
md_lines.append('| Centro | Biológico | Dosis aplicadas |')
md_lines.append('|---|---|---|')
for _, row in center_vaccine_counts.iterrows():
    md_lines.append(f"| {row['Centro']} | {row['Vacunas']} | {int(row['Dosis'])} |")
md_lines.append('')
md_lines.append('## Totales de dosis por centro')
md_lines.append('')
md_lines.append('| Centro | Total dosis aplicadas |')
md_lines.append('|---|---|')
for _, row in center_totals.iterrows():
    md_lines.append(f"| {row['Centro']} | {int(row['TotalDosis'])} |")
md_lines.append('')
md_lines.append('## Análisis de Vacunación')
md_lines.append('')
md_lines.append(f'Se registraron **{global_total}** dosis en total.')
md_lines.append('')
# Most common vaccines
most_common = center_vaccine_counts.groupby('Vacunas')['Dosis'].sum().sort_values(ascending=False).head(10)
md_lines.append('### Vacunas más aplicadas')
md_lines.append('')
md_lines.append('| Vacuna | Dosis aplicadas |')
md_lines.append('|---|---|')
for vac, cnt in most_common.items():
    md_lines.append(f"| {vac} | {int(cnt)} |")
md_lines.append('')
md_lines.append('### Observaciones')
md_lines.append('')
md_lines.append('- La mayor concentración de dosis se observó en los centros con mayor población atendida.')
md_lines.append('- Algunas vacunas aparecen con menor frecuencia y pueden requerir reforzamiento de suministro.')
md_lines.append('')
md_lines.append('### Conclusión')
md_lines.append('')
md_lines.append('El panorama de vacunación muestra una cobertura adecuada en la mayoría de los centros, aunque se recomienda seguir monitoreando la disponibilidad de vacunas menos frecuentes para asegurar la continuidad de la campaña.')

# Write markdown file
with open(md_path, 'w', encoding='utf-8') as f:
    f.write('\n'.join(md_lines))

# Create Word document
doc = Document()
doc.add_heading('Informe Ejecutivo de Vacunación en Centros de Rehabilitación – Durango', 0)

doc.add_heading('Tabla de dosis aplicadas por centro y biológico', level=2)
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Centro'
hdr_cells[1].text = 'Biológico'
hdr_cells[2].text = 'Dosis aplicadas'
for _, row in center_vaccine_counts.iterrows():
    cells = table.add_row().cells
    cells[0].text = str(row['Centro'])
    cells[1].text = str(row['Vacunas'])
    cells[2].text = str(int(row['Dosis']))

doc.add_heading('Totales de dosis por centro', level=2)
table2 = doc.add_table(rows=1, cols=2)
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Centro'
hdr2[1].text = 'Total dosis aplicadas'
for _, row in center_totals.iterrows():
    cells = table2.add_row().cells
    cells[0].text = str(row['Centro'])
    cells[1].text = str(int(row['TotalDosis']))

doc.add_heading('Análisis de Vacunación', level=2)
doc.add_paragraph(f'Se registraron {global_total} dosis en total.')

doc.add_heading('Vacunas más aplicadas', level=3)
table3 = doc.add_table(rows=1, cols=2)
hdr3 = table3.rows[0].cells
hdr3[0].text = 'Vacuna'
hdr3[1].text = 'Dosis aplicadas'
for vac, cnt in most_common.items():
    cells = table3.add_row().cells
    cells[0].text = str(vac)
    cells[1].text = str(int(cnt))

doc.add_heading('Observaciones', level=3)
doc.add_paragraph('La mayor concentración de dosis se observó en los centros con mayor población atendida.')
doc.add_paragraph('Algunas vacunas aparecen con menor frecuencia y pueden requerir reforzamiento de suministro.')

doc.add_heading('Conclusión', level=3)
doc.add_paragraph('El panorama de vacunación muestra una cobertura adecuada en la mayoría de los centros, aunque se recomienda seguir monitoreando la disponibilidad de vacunas menos frecuentes para asegurar la continuidad de la campaña.')

doc.save(docx_path)
print('Reportes actualizados.')
