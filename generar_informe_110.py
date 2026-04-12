import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Helper functions ──────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top','left','bottom','right','insideH','insideV']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'),   'single')
                border.set(qn('w:sz'),    '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'BDBDBD')
                tcBorders.append(border)
            tcPr.append(tcBorders)

def heading(doc, text, level=1, color='1A237E'):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(13 if level == 1 else 11)
    run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(doc, text, bold=False, size=10):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_image(doc, path, width_cm=16):
    try:
        doc.add_picture(path, width=Cm(width_cm))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        body(doc, f"[imagen no disponible: {e}]")

# ══════════════════════════════════════════════════════════════
#  PORTADA / ENCABEZADO
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SECRETARIA DE SALUD DEL ESTADO DE DURANGO")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor.from_string('1A237E')
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PROGRAMA DE VACUNACION UNIVERSAL")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor.from_string('1A237E')
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CAMPAÑA DE VACUNACION CONTRA SARAMPION 2026")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string('B71C1C')
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("INFORME EJECUTIVO – SEMANA 8 DE 10")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor.from_string('1A237E')
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Fecha de corte: 10 de abril de 2026")
run.font.size = Pt(10)
p.paragraph_format.space_after = Pt(10)

doc.add_paragraph()  # espacio

# ══════════════════════════════════════════════════════════════
#  1. RESUMEN EJECUTIVO
# ══════════════════════════════════════════════════════════════
heading(doc, "1. RESUMEN EJECUTIVO")

resumen = (
    "La Campaña de Vacunacion contra Sarampion SRP/SR en el estado de Durango "
    "se desarrollo del 19 de febrero al 29 de abril de 2026, con una duracion "
    "programada de 10 semanas. Al corte del 10 de abril de 2026 (Semana 8), "
    "la estrategia registra 297,251 dosis aplicadas a nivel sectorial, "
    "representando el 65.0% de la meta acumulada correspondiente a las primeras "
    "ocho semanas (457,475 dosis). Se cuenta con 110 brigadas activas operando "
    "en el territorio estatal. La campana tiene 51 dias transcurridos y 20 dias restantes "
    "para concluir el esquema de 10 semanas."
)
body(doc, resumen)

# ══════════════════════════════════════════════════════════════
#  2. INDICADORES GENERALES – SECTORIAL
# ══════════════════════════════════════════════════════════════
heading(doc, "2. INDICADORES GENERALES DE LA CAMPAÑA – NIVEL SECTORIAL")

# tabla resumen sectorial
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for i, txt in enumerate(["INDICADOR", "META", "LOGRADO", "AVANCE"]):
    hdr[i].text = txt
    hdr[i].paragraphs[0].runs[0].bold = True
    hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(hdr[i], '1A237E')
    hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string('FFFFFF')
    hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

rows_data = [
    ("Meta total 10 semanas",       "522,608 dosis",  "297,251 dosis", "56.9%"),
    ("Meta acumulada Sem 1-8",       "457,475 dosis",  "297,251 dosis", "65.0%"),
    ("Dias transcurridos / restantes","71 dias totales","51 dias",       "20 restantes"),
    ("Brigadas activas",             "—",              "110",            "—"),
    ("Ritmo requerido",              "52,260.8 dos/sem","—",            "—"),
    ("Ritmo actual",                 "—",              "40,799 dos/sem","—"),
]

for rd in rows_data:
    row = table.add_row().cells
    for i, val in enumerate(rd):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(9)
        row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i == 0:
            row[i].paragraphs[0].runs[0].bold = True
            set_cell_bg(row[i], 'E8EAF6')
set_cell_borders(table)

doc.add_paragraph()

# ── Dashboard sectorial imagen 1
body(doc, "Figura 1. Tablero sectorial – Semana 8 de 10 (corte 10 de abril de 2026).", bold=True, size=9)
add_image(doc, r'C:\Users\aicil\.gemini\antigravity\scratch\docx_images\word\media\image1.png', width_cm=15)
doc.add_paragraph()

# ── Grafico progreso sectorial
body(doc, "Figura 2. Progreso semanal sectorial – dosis aplicadas vs meta acumulada.", bold=True, size=9)
add_image(doc, r'C:\Users\aicil\.gemini\antigravity\scratch\docx_images\word\media\image2.png', width_cm=15)
doc.add_paragraph()

# ── Tabla desglose semanal sectorial
body(doc, "Figura 3. Desglose semanal detallado – nivel sectorial.", bold=True, size=9)
add_image(doc, r'C:\Users\aicil\.gemini\antigravity\scratch\docx_images\word\media\image3.png', width_cm=15)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  3. PROMEDIO DE DOSIS POR DIA
# ══════════════════════════════════════════════════════════════
heading(doc, "3. PROMEDIO DE DOSIS APLICADAS POR DIA (ESTRATEGIA 10 SEMANAS)")

# Calculos:
# 10 semanas = 70 dias calendario
# Total aplicadas al corte: 297,251  en 51 dias transcurridos
# Promedio real = 297,251 / 51 = 5,829/dia aprox
# Promedio requerido para cumplir meta = (522,608 - 297,251) / 20 = 11,268/dia
# Por brigada activa (60): 5,829/60 = 53 dosis/brigada/dia actual
#                          11,268/60 = 102 dosis/brigada/dia requeridas

prom_texto = (
    "Con base en los 51 dias transcurridos desde el inicio de la campana (19 de febrero de 2026) "
    "y las 297,251 dosis aplicadas a nivel sectorial al 10 de abril de 2026, "
    "se estiman los siguientes promedios diarios:"
)
body(doc, prom_texto)

# Tabla promedios
table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
for i, txt in enumerate(["INDICADOR", "VALOR TOTAL", "POR BRIGADA (60 activas)"]):
    hdr2[i].text = txt
    hdr2[i].paragraphs[0].runs[0].bold = True
    hdr2[i].paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(hdr2[i], '1A237E')
    hdr2[i].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string('FFFFFF')
    hdr2[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

prom_data = [
    ("Dosis aplicadas acumuladas (Sem 1-8)",    "297,251",     "—"),
    ("Dias transcurridos",                       "51",          "—"),
    ("Promedio real de dosis/dia",               "5,829 dosis", "53 dosis/brigada/dia"),
    ("Dosis faltantes para meta",               "225,357",     "—"),
    ("Dias restantes",                           "20",          "—"),
    ("Promedio requerido de dosis/dia",          "11,268 dosis","102 dosis/brigada/dia"),
    ("Promedio meta programada (10 sem/70 dias)","7,466 dosis", "68 dosis/brigada/dia"),
]

alt = False
for rd in prom_data:
    row = table2.add_row().cells
    bg = 'E8EAF6' if not alt else 'FFFFFF'
    alt = not alt
    for i, val in enumerate(rd):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(9)
        row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i == 0:
            row[i].paragraphs[0].runs[0].bold = True
            set_cell_bg(row[i], bg)
set_cell_borders(table2)

doc.add_paragraph()

nota = (
    "El ritmo actual de 40,799 dosis/semana sectorial (equivalente a 5,829 dosis/dia) "
    "es inferior al ritmo requerido de 52,261 dosis/semana (7,466 dosis/dia). "
    "Para alcanzar la meta total de 522,608 dosis en los 20 dias restantes, "
    "se requiere incrementar el promedio diario a 11,268 dosis, "
    "lo que implica que cada una de las 110 brigadas activas debe aplicar "
    "un minimo de 102 dosis por dia."
)
body(doc, nota)

# ══════════════════════════════════════════════════════════════
#  4. AVANCE POR INSTITUCION
# ══════════════════════════════════════════════════════════════
heading(doc, "4. AVANCE POR INSTITUCION")

inst_intro = (
    "A continuacion se presenta el avance por институcion participante en la campana sectorial, "
    "con corte al 10 de abril de 2026 (Semana 8 de 10):"
)
inst_intro = (
    "A continuacion se presenta el avance por institucion participante en la campana sectorial, "
    "con corte al 10 de abril de 2026 (Semana 8 de 10):"
)
body(doc, inst_intro)

# Tabla institucional
table3 = doc.add_table(rows=1, cols=6)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = table3.rows[0].cells
for i, txt in enumerate(["INSTITUCION","META 10 SEM","META ACUM SEM 1-8","APLICADAS","AVANCE vs META ACUM","RITMO ACTUAL/SEM"]):
    hdr3[i].text = txt
    hdr3[i].paragraphs[0].runs[0].bold = True
    hdr3[i].paragraphs[0].runs[0].font.size = Pt(8)
    set_cell_bg(hdr3[i], '1A237E')
    hdr3[i].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string('FFFFFF')
    hdr3[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

inst_data = [
    ("SSA (Jurisdicciones)", "170,247", "147,500", "53,183",  "36.1%",  "7,300"),
    ("IMSS",                 "276,504", "242,663", "213,901", "88.1%",  "29,359"),
    ("ISSSTE",               "75,857",  "67,312",  "30,061",  "44.7%",  "4,126"),
    ("SECTORIAL TOTAL",      "522,608", "457,475", "297,251", "65.0%",  "40,799"),
]

alt = False
for rd in inst_data:
    row  = table3.add_row().cells
    bold = rd[0].startswith("SECT")
    bg   = '1A237E' if bold else ('E8EAF6' if not alt else 'FFFFFF')
    alt  = not alt
    for i, val in enumerate(rd):
        row[i].text = val
        r = row[i].paragraphs[0].runs[0]
        r.font.size = Pt(8)
        r.bold = bold
        row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if bold:
            r.font.color.rgb = RGBColor.from_string('FFFFFF')
            set_cell_bg(row[i], '283593')
        elif i == 0:
            set_cell_bg(row[i], bg)
set_cell_borders(table3)

doc.add_paragraph()

# ── SSA dashboard
body(doc, "Figura 4. Tablero SSA – Semana 8 de 10.", bold=True, size=9)
add_image(doc, r'C:\Users\aicil\.gemini\antigravity\scratch\docx_images\word\media\image5.png', width_cm=15)
doc.add_paragraph()

body(doc, "Figura 5. Progreso semanal SSA.", bold=True, size=9)
add_image(doc, r'C:\Users\aicil\.gemini\antigravity\scratch\docx_images\word\media\image6.png', width_cm=15)
doc.add_paragraph()

body(doc, "Figura 6. Desglose semanal detallado – SSA.", bold=True, size=9)
add_image(doc, r'C:\Users\aicil\.gemini\antigravity\scratch\docx_images\word\media\image7.png', width_cm=15)
doc.add_paragraph()

body(doc, "Figura 7. Tablero IMSS – Semana 8 de 10.", bold=True, size=9)
add_image(doc, r'C:\Users\aicil\.gemini\antigravity\scratch\docx_images\word\media\image9.png', width_cm=15)
doc.add_paragraph()

body(doc, "Figura 8. Progreso semanal IMSS.", bold=True, size=9)
add_image(doc, r'C:\Users\aicil\.gemini\antigravity\scratch\docx_images\word\media\image10.png', width_cm=15)
doc.add_paragraph()

body(doc, "Figura 9. Desglose semanal detallado – IMSS.", bold=True, size=9)
add_image(doc, r'C:\Users\aicil\.gemini\antigravity\scratch\docx_images\word\media\image11.png', width_cm=15)
doc.add_paragraph()

body(doc, "Figura 10. Tablero ISSSTE – Semana 8 de 10.", bold=True, size=9)
add_image(doc, r'C:\Users\aicil\.gemini\antigravity\scratch\docx_images\word\media\image13.png', width_cm=15)
doc.add_paragraph()

body(doc, "Figura 11. Progreso semanal ISSSTE.", bold=True, size=9)
add_image(doc, r'C:\Users\aicil\.gemini\antigravity\scratch\docx_images\word\media\image14.png', width_cm=15)
doc.add_paragraph()

body(doc, "Figura 12. Desglose semanal detallado – ISSSTE.", bold=True, size=9)
add_image(doc, r'C:\Users\aicil\.gemini\antigravity\scratch\docx_images\word\media\image15.png', width_cm=15)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  5. DOSIS POR GRUPO DE EDAD
# ══════════════════════════════════════════════════════════════
heading(doc, "5. DISTRIBUCION DE DOSIS POR GRUPO DE EDAD")

edad_texto = (
    "La distribucion de dosis aplicadas por grupo de edad muestra que el grupo de 2 a 12 anos "
    "concentra el mayor volumen de aplicacion tanto en primera como en segunda dosis SRP, "
    "seguido por el grupo de 20 a 39 anos. Los grupos de 6 a 11 meses y 18 meses presentan "
    "coberturas importantes en el marco de la estrategia de rescate de susceptibles."
)
body(doc, edad_texto)

body(doc, "Figura 13. Dosis aplicadas por grupo de edad – primera y segunda dosis SRP (acumulado campana).", bold=True, size=9)
add_image(doc, r'C:\Users\aicil\.gemini\antigravity\scratch\docx_images\word\media\image4.png', width_cm=14)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  6. EXISTENCIAS DE VACUNA CONTRA SARAMPION
# ══════════════════════════════════════════════════════════════
heading(doc, "6. EXISTENCIAS DE VACUNA CONTRA SARAMPION")

exist_intro = (
    "A continuacion se presentan las existencias de biológico SRP y SR disponibles "
    "en el nivel estatal y por jurisdiccion sanitaria, con corte al 10 de abril de 2026. "
    "El total estatal disponible asciende a 108,783 dosis entre ambos biologicos."
)
body(doc, exist_intro)

# Tabla existencias
table4 = doc.add_table(rows=1, cols=4)
table4.style = 'Table Grid'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr4 = table4.rows[0].cells
for i, txt in enumerate(["LUGAR", "VACUNA SRP", "VACUNA SR", "TOTAL"]):
    hdr4[i].text = txt
    hdr4[i].paragraphs[0].runs[0].bold = True
    hdr4[i].paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(hdr4[i], '1A237E')
    hdr4[i].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string('FFFFFF')
    hdr4[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

exist_data = [
    ("Estatal",                    "58,300",  "19,750", "78,050"),
    ("Jurisdiccion Sanitaria No. 1","209",    "16,590", "16,799"),
    ("Jurisdiccion Sanitaria No. 2","479",    "6,660",  "7,139"),
    ("Jurisdiccion Sanitaria No. 3","914",    "2,910",  "3,824"),
    ("Jurisdiccion Sanitaria No. 4","741",    "2,230",  "2,971"),
    ("TOTAL",                       "60,643", "48,140", "108,783"),
]

alt = False
for rd in exist_data:
    row  = table4.add_row().cells
    bold = rd[0] == "TOTAL"
    bg   = 'E3F2FD' if not alt else 'FFFFFF'
    alt  = not alt
    for i, val in enumerate(rd):
        row[i].text = val
        r = row[i].paragraphs[0].runs[0]
        r.font.size = Pt(9)
        r.bold = bold
        row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if bold:
            r.font.color.rgb = RGBColor.from_string('FFFFFF')
            set_cell_bg(row[i], '1565C0')
        elif i == 0:
            set_cell_bg(row[i], bg)
set_cell_borders(table4)

doc.add_paragraph()

exist_analisis = (
    "Las existencias en el nivel estatal representan la mayor concentracion de biologico disponible "
    "(78,050 dosis), lo que permite la redistribucion oportuna a jurisdicciones de acuerdo con "
    "ritmo de aplicacion y dias restantes de campana. Las dosis totales disponibles (108,783) "
    "son suficientes para cubrir el deficit de 225,357 dosis si se complementa con la cadena "
    "de suministro federal. Se requiere gestion inmediata de reabasto para las jurisdicciones "
    "con menores existencias (JS No. 1: 16,799 dosis; JS No. 4: 2,971 dosis), "
    "priorizando la continuidad operativa de las 110 brigadas activas."
)
body(doc, exist_analisis)

# ══════════════════════════════════════════════════════════════
#  7. BRIGADAS ACTIVAS Y CAPACIDAD OPERATIVA
# ══════════════════════════════════════════════════════════════
heading(doc, "7. BRIGADAS ACTIVAS Y CAPACIDAD OPERATIVA")

brigadas_texto = (
    "Al corte del 10 de abril de 2026, se encuentran en operacion 110 brigadas activas "
    "distribuidas en el territorio estatal. Con base en el ritmo actual de apliacion "
    "(5,829 dosis/dia a nivel sectorial), cada brigada aplica en promedio 53 dosis por dia. "
    "Para alcanzar la meta estatal en los 20 dias restantes, el ritmo por brigada debe "
    "incrementarse a 102 dosis/brigada/dia, lo que requiere intensificar las actividades "
    "de barrido, casa a casa y concentracion en puntos de alta affluencia poblacional."
)
body(doc, brigadas_texto)

# tabla brigadas/capacidad
table5 = doc.add_table(rows=1, cols=3)
table5.style = 'Table Grid'
table5.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr5 = table5.rows[0].cells
for i, txt in enumerate(["INDICADOR OPERATIVO", "SITUACION ACTUAL", "REQUERIDO PARA META"]):
    hdr5[i].text = txt
    hdr5[i].paragraphs[0].runs[0].bold = True
    hdr5[i].paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(hdr5[i], '1A237E')
    hdr5[i].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string('FFFFFF')
    hdr5[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

op_data = [
    ("Brigadas activas",              "110 brigadas",           "110 brigadas"),
    ("Dias restantes de campana",    "20 dias",               "20 dias"),
    ("Dosis faltantes",              "225,357",               "225,357"),
    ("Ritmo diario sectorial",       "5,829 dosis/dia",       "11,268 dosis/dia"),
    ("Promedio por brigada/dia",     "53 dosis",              "102 dosis"),
    ("Ritmo semanal requerido",      "40,799 dos/sem (actual)","52,261 dos/sem"),
]

alt = False
for rd in op_data:
    row = table5.add_row().cells
    bg  = 'E8EAF6' if not alt else 'FFFFFF'
    alt = not alt
    for i, val in enumerate(rd):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(9)
        row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i == 0:
            row[i].paragraphs[0].runs[0].bold = True
            set_cell_bg(row[i], bg)
set_cell_borders(table5)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  8. CONCLUSIONES Y LINEAS DE ACCION
# ══════════════════════════════════════════════════════════════
heading(doc, "8. CONCLUSIONES Y LINEAS DE ACCION")

conclusiones = [
    ("Brecha de cobertura critica:", 
     "La campaña sectorial acumula 297,251 dosis (65.0% de la meta semanas 1-8). "
     "Con 20 dias restantes se requiere casi duplicar el ritmo diario actual."),
    ("SSA con mayor deficit:",
     "La SSA presenta el avance mas bajo (36.1%), con 53,183 dosis aplicadas de meta "
     "acumulada de 147,500. Es prioritario reforzar las acciones operativas de las "
     "jurisdicciones sanitarias."),
    ("IMSS lidera el avance:",
     "Con 88.1% de meta acumulada (213,901 dosis de 242,663), el IMSS mantiene el "
     "ritmo mas alto y es la institucion con mejor desempeno en la campana."),
    ("Existencias suficientes a nivel estatal:",
     "El total disponible de 108,783 dosis (SRP + SR) permite la continuidad de la "
     "campana si se gestiona una redistribucion oportuna desde el nivel estatal."),
    ("Intensificacion operativa urgente:",
     "Las 110 brigadas activas deben incrementar su rendimiento de 97 a 102 dosis/dia. "
     "Se recomienda activar estrategias de barrido intensivo, jornadas sabatinas y "
     "acciones intersectoriales para cerrar la brecha en las semanas 9 y 10."),
]

for titulo, texto in conclusiones:
    p = doc.add_paragraph(style='List Bullet')
    run_t = p.add_run(titulo + " ")
    run_t.bold = True
    run_t.font.size = Pt(10)
    run_c = p.add_run(texto)
    run_c.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

# ── Firma
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("Durango, Dgo., a 10 de abril de 2026")
run.font.size = Pt(9)
run.italic    = True

# ══════════════════════════════════════════════════════════════
#  Guardar
# ══════════════════════════════════════════════════════════════
out = r'C:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\CAMPAÑA SARAMPIÓN 10 SEMANAS\Informe_Ejecutivo_Sarampion_110brigadas_10-04-2026.docx'
doc.save(out)
print(f"Guardado en: {out}")
