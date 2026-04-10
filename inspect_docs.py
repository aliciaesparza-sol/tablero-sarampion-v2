import docx
import os

def read_docx(file_path):
    if not os.path.exists(file_path):
        return f"File not found: {file_path}"
    try:
        doc = docx.Document(file_path)
        content = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                content.append(f"P{i}: {para.text}")
        
        # Check tables too
        for i, table in enumerate(doc.tables):
            content.append(f"\n--- Table {i} ---")
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                content.append(" | ".join(row_text))
        
        return "\n".join(content)
    except Exception as e:
        return str(e)

model_path = r"c:\Users\aicil\OneDrive\Escritorio\PVU\CAMPAÑA DE VACUNACION MASIVA INTERINSTITUCIONAL DICIEMBRE 2025\MACROCENTROS UNIVERSITARIOS 25 MARZO 2026\FICHA TECNICA MACROCENTROS DE VACUNACIÓN UNIVERSITARIOS 25 MARZO 2026.docx"
letterhead_path = r"c:\Users\aicil\OneDrive\Escritorio\hoja MEMBRETADA GIGANTE2026_carta.docx"

print("--- MODEL CONTENT ---")
print(read_docx(model_path))
print("\n--- LETTERHEAD INFO ---")
# For letterhead, we mostly care if it has headers/footers or images
doc_lh = docx.Document(letterhead_path)
print(f"Paragraphs: {len(doc_lh.paragraphs)}")
print(f"Sections: {len(doc_lh.sections)}")
for i, section in enumerate(doc_lh.sections):
    print(f"Section {i} Header: {len(section.header.paragraphs)} paras")
    print(f"Section {i} Footer: {len(section.footer.paragraphs)} paras")
