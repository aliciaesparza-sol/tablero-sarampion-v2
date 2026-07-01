import os
import sys

# Ensure python-docx is installed
try:
    import docx
except ImportError:
    import subprocess, sys
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    import docx

from docx import Document
from docx.shared import Inches

# Paths
md_path = r"C:/Users/aicil/.gemini/antigravity-ide/scratch/vaccination_report.md"
output_docx = r"C:/Users/aicil/.gemini/antigravity-ide/scratch/vaccination_report.docx"

# Create a new Word document
doc = Document()

# Read markdown content
with open(md_path, 'r', encoding='utf-8') as f:
    lines = [line.rstrip('\n') for line in f]

# Simple parser for headings and tables
i = 0
while i < len(lines):
    line = lines[i]
    if line.startswith('# '):
        # Title
        doc.add_heading(line[2:], level=0)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=2)
    elif line.startswith('|'):
        # Table start: collect rows until a blank line or non-table line
        table_rows = []
        while i < len(lines) and lines[i].startswith('|'):
            table_rows.append(lines[i])
            i += 1
        # The first row is header, second is separator
        if len(table_rows) >= 2:
            header = [h.strip() for h in table_rows[0].split('|')[1:-1]]
            # Create table with number of rows = data rows + 1 header
            data_rows = table_rows[2:]
            table = doc.add_table(rows=1 + len(data_rows), cols=len(header))
            # Header row
            hdr_cells = table.rows[0].cells
            for idx, col_name in enumerate(header):
                hdr_cells[idx].text = col_name
            # Data rows
            for r_idx, row_line in enumerate(data_rows):
                cells = [c.strip() for c in row_line.split('|')[1:-1]]
                row_cells = table.rows[r_idx + 1].cells
                for c_idx, cell_text in enumerate(cells):
                    row_cells[c_idx].text = cell_text
        continue  # continue to avoid i increment at end
    else:
        if line.strip():
            doc.add_paragraph(line)
    i += 1

# Save the document
doc.save(output_docx)
print('Word report generated at', output_docx)
