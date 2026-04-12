import docx
import os
import json

file_path = r"c:\Users\aicil\.gemini\antigravity\scratch\report.docx"
data = {"paragraphs": [], "tables": []}

try:
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
    else:
        doc = docx.Document(file_path)
        
        for para in doc.paragraphs:
            if para.text.strip():
                data["paragraphs"].append(para.text.strip())
                
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    text = cell.text.strip()
                    if not row_data or text != row_data[-1]:
                        row_data.append(text)
                table_data.append(row_data)
            data["tables"].append(table_data)
            
        with open(r"c:\Users\aicil\.gemini\antigravity\scratch\all_data.json", "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4, ensure_ascii=False)
        print("Data saved to all_data.json")
except Exception as e:
    print(f"Error: {e}")
