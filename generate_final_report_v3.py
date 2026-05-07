import json
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load data
with open("locality_doses_v2.json", "r") as f:
    locality_data = json.load(f)
with open("mezquital_pop_conapo.json", "r") as f:
    pop_data = json.load(f)
with open("mezquital_doses_age_v3.json", "r") as f:
    doses_age_data = json.load(f)

# Base report path
base_report = r"C:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\mezquital\INFORMES\Informe_Vacunacion_Mezquital_2026_Final_Heatmap.docx"
output_report = r"C:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\mezquital\INFORMES\Informe_Vacunacion_Mezquital_2026_Completo_v2.docx"

try:
    doc = docx.Document(base_report)
    
    # --- SECTION: LOCALITY DOSES ---
    doc.add_page_break()
    doc.add_heading('Detalle de Dosis Aplicadas por Localidad (Mezquital 2026)', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Localidad'
    hdr_cells[1].text = 'Dosis Aplicadas'
    hdr_cells[2].text = 'Población (INEGI)'
    hdr_cells[3].text = 'Alcance (%)'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    for item in locality_data:
        row_cells = table.add_row().cells
        row_cells[0].text = item['Localidad']
        row_cells[1].text = f"{int(item['Doses']):,}"
        row_cells[2].text = f"{int(item['Population']):,}"
        val = item['Reach']
        if isinstance(val, (int, float)):
            # Reach in Excel is already multiplied by 100 in the script but stored as number
            # Wait, add_coverage.py did: alcance = (doses / pob) * 100; df.at[i, 83] = round(alcance, 2)
            # So it's already a percentage value (e.g. 160.61)
            row_cells[3].text = f"{val:.2f}%"
        else:
            row_cells[3].text = str(val)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Fuente: Concentrado Vacunación Mezquital 2026 / INEGI ITER 2020.")
    run.font.size = Pt(8)
    run.italic = True

    # --- SECTION: COVERAGE BY AGE GROUP ---
    doc.add_page_break()
    doc.add_heading('Cobertura de Vacunación - Municipio de Mezquital (CONAPO 2026)', level=1)
    
    table_age = doc.add_table(rows=1, cols=6)
    table_age.style = 'Table Grid'
    hdr_cells = table_age.rows[0].cells
    hdr_cells[0].text = 'Grupo de Edad'
    hdr_cells[1].text = 'Población (CONAPO)'
    hdr_cells[2].text = 'Dosis 2025'
    hdr_cells[3].text = 'Dosis 2026'
    hdr_cells[4].text = 'Total 25-26'
    hdr_cells[5].text = 'Cobertura (%)'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    for label in pop_data.keys():
        pop = pop_data[label]
        d_info = doses_age_data.get(label, {"2025": 0, "2026": 0, "Total": 0})
        total_doses = d_info["Total"]
        coverage = (total_doses / pop * 100) if pop > 0 else 0
        
        row_cells = table_age.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = f"{int(pop):,}"
        row_cells[2].text = f"{int(d_info['2025']):,}"
        row_cells[3].text = f"{int(d_info['2026']):,}"
        row_cells[4].text = f"{int(total_doses):,}"
        row_cells[5].text = f"{coverage:.2f}%"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Fuente: SIS/CeNSIA, consultado el 04 de mayo de 2026; Proyecciones de Población CONAPO 2026.")
    run.font.size = Pt(8)
    run.italic = True
    
    doc.save(output_report)
    print(f"Final report saved to {output_report}")

except Exception as e:
    print(f"Error: {e}")
