
from docx import Document

doc = Document(r'C:\Users\aicil\.gemini\antigravity\scratch\informe_original.docx')

print('=== PÁRRAFOS ===')
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f'[{i}] ESTILO={p.style.name!r} | {p.text[:300]}')

print()
print('=== TABLAS ===')
for ti, t in enumerate(doc.tables):
    print(f'\n--- TABLA {ti} ({len(t.rows)} filas x {len(t.columns)} cols) ---')
    for ri, row in enumerate(t.rows):
        cells = [c.text.strip()[:60] for c in row.cells]
        print(f'  F{ri}: {cells}')
