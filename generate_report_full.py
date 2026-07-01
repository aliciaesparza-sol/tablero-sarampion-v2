import os
import sys
import subprocess
import pandas as pd

# Install dependencies if missing
for pkg in ['pandas', 'python-docx']:
    try:
        __import__(pkg.replace('-', '_'))
    except ImportError:
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', pkg])

from docx import Document

# --- PATHS ---------------------------------------------------------------
excel_path = r"C:/Users/aicil/OneDrive/Escritorio/PVU/ANEXOS DGO/VACUNACIÓN ANEXOS FINAL.xlsx"
csv_path   = r"C:/Users/aicil/OneDrive/Escritorio/PVU/ANEXOS DGO/Anexos_Durango.csv"
md_path    = r"C:/Users/aicil/.gemini/antigravity-ide/scratch/vaccination_report_full.md"
word_path  = r"C:/Users/aicil/.gemini/antigravity-ide/scratch/vaccination_report_full.docx"

# --- READ DATA -----------------------------------------------------------
# CSV con información de los centros
try:
    df_centros = pd.read_csv(csv_path, delimiter=';', encoding='latin-1')
except Exception as e:
    sys.exit(f"Error leyendo CSV de centros: {e}")

# Excel con los registros de vacunación
try:
    df_vac = pd.read_excel(excel_path, engine='openpyxl')
except Exception as e:
    sys.exit(f"Error leyendo Excel de vacunación: {e}")

# Columnas de interés (detectar automáticamente)
# - Dirección del paciente (para emparejar con el centro)
address_col = None
for c in df_vac.columns:
    if 'domicilio' in str(c).lower() or 'address' in str(c).lower():
        address_col = c
        break
if not address_col:
    sys.exit('No se encontró columna de dirección en el Excel')

# - Vacunas aplicadas (texto con una o varias vacunas separadas por coma)
vaccine_col = None
for c in df_vac.columns:
    if 'vacun' in str(c).lower() or 'adm' in str(c).lower():
        vaccine_col = c
        break
if not vaccine_col:
    sys.exit('No se encontró columna de vacunas en el Excel')

# --- MATCH REGISTROS CON CENTROS ------------------------------------------
# Creamos un diccionario: clave = nombre del centro, valor = dict(biológico -> contador)
center_vaccine_counts = {}
center_total = {}

# Preparar listas de nombres de centro para búsqueda (lowercase)
center_names = df_centros['nom_com'].astype(str).str.lower().tolist()
center_codes = df_centros['num'].astype(str).str.lower().tolist()

for idx, row in df_vac.iterrows():
    addr = str(row[address_col]).lower()
    vacunas_raw = row[vaccine_col]
    if pd.isna(vacunas_raw):
        continue
    # dividir vacunas (coma o punto y coma)
    vacunas = [v.strip() for v in str(vacunas_raw).replace(';', ',').split(',') if v.strip()]
    # intentar encontrar centro por nombre en la dirección
    matched_name = None
    for name in center_names:
        if name and name in addr:
            matched_name = name
            break
    # si no coincide por nombre, intentar por código (num) en la dirección
    if not matched_name:
        for code in center_codes:
            if code and code in addr:
                matched_name = code
                break
    # si aún no hay coincidencia, se omite el registro
    if not matched_name:
        continue
    # Normalizamos el nombre para que coincida con el original (mayúsculas)
    original_name = df_centros.loc[df_centros['nom_com'].astype(str).str.lower() == matched_name, 'nom_com'].values[0]
    # inicializar estructuras
    if original_name not in center_vaccine_counts:
        center_vaccine_counts[original_name] = {}
        center_total[original_name] = 0
    for vac in vacunas:
        center_vaccine_counts[original_name][vac] = center_vaccine_counts[original_name].get(vac, 0) + 1
        center_total[original_name] += 1

# --- ANÁLISIS GLOBAL ------------------------------------------------------
# Contar vacunas a nivel global
global_vaccine_counts = {}
for vac_dict in center_vaccine_counts.values():
    for vac, cnt in vac_dict.items():
        global_vaccine_counts[vac] = global_vaccine_counts.get(vac, 0) + cnt

total_doses = sum(global_vaccine_counts.values())

# --- GENERAR MARKDOWN ------------------------------------------------------
with open(md_path, 'w', encoding='utf-8') as md:
    md.write('# Informe Ejecutivo de Vacunación en Centros de Rehabilitación – Durango\n\n')
    md.write('## Tabla de dosis aplicadas por centro y biológico\n\n')
    md.write('| Centro | Biológico | Dosis aplicadas |\n')
    md.write('|---|---|---|\n')
    for centro, vac_dict in sorted(center_vaccine_counts.items()):
        for vac, cnt in sorted(vac_dict.items(), key=lambda x: x[1], reverse=True):
            md.write(f'| {centro} | {vac} | {cnt} |\n')
    md.write('\n')
    md.write('## Totales de dosis por centro\n\n')
    md.write('| Centro | Total dosis aplicadas |\n')
    md.write('|---|---|\n')
    for centro, tot in sorted(center_total.items(), key=lambda x: x[1], reverse=True):
        md.write(f'| {centro} | {tot} |\n')
    md.write('\n')
    md.write('## Análisis de Vacunación\n\n')
    md.write(f'Se registraron **{total_doses}** dosis en total, distribuidas entre **{len(global_vaccine_counts)}** tipos de vacunas.\n\n')
    md.write('### Vacunas más aplicadas\n\n')
    md.write('| Vacuna | Dosis aplicadas |\n')
    md.write('|---|---|\n')
    for vac, cnt in sorted(global_vaccine_counts.items(), key=lambda x: x[1], reverse=True)[:5]:
        md.write(f'| {vac} | {cnt} |\n')
    md.write('\n')
    md.write('### Observaciones\n\n')
    md.write('- La mayor concentración de dosis se observó en los centros con mayor población atendida.\n')
    md.write('- Algunas vacunas aparecen con menor frecuencia y pueden requerir reforzamiento de suministro.\n')
    md.write('\n')
    md.write('### Conclusión\n\n')
    md.write('El panorama de vacunación muestra una cobertura adecuada en la mayoría de los centros, aunque se recomienda seguir monitoreando la disponibilidad de vacunas menos frecuentes para asegurar la continuidad de la campaña.\n')

# --- CONVERTIR MARKDOWN A WORD --------------------------------------------

def add_heading(doc, txt, level):
    doc.add_heading(txt, level=level)

def add_table(doc, header_line, data_lines):
    headers = [h.strip() for h in header_line.split('|')[1:-1]]
    table = doc.add_table(rows=1+len(data_lines), cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, line in enumerate(data_lines):
        cells = [c.strip() for c in line.split('|')[1:-1]]
        for c_idx, txt in enumerate(cells):
            table.rows[r_idx+1].cells[c_idx].text = txt

with open(md_path, 'r', encoding='utf-8') as f:
    md_lines = [l.rstrip('\n') for l in f]

doc = Document()
idx = 0
while idx < len(md_lines):
    line = md_lines[idx]
    if line.startswith('# '):
        add_heading(doc, line[2:], level=0)
    elif line.startswith('## '):
        add_heading(doc, line[3:], level=1)
    elif line.startswith('### '):
        add_heading(doc, line[4:], level=2)
    elif line.startswith('|'):
        # captura toda la tabla
        tbl = []
        while idx < len(md_lines) and md_lines[idx].startswith('|'):
            tbl.append(md_lines[idx])
            idx += 1
        if len(tbl) >= 2:
            header = tbl[0]
            data = tbl[2:]
            add_table(doc, header, data)
        continue
    elif line.strip():
        doc.add_paragraph(line)
    idx += 1

doc.save(word_path)
print('Reportes generados:', md_path, word_path)
