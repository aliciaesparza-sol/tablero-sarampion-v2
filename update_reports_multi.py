import pandas as pd
import os
from pathlib import Path
from docx import Document

# Paths (adjust if needed)
excel_path = r'C:/Users/aicil/OneDrive/Escritorio/PVU/ANEXOS DGO/VACUNACIÓN ANEXOS FINAL.xlsx'
md_path = r'C:/Users/aicil/OneDrive/Escritorio/PVU/ANEXOS DGO/vaccination_report_multi.md'
# Use a different docx name to avoid permission issues
docx_path = r'C:/Users/aicil/OneDrive/Escritorio/PVU/ANEXOS DGO/vaccination_report_multi.docx'

# Load all sheets – each sheet corresponds to a rehabilitation centre
all_sheets = pd.read_excel(excel_path, sheet_name=None, engine='openpyxl')

# List to collect processed rows
rows = []
for sheet_name, df in all_sheets.items():
    # Clean column names
    df.columns = [c.strip() for c in df.columns]
    # Identify vaccine column (contains "admnistr" or "vacun")
    vac_col = None
    for col in df.columns:
        if 'admnistr' in col.lower() or 'vacun' in col.lower():
            vac_col = col
            break
    if not vac_col:
        continue  # skip sheet if no vaccine column
    # Split vaccines (comma‑separated) and explode
    df['Vacunas'] = df[vac_col].astype(str).str.split(',')
    exploded = df.explode('Vacunas')
    exploded['Vacunas'] = exploded['Vacunas'].str.strip()
    exploded['Centro'] = sheet_name  # sheet name is the centre name
    rows.append(exploded[['Centro', 'Vacunas']])

if not rows:
    raise RuntimeError('No data extracted from any sheet')

full_df = pd.concat(rows, ignore_index=True)
# Count doses per centre and vaccine
center_vaccine_counts = full_df.groupby(['Centro', 'Vacunas']).size().reset_index(name='Dosis')
# Totals per centre
center_totals = full_df.groupby('Centro').size().reset_index(name='TotalDosis')
# Global total
global_total = len(full_df)

# Build markdown report
md_lines = []
md_lines.append('# Informe de Dosis por Centro (hojas del Excel)')
md_lines.append('')
md_lines.append('## Tabla de dosis por centro y biológico')
md_lines.append('| Centro | Biológico | Dosis aplicadas |')
md_lines.append('|---|---|---|')
for _, row in center_vaccine_counts.iterrows():
    md_lines.append(f"| {row['Centro']} | {row['Vacunas']} | {int(row['Dosis'])} |")
md_lines.append('')
md_lines.append('## Totales por centro')
md_lines.append('| Centro | Total dosis aplicadas |')
md_lines.append('|---|---|')
for _, row in center_totals.iterrows():
    md_lines.append(f"| {row['Centro']} | {int(row['TotalDosis'])} |")
md_lines.append('')
md_lines.append('## Análisis global')
md_lines.append(f'Se registraron **{global_total}** dosis en total, distribuidas entre **{center_vaccine_counts["Vacunas"].nunique()}** tipos de vacunas.')

with open(md_path, 'w', encoding='utf-8') as f:
    f.write('\n'.join(md_lines))

# Create Word document with same tables
doc = Document()
doc.add_heading('Informe de Dosis por Centro (hojas del Excel)', 0)

# Table centre‑vaccine
doc.add_heading('Dosis por centro y biológico', level=2)
table = doc.add_table(rows=1, cols=3)
hdr = table.rows[0].cells
hdr[0].text = 'Centro'
hdr[1].text = 'Biológico'
hdr[2].text = 'Dosis aplicadas'
for _, row in center_vaccine_counts.iterrows():
    cells = table.add_row().cells
    cells[0].text = str(row['Centro'])
    cells[1].text = str(row['Vacunas'])
    cells[2].text = str(int(row['Dosis']))

# Totals per centre
doc.add_heading('Totales por centro', level=2)
table2 = doc.add_table(rows=1, cols=2)
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Centro'
hdr2[1].text = 'Total dosis aplicadas'
for _, row in center_totals.iterrows():
    cells = table2.add_row().cells
    cells[0].text = str(row['Centro'])
    cells[1].text = str(int(row['TotalDosis']))

# Global analysis
doc.add_heading('Análisis global', level=2)
doc.add_paragraph(f'Se registraron {global_total} dosis en total.')

doc.save(docx_path)
print('Reportes multi-hoja generados con exito')
