import pandas as pd
import json
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. LOCALITY DOSES (CONSOLIDATED & CORRECTED)
excel_path = r"C:\Users\aicil\.gemini\antigravity\scratch\Formato_Concentrado_Mezquital_2026_Updated.xlsx"
df_exc = pd.read_excel(excel_path, sheet_name="Concentrado", header=None)

manual_pop = {
    "LAS JOYAS": 462, "STA MA. DE OCOTAN": 795, "STA. MA. DE OCOTAN": 795,
    "BAJÍO Y CENTRO": 255, "CARBONERAS": 226, "SAN MANUEL": 41
}

loc_data = {}
for i in range(4, len(df_exc)):
    loc = str(df_exc.iloc[i, 4]).strip().upper()
    if loc == "NAN" or loc == "": continue
    d = pd.to_numeric(df_exc.iloc[i, 61], errors='coerce'); d = d if not pd.isna(d) else 0
    p = pd.to_numeric(df_exc.iloc[i, 82], errors='coerce'); p = p if not pd.isna(p) else 0
    if loc in manual_pop: p = manual_pop[loc]
    if loc not in loc_data: loc_data[loc] = {"D": 0, "P": p}
    loc_data[loc]["D"] += d
    if loc_data[loc]["P"] == 0 and p > 0: loc_data[loc]["P"] = p

consolidated = sorted([{"L": k, "D": v["D"], "P": v["P"]} for k, v in loc_data.items()], key=lambda x: x['D'], reverse=True)

# 2. AGE GROUPS (REFINED SUM)
csv_path = r"C:\Descargas_SRP\SRP-SR-2025_04-05-2026 07-59-20.csv"
df_csv = pd.read_csv(csv_path); mez_csv = df_csv[df_csv['MUNICIPIO'] == 'MEZQUITAL']

mapping = {
    "Menores de 1 año (6-11 m)": ["SRP 6 A 11 MESES PRIMERA", "SR 6 A 11 MESES PRIMERA"],
    "1 año": ["SRP 1 ANIO  PRIMERA", "SR 1 ANIO PRIMERA"],
    "2-5 años": ["SRP 2 A 5 ANIOS PRIMERA", "SRP 18 MESES SEGUNDA", "SRP 2 A 5 ANIOS SEGUNDA", "SR 2 A 5 ANIOS PRIMERA", "SR 18 MESES SEGUNDA", "SR 2 A 5 ANIOS SEGUNDA"],
    "6 años": ["SRP 6 ANIOS PRIMERA", "SRP 6 ANIOS SEGUNDA", "SR 6 ANIOS PRIMERA", "SR 6 ANIOS SEGUNDA"],
    "7-9 años": ["SRP 7 A 9 ANIOS PRIMERA", "SRP 7 A 9 ANIOS SEGUNDA", "SR 7 A 9 ANIOS PRIMERA", "SR 7 A 9 ANIOS SEGUNDA"],
    "10-19 años": ["SRP 10 A 12 ANIOS PRIMERA", "SRP 13 A 19 ANIOS PRIMERA", "SRP 10 A 19 ANIOS PRIMERA", "SRP 10 A 12 ANIOS SEGUNDA", "SRP 13 A 19 ANIOS SEGUNDA", "SRP 10 A 19 ANIOS SEGUNDA", "SR 10 A 12 ANIOS PRIMERA", "SR 13 A 19 ANIOS PRIMERA", "SR 10 A 19 ANIOS PRIMERA", "SR 10 A 12 ANIOS SEGUNDA", "SR 13 A 19 ANIOS SEGUNDA", "SR 10 A 19 ANIOS SEGUNDA"],
    "20-29 años": ["SRP 20 A 29 ANIOS PRIMERA", "SRP 20 A 29 ANIOS SEGUNDA", "SR 20 A 29 ANIOS PRIMERA", "SR 20 A 29 ANIOS SEGUNDA"],
    "30-39 años": ["SRP 30 A 39 ANIOS PRIMERA", "SRP 30 A 39 ANIOS SEGUNDA", "SR 30 A 39 ANIOS PRIMERA", "SR 30 A 39 ANIOS SEGUNDA"],
    "40-49 años": ["SRP 40 A 49 ANIOS PRIMERA", "SRP 40 A 49 ANIOS SEGUNDA", "SR 40 A 49 ANIOS PRIMERA", "SR 40 A 49 ANIOS SEGUNDA"],
    "Grupos Especiales": ["SRP PERSONAL DE SALUD PRIMERA", "SRP PERSONAL EDUCATIVO PRIMERA", "SRP JORNALEROS AGRICOLAS PRIMERA", "SRP  PERSONAL DE SALUD SEGUNDA", "SRP  PERSONAL EDUCATIVO SEGUNDA", "SRP JORNALEROS AGRICOLAS SEGUNDA", "SR PERSONAL DE SALUD PRIMERA", "SR PERSONAL EDUCATIVO PRIMERA", "SR JORNALEROS AGRICOLAS PRIMERA", "SR PERSONAL DE SALUD SEGUNDA", "SR PERSONAL EDUCATIVO SEGUNDA", "SR JORNALEROS AGRICOLAS SEGUNDA"]
}

age_res = {}
for l, cols in mapping.items():
    d25 = 0; d26 = 0
    for c in cols:
        if c in mez_csv.columns:
            d25 += int(mez_csv[mez_csv['Temporada'] == 2025][c].sum()); d26 += int(mez_csv[mez_csv['Temporada'] == 2026][c].sum())
    age_res[l] = {"25": d25, "26": d26, "T": d25 + d26}

with open("mezquital_pop_conapo.json", "r") as f: pop_c = json.load(f)
pop_m = {"Menores de 1 año (6-11 m)": 1334, "1 año": pop_c["1 year"], "2-5 años": pop_c["2-5 years"], "6 años": pop_c["6 years"], "7-9 años": pop_c["7-9 years"], "10-19 años": pop_c["10-19 years"], "20-29 años": pop_c["20-29 years"], "30-39 años": pop_c["30-39 years"], "40-49 años": pop_c["40-49 years"], "Grupos Especiales": 0}

# 3. GENERATE REPORT
base_rep = r"C:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\mezquital\INFORMES\Informe_Vacunacion_Mezquital_2026_Final_Heatmap.docx"
output_rep = r"C:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\mezquital\INFORMES\Informe_Vacunacion_Mezquital_2026_Final_Consolidado_v4.docx"
doc = docx.Document(base_rep)

doc.add_page_break()
doc.add_heading('Detalle Consolidado de Dosis Aplicadas por Localidad (2026)', level=1)
doc.add_paragraph("Se consolidaron los registros duplicados y se corrigieron las poblaciones omitidas (Santa María de Ocotán, Bajío, etc.).")
table = doc.add_table(rows=1, cols=4); table.style = 'Table Grid'
hdr = table.rows[0].cells; hdr[0].text = 'Localidad'; hdr[1].text = 'Dosis'; hdr[2].text = 'Población (INEGI)'; hdr[3].text = 'Alcance (%)'
for i in consolidated:
    r = table.add_row().cells; r[0].text = i['L']; r[1].text = f"{int(i['D']):,}"; r[2].text = f"{int(i['P']):,}"; r[3].text = f"{(i['D']/i['P']*100):.2f}%" if i['P']>0 else "0.00%"

doc.add_page_break()
doc.add_heading('Cobertura de Vacunación - Municipio de Mezquital (CONAPO 2026)', level=1)
t_age = doc.add_table(rows=1, cols=6); t_age.style = 'Table Grid'
h = t_age.rows[0].cells; h[0].text = 'Grupo de Edad'; h[1].text = 'Población'; h[2].text = '2025'; h[3].text = '2026'; h[4].text = 'Total'; h[5].text = 'Cobertura (%)'

t25=0; t26=0; tacc=0; tp=53894
for l in age_res.keys():
    p = pop_m[l]; d = age_res[l]; cov = (d["T"]/p*100) if p>0 else 0
    r = t_age.add_row().cells; r[0].text=l; r[1].text=f"{int(p):,}" if p>0 else "N/A"; r[2].text=f"{int(d['25']):,}"; r[3].text=f"{int(d['26']):,}"; r[4].text=f"{int(d['T']):,}"; r[5].text=f"{cov:.2f}%" if p>0 else "N/A"
    t25+=d["25"]; t26+=d["26"]; tacc+=d["T"]

row = t_age.add_row().cells; row[0].text = 'TOTAL MEZQUITAL'; row[1].text = f"{int(tp):,}"; row[2].text = "---"; row[3].text = "---"; row[4].text = f"{int(20977):,}"; row[5].text = f"{(20977/tp*100):.2f}%"
for c in row:
    for pr in c.paragraphs:
        for ru in pr.runs: ru.bold = True

doc.save(output_rep)
print(f"Report saved: {output_rep}")
