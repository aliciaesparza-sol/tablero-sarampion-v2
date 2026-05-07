import pandas as pd
import json

# 1. DOSES BY AGE (CSV)
csv_path = r"C:\Descargas_SRP\SRP-SR-2025_04-05-2026 07-59-20.csv"
df_csv = pd.read_csv(csv_path)
mez_csv = df_csv[df_csv['MUNICIPIO'] == 'MEZQUITAL']

mapping = {
    "1 año": ["SRP 1 ANIO  PRIMERA", "SR 1 ANIO PRIMERA"],
    "2-5 años": ["SRP 2 A 5 ANIOS PRIMERA", "SRP 18 MESES SEGUNDA", "SRP 2 A 5 ANIOS SEGUNDA", "SR 2 A 5 ANIOS PRIMERA", "SR 18 MESES SEGUNDA", "SR 2 A 5 ANIOS SEGUNDA"],
    "6 años": ["SRP 6 ANIOS PRIMERA", "SRP 6 ANIOS SEGUNDA", "SR 6 ANIOS PRIMERA", "SR 6 ANIOS SEGUNDA"],
    "7-9 años": ["SRP 7 A 9 ANIOS PRIMERA", "SRP 7 A 9 ANIOS SEGUNDA", "SR 7 A 9 ANIOS PRIMERA", "SR 7 A 9 ANIOS SEGUNDA"],
    "10-19 años": [
        "SRP 10 A 12 ANIOS PRIMERA", "SRP 13 A 19 ANIOS PRIMERA", "SRP 10 A 19 ANIOS PRIMERA",
        "SRP 10 A 12 ANIOS SEGUNDA", "SRP 13 A 19 ANIOS SEGUNDA", "SRP 10 A 19 ANIOS SEGUNDA",
        "SR 10 A 12 ANIOS PRIMERA", "SR 13 A 19 ANIOS PRIMERA", "SR 10 A 19 ANIOS PRIMERA",
        "SR 10 A 12 ANIOS SEGUNDA", "SR 13 A 19 ANIOS SEGUNDA", "SR 10 A 19 ANIOS SEGUNDA"
    ],
    "20-29 años": ["SRP 20 A 29 ANIOS PRIMERA", "SRP 20 A 29 ANIOS SEGUNDA", "SR 20 A 29 ANIOS PRIMERA", "SR 20 A 29 ANIOS SEGUNDA"],
    "30-39 años": ["SRP 30 A 39 ANIOS PRIMERA", "SRP 30 A 39 ANIOS SEGUNDA", "SR 30 A 39 ANIOS PRIMERA", "SR 30 A 39 ANIOS SEGUNDA"],
    "40-49 años": ["SRP 40 A 49 ANIOS PRIMERA", "SRP 40 A 49 ANIOS SEGUNDA", "SR 40 A 49 ANIOS PRIMERA", "SR 40 A 49 ANIOS SEGUNDA"]
}

age_results = {}
for label, cols in mapping.items():
    d25 = 0
    d26 = 0
    for col in cols:
        if col in mez_csv.columns:
            d25 += int(mez_csv[mez_csv['Temporada'] == 2025][col].sum())
            d26 += int(mez_csv[mez_csv['Temporada'] == 2026][col].sum())
    age_results[label] = {"2025": d25, "2026": d26, "Total": d25 + d26}

# 2. LOCALITY DOSES (EXCEL) - CONSOLIDATED
excel_path = r"C:\Users\aicil\.gemini\antigravity\scratch\Formato_Concentrado_Mezquital_2026_Updated.xlsx"
df_exc = pd.read_excel(excel_path, sheet_name="Concentrado", header=None)

loc_data = {}
for i in range(4, len(df_exc)):
    loc = str(df_exc.iloc[i, 4]).strip().upper()
    if loc == "NAN" or loc == "": continue
    
    doses = pd.to_numeric(df_exc.iloc[i, 61], errors='coerce')
    if pd.isna(doses): doses = 0
    
    pop = pd.to_numeric(df_exc.iloc[i, 82], errors='coerce')
    if pd.isna(pop): pop = 0
    
    if loc not in loc_data:
        loc_data[loc] = {"Doses": 0, "Population": 0}
    
    loc_data[loc]["Doses"] += doses
    # Population might be the same for all entries of the same locality, or different.
    # We take the maximum population found for that name (usually the most complete record)
    loc_data[loc]["Population"] = max(loc_data[loc]["Population"], pop)

consolidated_locs = []
for loc, vals in loc_data.items():
    reach = (vals["Doses"] / vals["Population"]) * 100 if vals["Population"] > 0 else 0
    consolidated_locs.append({
        "Localidad": loc,
        "Doses": vals["Doses"],
        "Population": vals["Population"],
        "Reach": reach
    })
consolidated_locs.sort(key=lambda x: x['Doses'], reverse=True)

# 3. POPULATION (CONAPO)
with open("mezquital_pop_conapo.json", "r") as f:
    pop_conapo = json.load(f)
# Map keys to match age_results
pop_map = {
    "1 año": pop_conapo["1 year"],
    "2-5 años": pop_conapo["2-5 years"],
    "6 años": pop_conapo["6 years"],
    "7-9 años": pop_conapo["7-9 years"],
    "10-19 años": pop_conapo["10-19 years"],
    "20-29 años": pop_conapo["20-29 years"],
    "30-39 años": pop_conapo["30-39 years"],
    "40-49 años": pop_conapo["40-49 years"]
}

# 4. GENERATE REPORT
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

base_report = r"C:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\mezquital\INFORMES\Informe_Vacunacion_Mezquital_2026_Final_Heatmap.docx"
output_report = r"C:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\mezquital\INFORMES\Informe_Vacunacion_Mezquital_2026_Final_Consolidado.docx"

doc = docx.Document(base_report)

# LOCALITIES
doc.add_page_break()
doc.add_heading('Dosis Aplicadas por Localidad - Municipio de Mezquital (2026)', level=1)
p = doc.add_paragraph("Resumen consolidado por nombre de localidad según Formato Concentrado.")
p.font.size = Pt(10)

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Localidad'
hdr[1].text = 'Dosis Aplicadas'
hdr[2].text = 'Población (INEGI)'
hdr[3].text = 'Alcance (%)'

for item in consolidated_locs:
    row = table.add_row().cells
    row[0].text = item['Localidad']
    row[1].text = f"{int(item['Doses']):,}"
    row[2].text = f"{int(item['Population']):,}"
    row[3].text = f"{item['Reach']:.2f}%"

# AGE GROUPS
doc.add_page_break()
doc.add_heading('Cobertura de Vacunación por Grupo de Edad - Municipio de Mezquital (CONAPO 2026)', level=1)

table_age = doc.add_table(rows=1, cols=6)
table_age.style = 'Table Grid'
hdr = table_age.rows[0].cells
hdr[0].text = 'Grupo de Edad'
hdr[1].text = 'Población (CONAPO)'
hdr[2].text = 'Dosis 2025'
hdr[3].text = 'Dosis 2026'
hdr[4].text = 'Total 25-26'
hdr[5].text = 'Cobertura (%)'

for label in age_results.keys():
    pop = pop_map[label]
    d_info = age_results[label]
    cov = (d_info["Total"] / pop * 100) if pop > 0 else 0
    
    row = table_age.add_row().cells
    row[0].text = label
    row[1].text = f"{int(pop):,}"
    row[2].text = f"{int(d_info['2025']):,}"
    row[3].text = f"{int(d_info['2026']):,}"
    row[4].text = f"{int(d_info['Total']):,}"
    row[5].text = f"{cov:.2f}%"

doc.save(output_report)
print(f"Report saved: {output_report}")
