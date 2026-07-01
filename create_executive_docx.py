import pathlib, re
from docx import Document
from docx.shared import Inches

md_path = pathlib.Path(r"C:/Users/aicil/OneDrive/Escritorio/PVU/ANEXOS DGO/executive_report.md")
output_path = pathlib.Path(r"C:/Users/aicil/OneDrive/Escritorio/PVU/ANEXOS DGO/Informe_Ejecutivo_Salud.docx")

def add_heading(doc, text, level):
    if level == 1:
        doc.add_heading(text, level=0)
    else:
        doc.add_heading(text, level=level)

def parse_table(lines, start_idx):
    # lines contain header, separator, then rows until a blank line or end
    header = [h.strip() for h in lines[start_idx].strip().split('|') if h]
    rows = []
    i = start_idx + 2  # skip separator line
    while i < len(lines) and lines[i].strip():
        row = [c.strip() for c in lines[i].strip().split('|') if c]
        if row:
            rows.append(row)
        i += 1
    return header, rows, i

def main():
    doc = Document()
    lines = md_path.read_text(encoding='utf-8').splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith('# '):
            add_heading(doc, line[2:].strip(), 1)
            i += 1
        elif line.startswith('## '):
            add_heading(doc, line[3:].strip(), 2)
            i += 1
        elif line.startswith('|'):
            # start of a markdown table
            header, rows, new_i = parse_table(lines, i)
            table = doc.add_table(rows=1, cols=len(header))
            hdr_cells = table.rows[0].cells
            for idx, h in enumerate(header):
                hdr_cells[idx].text = h
            for row in rows:
                row_cells = table.add_row().cells
                for idx, cell in enumerate(row):
                    row_cells[idx].text = cell
            i = new_i
        elif line.strip():
            doc.add_paragraph(line.strip())
            i += 1
        else:
            i += 1
    doc.save(output_path)

if __name__ == "__main__":
    main()
