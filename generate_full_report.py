import docx
import json
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_table_border(table):
    tbl = table._tbl
    for row in tbl.tr_lst:
        for cell in row.tc_lst:
            tcPr = cell.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'bottom', 'left', 'right']:
                edge = OxmlElement(f'w:{side}')
                edge.set(qn('w:val'), 'single')
                edge.set(qn('w:sz'), '4')
                edge.set(qn('w:space'), '0')
                edge.set(qn('w:color'), '000000')
                tcBorders.append(edge)
            tcPr.append(tcBorders)

def calculate_increment(v2025, v2024):
    try:
        v25 = float(str(v2025).replace(',', '').replace('%', '').replace(' ', ''))
        v24 = float(str(v2024).replace(',', '').replace('%', '').replace(' ', ''))
        if v24 == 0: return "N/A"
        inc = ((v25 - v24) / v24) * 100
        return f"{inc:+.1f}%"
    except:
        return "N/A"

def generate_full_report():
    json_path = r"c:\Users\aicil\.gemini\antigravity\scratch\all_data.json"
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    doc = docx.Document()
    
    # Clean up default styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Title
    title = doc.add_paragraph('INFORME DE GOBIERNO 2025: PROGRAMA DE VACUNACIÓN UNIVERSAL')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)
    doc.add_paragraph('Periodo Comparativo: Octubre - Diciembre (2024 vs 2025)').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro Paragraphs
    doc.add_paragraph("2025 representa un año de consolidación estratégica para la Secretaría de Salud en Durango, logrando optimizar la captación de población de responsabilidad y blanco a través de brigadas operativas eficientes.")
    doc.add_paragraph("2025 se distingue por el uso de herramientas tecnológicas para el seguimiento en tiempo real de las coberturas, permitiendo identificar áreas de oportunidad y aplicar refuerzos logísticos de manera inmediata.")

    # Process individual tables (skipping the first general population table and handling others)
    for i, table_raw in enumerate(data['tables']):
        if i == 0: continue # Skip general population table
        
        # Identify Vaccine Name from the first row or surrounding text
        # Usually row 0 has the name
        vaccine_name = ""
        if len(table_raw[0]) == 1:
            vaccine_name = table_raw[0][0]
            rows_start = 1
        else:
            vaccine_name = f"Tabla de Vacunación {i}"
            rows_start = 0
            
        doc.add_heading(vaccine_name, 2).runs[0].font.color.rgb = docx.shared.RGBColor(0,0,0)
        
        # Add analysis paragraph before each table starting with 2025
        doc.add_paragraph(f"2025 presenta un monitoreo detallado para la {vaccine_name}, analizando el comportamiento de las dosis aplicadas sectoriales frente a las metas programadas para este cuarto trimestre.")

        # Re-structure table data with Increment Column
        # Headings
        header = table_raw[rows_start]
        new_header = header + ["Incremento (%)"]
        
        new_table_rows = [new_header]
        
        # We need to find matching 2024 and 2025 lines
        # This is tricky because the tables have different structures.
        # Simple strategy: If it has 2 years in separate rows, compare them.
        
        temp_rows = table_raw[rows_start+1:]
        
        processed_rows = []
        # Attempt to match 2024 and 2025 rows
        for r_idx, r_data in enumerate(temp_rows):
            if not r_data or "Fuente:" in r_data[0]: continue
            
            # If row is for 2025, look for a matching 2024 row to calc increment
            if "2025" in r_data[0]:
                inc_val = "N/A"
                # Search for 2024 matching row (same dose desc or just looking below)
                for r_search in temp_rows:
                    if "2024" in r_search[0]:
                        # Check if it's the same dose (e.g. "Primera dosis")
                        dose_type = r_data[0].split(',')[0].strip() if ',' in r_data[0] else ""
                        dose_type_search = r_search[0].split(',')[0].strip() if ',' in r_search[0] else ""
                        
                        if dose_type == dose_type_search or (not dose_type and not dose_type_search):
                            # Calc increment on "Dosis aplicada" (usually index 1)
                            if len(r_data) > 1 and len(r_search) > 1:
                                inc_val = calculate_increment(r_data[1], r_search[1])
                            break
                processed_rows.append(r_data + [inc_val])
            elif "2024" in r_data[0]:
                processed_rows.append(r_data + ["-"])
            else:
                # Other rows (labels, etc)
                processed_rows.append(r_data + [""])

        # Create Table in Word
        if not processed_rows: continue
        
        word_table = doc.add_table(rows=len(processed_rows)+1, cols=len(new_header))
        word_table.style = 'Table Grid'
        
        # Fill Headers
        for h_idx, h_text in enumerate(new_header):
            cell = word_table.cell(0, h_idx)
            cell.text = h_text
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Fill Content
        for r_idx, r_data in enumerate(processed_rows):
            for c_idx, c_val in enumerate(r_data):
                if c_idx >= len(new_header): break
                cell = word_table.cell(r_idx+1, c_idx)
                cell.text = str(c_val)
                para = cell.paragraphs[0]
                if para.runs:
                    run = para.runs[0]
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        set_table_border(word_table)
        doc.add_paragraph("") # Space after table

    output_path = r"c:\Users\aicil\OneDrive\Escritorio\PVU\INFORMES PVU\INFORME DE GOBIERNO\INFORME DE GOBIERNO 2025\INFORME_GOBIERNO_2025_COMPLETO_CON_TABLAS_E_INCREMENTOS.docx"
    doc.save(output_path)
    print(f"Full report generated at: {output_path}")

if __name__ == "__main__":
    generate_full_report()
