import json
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load data
with open("locality_doses.json", "r") as f:
    locality_data = json.load(f)
with open("mezquital_pop_conapo.json", "r") as f:
    pop_data = json.load(f)
with open("mezquital_doses_age.json", "r") as f:
    doses_data = json.load(f)

# Base report path
base_report = r"C:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\mezquital\INFORMES\Informe_Vacunacion_Mezquital_2026_Final_Heatmap.docx"
output_report = r"C:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\mezquital\INFORMES\Informe_Vacunacion_Mezquital_2026_Completo.docx"

try:
    doc = docx.Document(base_report)
    
    # Add a new section for detailed locality data
    doc.add_page_break()
    doc.add_heading('Dosis Aplicadas por Localidad (Mezquital 2026)', level=1)
    
    # Add locality table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Localidad'
    hdr_cells[1].text = 'Dosis Aplicadas'
    hdr_cells[2].text = 'Población (INEGI)'
    hdr_cells[3].text = 'Alcance (%)'
    
    # Format headers
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    for item in locality_data:
        row_cells = table.add_row().cells
        row_cells[0].text = item['Localidad']
        row_cells[1].text = f"{int(item['Doses']):,}"
        row_cells[2].text = f"{int(item['Population']):,}"
        # Reach is a float (0 to 1 usually, or more)
        val = item['Reach']
        if isinstance(val, (int, float)):
            row_cells[3].text = f"{val*100:.2f}%" if val != 0 else "0.00%"
        else:
            row_cells[3].text = str(val)

    # Source for this table
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Fuente: Concentrado Vacunación Mezquital 2026 / INEGI ITER 2020.")
    run.font.size = Pt(8)
    run.italic = True

    # Add coverage by age group section
    doc.add_page_break()
    doc.add_heading('Cobertura de Vacunación por Grupo de Edad (CONAPO 2026)', level=1)
    
    # Table for age groups
    table_age = doc.add_table(rows=1, cols=4)
    table_age.style = 'Table Grid'
    hdr_cells = table_age.rows[0].cells
    hdr_cells[0].text = 'Grupo de Edad'
    hdr_cells[1].text = 'Población (CONAPO 2026)'
    hdr_cells[2].text = 'Dosis Aplicadas (2026)'
    hdr_cells[3].text = 'Cobertura (%)'
    
    # Format headers
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    for label in pop_data.keys():
        pop = pop_data[label]
        doses = doses_data.get(label, 0)
        coverage = (doses / pop * 100) if pop > 0 else 0
        
        row_cells = table_age.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = f"{int(pop):,}"
        row_cells[2].text = f"{int(doses):,}"
        row_cells[3].text = f"{coverage:.2f}%"

    # Source for this table
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Fuente: SIS/CeNSIA, consultado el 04 de mayo de 2026; Proyecciones de Población CONAPO 2026.")
    run.font.size = Pt(8)
    run.italic = True
    
    doc.save(output_report)
    print(f"Final report saved to {output_report}")

except Exception as e:
    print(f"Error: {e}")
