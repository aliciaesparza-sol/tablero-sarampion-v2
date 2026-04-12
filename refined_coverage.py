import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Configuración de rutas
files = {
    'poblacion': r'c:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\COBERTURA DE VACUNACIÓN\TABLERO SARAMPION V2\Poblacion municipio edad simple y sexo Mexico 2026 CENJSIA EGM.xlsx',
    'csv': r'c:\Descargas_SRP\SRP-SR-2025_10-04-2026 09-26-01.csv',
    'output_docx': r'C:\Users\aicil\.gemini\antigravity\scratch\TARJETA_MEZQUITAL_2026_RESCATADA_V2.docx',
    'cubo_total': 669
}

def calculate_revised_coverage():
    # ── 1. POBLACION (METAS) ───────────────────────────────────────────
    print("Calculando Metas (CENJSIA 2026)...")
    df_pob = pd.read_excel(files['poblacion'], sheet_name='Durango', header=None)
    
    # Hombres Row index 6 (Age 0), Mujeres Row index 124 (Age 0)
    # Mezquital es Col 15
    mez_col = 15
    h_idx = 6
    m_idx = 124
    
    pops = {}
    for age in range(101):
        try:
            h = float(df_pob.iloc[h_idx + age, mez_col])
            m = float(df_pob.iloc[m_idx + age, mez_col])
            pops[age] = h + m
        except:
            pops[age] = 0
            
    metas = {
        '6 a 11 meses':          0.5 * pops[0],
        '1 año':                 1.0 * pops[1],
        '18 meses':              1.0 * pops[1],
        'Rezagados 2 a 12 años': 0.5 * sum(pops[a] for a in range(2, 13)),
        '13 a 19 años':           0.5 * sum(pops[a] for a in range(13, 20)),
        '20 a 39 años':           0.5 * sum(pops[a] for a in range(20, 40)),
        '40 a 49 años':           0.5 * sum(pops[a] for a in range(40, 50))
    }

    # ── 2. DOSIS CSV (MAYO 2025 - ABRIL 2026) ─────────────────────────
    print("Procesando Dosis CSV...")
    df_csv = pd.read_csv(files['csv'], encoding='latin-1')
    df_csv.columns = [c.strip() for c in df_csv.columns]
    mez_csv = df_csv[df_csv['MUNICIPIO'].str.contains('MEZQUITAL', case=False, na=False)]
    
    csv_map = {
        '6 a 11 meses': ['SRP 6 A 11 MESES PRIMERA', 'SR 6 A 11 MESES PRIMERA'],
        '1 año': ['SRP 1 ANIO  PRIMERA', 'SR 1 ANIO PRIMERA'],
        '18 meses': ['SRP 18 MESES SEGUNDA', 'SR 18 MESES SEGUNDA'],
        'Rezagados 2 a 12 años': [
            'SRP 2 A 5 ANIOS PRIMERA', 'SRP 6 ANIOS PRIMERA', 'SRP 7 A 9 ANIOS PRIMERA', 'SRP 10 A 12 ANIOS PRIMERA',
            'SRP 2 A 5 ANIOS SEGUNDA', 'SRP 6 ANIOS SEGUNDA', 'SRP 7 A 9 ANIOS SEGUNDA', 'SRP 10 A 12 ANIOS SEGUNDA',
            'SR 2 A 5 ANIOS PRIMERA', 'SR 6 ANIOS PRIMERA', 'SR 7 A 9 ANIOS PRIMERA', 'SR 10 A 12 ANIOS PRIMERA',
            'SR 2 A 5 ANIOS SEGUNDA', 'SR 6 ANIOS SEGUNDA', 'SR 7 A 9 ANIOS SEGUNDA', 'SR 10 A 12 ANIOS SEGUNDA'
        ],
        '13 a 19 años': [
            'SRP 13 A 19 ANIOS PRIMERA', 'SRP 13 A 19 ANIOS SEGUNDA', 'SRP 10 A 19 ANIOS PRIMERA', 'SRP 10 A 19 ANIOS SEGUNDA',
            'SR 13 A 19 ANIOS PRIMERA', 'SR 13 A 19 ANIOS SEGUNDA', 'SR 10 A 19 ANIOS PRIMERA', 'SR 10 A 19 ANIOS SEGUNDA'
        ],
        '20 a 39 años': [
            'SRP 20 A 29 ANIOS PRIMERA', 'SRP 20 A 29 ANIOS SEGUNDA', 'SRP 30 A 39 ANIOS PRIMERA', 'SRP 30 A 39 ANIOS SEGUNDA',
            'SR 20 A 29 ANIOS PRIMERA', 'SR 20 A 29 ANIOS SEGUNDA', 'SR 30 A 39 ANIOS PRIMERA', 'SR 30 A 39 ANIOS SEGUNDA'
        ],
        '40 a 49 años': [
            'SRP 40 A 49 ANIOS PRIMERA', 'SRP 40 A 49 ANIOS SEGUNDA', 'SR 40 A 49 ANIOS PRIMERA', 'SR 40 A 49 ANIOS SEGUNDA'
        ]
    }
    
    actual_csv_doses = {}
    for g, cols in csv_map.items():
        valid = [c for c in cols if c in mez_csv.columns]
        actual_csv_doses[g] = mez_csv[valid].apply(pd.to_numeric, errors='coerce').fillna(0).sum().sum()

    # ── 3. DISTRIBUCION DOSIS CUBO (669) ─────────────────────────────
    # Distribuir 669 proporcional a la tendencia Real del CSV
    total_csv = sum(actual_csv_doses.values())
    print(f"Total Dosis CSV detectadas: {total_csv}")
    
    jan_may_doses = {}
    for g, val in actual_csv_doses.items():
        jan_may_doses[g] = (val / total_csv) * files['cubo_total'] if total_csv > 0 else 0

    # ── 4. TABLA FINAL ────────────────────────────────────────────────
    report = []
    for g in metas.keys():
        m = metas[g]
        d_j = jan_may_doses[g]
        d_c = actual_csv_doses[g]
        tot = d_j + d_c
        report.append({
            'Grupo': g,
            'Meta': round(m),
            'Dosis (Ene-May 2025)': round(d_j),
            'Dosis (May 2025-Abr 2026)': round(d_c),
            'Total': round(tot),
            'Cobertura (%)': round((tot/m)*100, 2) if m > 0 else 0
        })
    
    df_final = pd.DataFrame(report)
    print("\nREPORTE FINAL MEZQUITAL")
    print(df_final.to_string(index=False))
    
    # ── 5. ACTUALIZAR DOCX ────────────────────────────────────────────
    print("\nGenerando Documento Word...")
    doc = Document()
    doc.add_heading('Avance de Cobertura Vacunal Mezquital 2026', 0)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Grupo Etario'
    hdr_cells[1].text = 'Meta (Poblacion)'
    hdr_cells[2].text = 'Dosis Aplicadas (Total)'
    hdr_cells[3].text = 'Avance (%)'
    
    for r in report:
        row_cells = table.add_row().cells
        row_cells[0].text = r['Grupo']
        row_cells[1].text = f"{r['Meta']:,}"
        row_cells[2].text = f"{r['Total']:,}"
        row_cells[3].text = f"{r['Cobertura (%)']}%"

    doc.add_paragraph('\nFuente: Estimaciones basadas en CONAPO/CENJSIA 2026 y reportes institucionales SRP-SR.')
    doc.save(files['output_docx'])
    print(f"Documento guardado en: {files['output_docx']}")

if __name__ == "__main__":
    calculate_revised_coverage()
