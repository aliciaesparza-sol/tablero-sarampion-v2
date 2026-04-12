from docx import Document

doc = Document('mezquital_card.docx')
if len(doc.tables) > 1:
    table = doc.tables[1]
    print("Table 1 (Age Groups):")
    for i in range(min(15, len(table.rows))):
        cells = [cell.text.strip().replace('\n', ' ') for cell in table.rows[i].cells]
        # Remove duplicate adjacent texts due to merging
        unique_cells = []
        if cells:
            unique_cells.append(cells[0])
            for c in cells[1:]:
                if c != unique_cells[-1]:
                    unique_cells.append(c)
        print(f"  Row {i}: {unique_cells}")
else:
    print("Only one table found.")
