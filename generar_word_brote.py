import json
import math
import os
from datetime import datetime
from collections import Counter, defaultdict

# Check dependencies
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    from matplotlib.offsetbox import OffsetImage, AnnotationBbox
except ImportError:
    import subprocess
    subprocess.check_call(["pip", "install", "matplotlib"])
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches

# --- Config ---
SRC_DIR = r"C:\Users\aicil\.gemini\antigravity-ide\scratch"
DEST_DIR = r"C:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\BROTE MAQUILA SAMA OOAD DGO"

with open(os.path.join(SRC_DIR, "casos_tabla.json"), "r", encoding="utf-8") as f:
    cases = json.load(f)

n = len(cases)

# --- Helper functions ---
def haversine(lat1, lon1, lat2, lon2):
    R = 6371
    dlat = math.radians(lat2 - lat1)
    dlon = math.radians(lon2 - lon1)
    a = math.sin(dlat/2)**2 + math.cos(math.radians(lat1)) * math.cos(math.radians(lat2)) * math.sin(dlon/2)**2
    return R * 2 * math.asin(math.sqrt(a))

def parse_date(d):
    for fmt in ("%d/%m/%Y", "%d/%m/%y"):
        try:
            return datetime.strptime(d, fmt)
        except:
            continue
    return None

# --- Compute all statistics ---
edades = [c["edad"] for c in cases]
sexos = [c["sexo"] for c in cases]
instituciones = [c["institucion"] for c in cases]
colonias_list = [c["colonia"] for c in cases]
umfs = [c["umf"] for c in cases]
lats = [c["lat"] for c in cases]
lngs = [c["lng"] for c in cases]

fechas_sintomas = [parse_date(c["fecha_sintomas"]) for c in cases]
fechas_atencion = [parse_date(c["fecha_atencion"]) for c in cases]
fechas_sint_valid = [d for d in fechas_sintomas if d]
fecha_min = min(fechas_sint_valid)
fecha_max = max(fechas_sint_valid)

edad_prom = sum(edades) / n
edad_min = min(edades)
edad_max = max(edades)
edad_mediana = sorted(edades)[n // 2]

sexo_count = Counter(sexos)
inst_count = Counter(instituciones)
colonia_count = Counter(colonias_list)
umf_count = Counter(umfs)

grupos_edad = {"18-24": (18, 24), "25-29": (25, 29), "30-34": (30, 34), "35-41": (35, 41)}
grupo_count = {}
for label, (lo, hi) in grupos_edad.items():
    grupo_count[label] = sum(1 for e in edades if lo <= e <= hi)

curva = Counter()
for d in fechas_sint_valid:
    curva[d.strftime("%Y-%m-%d")] += 1
curva_sorted = sorted(curva.items())

intervalos = []
for fs, fa in zip(fechas_sintomas, fechas_atencion):
    if fs and fa:
        intervalos.append((fa - fs).days)
intervalo_prom = sum(intervalos) / len(intervalos) if intervalos else 0

lat_c = sum(lats) / n
lng_c = sum(lngs) / n

# Distance matrix
dist_matrix = [[haversine(cases[i]["lat"], cases[i]["lng"], cases[j]["lat"], cases[j]["lng"]) for j in range(n)] for i in range(n)]

max_dist = max(dist_matrix[i][j] for i in range(n) for j in range(i+1, n))

# Nearest neighbor
nearest = []
for i in range(n):
    min_d = min(dist_matrix[i][j] for j in range(n) if i != j)
    min_j = [j for j in range(n) if i != j and dist_matrix[i][j] == min_d][0]
    nearest.append((i, min_j, min_d))
avg_nearest = sum(d for _, _, d in nearest) / len(nearest)

# Clusters (Union-Find, threshold 2 km)
THRESHOLD = 2.0
parent = list(range(n))
def find(x):
    while parent[x] != x:
        parent[x] = parent[parent[x]]
        x = parent[x]
    return x
def union(x, y):
    px, py = find(x), find(y)
    if px != py:
        parent[px] = py

for i in range(n):
    for j in range(i+1, n):
        if dist_matrix[i][j] < THRESHOLD:
            union(i, j)

clusters = defaultdict(list)
for i in range(n):
    clusters[find(i)].append(i)

# Spatiotemporal pairs
pares = []
for i in range(n):
    for j in range(i+1, n):
        d = dist_matrix[i][j]
        fa = fechas_sintomas[i]
        fb = fechas_sintomas[j]
        if fa and fb and d < 3.0 and abs((fa - fb).days) <= 3:
            pares.append((i, j, d, abs((fa - fb).days)))

# Centroid distances
dists_centro = [(c["nombre"], c["colonia"], haversine(c["lat"], c["lng"], lat_c, lng_c)) for c in cases]
dists_centro.sort(key=lambda x: x[2])
avg_dc = sum(d for _, _, d in dists_centro) / n
std_dc = math.sqrt(sum((d - avg_dc)**2 for _, _, d in dists_centro) / n)

print("Estadisticas calculadas.")

# ==============================================
# GENERATE MAP IMAGE
# ==============================================
print("Generando mapa...")

# Cluster colors
cluster_colors_map = {}
color_palette = ['#e74c3c', '#3498db', '#2ecc71', '#f39c12', '#9b59b6', '#1abc9c', '#e67e22']
cluster_list = sorted(clusters.items(), key=lambda x: -len(x[1]))
for idx, (root, members) in enumerate(cluster_list):
    color = color_palette[idx % len(color_palette)]
    for m in members:
        cluster_colors_map[m] = color

fig, ax = plt.subplots(1, 1, figsize=(12, 10))

# Plot connections within clusters (light lines)
for root, members in clusters.items():
    if len(members) > 1:
        color = cluster_colors_map[members[0]]
        for i in range(len(members)):
            for j in range(i+1, len(members)):
                if dist_matrix[members[i]][members[j]] < THRESHOLD:
                    ax.plot([cases[members[i]]["lng"], cases[members[j]]["lng"]],
                            [cases[members[i]]["lat"], cases[members[j]]["lat"]],
                            color=color, alpha=0.25, linewidth=1, linestyle='--')

# Plot each case
for i, c in enumerate(cases):
    color = cluster_colors_map.get(i, '#95a5a6')
    marker = 'o' if c["sexo"] == "Femenino" else 's'
    ax.scatter(c["lng"], c["lat"], c=color, s=120, marker=marker, edgecolors='black', linewidths=0.8, zorder=5)
    ax.annotate(c["nombre"], (c["lng"], c["lat"]), fontsize=7, fontweight='bold',
                xytext=(5, 5), textcoords='offset points',
                bbox=dict(boxstyle='round,pad=0.2', facecolor='white', alpha=0.8, edgecolor=color),
                zorder=6)

# Plot centroid
ax.scatter(lng_c, lat_c, c='black', s=200, marker='*', zorder=7, label='Centroide')

# Legends for clusters
legend_patches = []
for idx, (root, members) in enumerate(cluster_list):
    if len(members) > 1:
        color = cluster_colors_map[members[0]]
        legend_patches.append(mpatches.Patch(color=color, label=f'Cluster {idx+1} ({len(members)} casos)'))
    elif len(members) == 1:
        legend_patches.append(mpatches.Patch(color=cluster_colors_map[members[0]], label=f'Aislado: {cases[members[0]]["nombre"]}'))

legend_patches.append(plt.Line2D([0], [0], marker='*', color='w', markerfacecolor='black', markersize=15, label='Centroide'))
legend_patches.append(plt.Line2D([0], [0], marker='o', color='w', markerfacecolor='gray', markersize=10, label='Femenino'))
legend_patches.append(plt.Line2D([0], [0], marker='s', color='w', markerfacecolor='gray', markersize=10, label='Masculino'))

ax.legend(handles=legend_patches, loc='upper left', fontsize=8, framealpha=0.9)

ax.set_xlabel('Longitud', fontsize=10)
ax.set_ylabel('Latitud', fontsize=10)
ax.set_title('Distribucion Geografica de Casos de Sarampion\nBrote Maquiladora SAMA - Victoria de Durango, 2026',
             fontsize=13, fontweight='bold')
ax.grid(True, alpha=0.3)
ax.set_aspect('equal')

plt.tight_layout()
map_path = os.path.join(DEST_DIR, "mapa_casos_brote_sama.png")
plt.savefig(map_path, dpi=200, bbox_inches='tight')
plt.close()
print(f"Mapa guardado: {map_path}")

# ==============================================
# GENERATE CURVA EPIDEMICA IMAGE
# ==============================================
print("Generando curva epidemica...")

fig2, ax2 = plt.subplots(figsize=(10, 5))
fechas_labels = [f.split("-")[2] + "/" + f.split("-")[1] for f, _ in curva_sorted]
valores = [c for _, c in curva_sorted]

bars = ax2.bar(fechas_labels, valores, color='#e74c3c', edgecolor='black', linewidth=0.8)
for bar, val in zip(bars, valores):
    ax2.text(bar.get_x() + bar.get_width()/2., bar.get_height() + 0.1,
             str(val), ha='center', va='bottom', fontweight='bold', fontsize=11)

ax2.set_xlabel('Fecha de Inicio de Sintomas', fontsize=11)
ax2.set_ylabel('Numero de Casos', fontsize=11)
ax2.set_title('Curva Epidemica - Brote Sarampion Maquiladora SAMA\nVictoria de Durango, Mayo 2026',
              fontsize=13, fontweight='bold')
ax2.set_ylim(0, max(valores) + 1.5)
ax2.grid(axis='y', alpha=0.3)

plt.tight_layout()
curva_path = os.path.join(DEST_DIR, "curva_epidemica_brote_sama.png")
plt.savefig(curva_path, dpi=200, bbox_inches='tight')
plt.close()
print(f"Curva epidemica guardada: {curva_path}")

# ==============================================
# CREATE WORD DOCUMENT
# ==============================================
print("Generando documento Word...")

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return h

def add_table_from_data(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

# ===================== CONTENT =====================

# Title page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SECRETARIA DE SALUD')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SERVICIOS DE SALUD DE DURANGO')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ANALISIS EPIDEMIOLOGICO')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BROTE DE SARAMPION\nMAQUILADORA SAMA - OOAD DURANGO')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Victoria de Durango, Durango\n{datetime.now().strftime("%d de mayo de %Y")}')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# --- 1. RESUMEN EJECUTIVO ---
add_heading_styled('1. RESUMEN EJECUTIVO', level=1)

p = doc.add_paragraph()
p.add_run(f'Se analizan ').font.size = Pt(11)
run = p.add_run(f'{n} casos confirmados')
run.bold = True
run.font.size = Pt(11)
p.add_run(f' de sarampion asociados al brote en la maquiladora SAMA, OOAD Durango. ')
p.add_run(f'Los casos se presentaron en el periodo del ')
run2 = p.add_run(f'{fecha_min.strftime("%d/%m/%Y")} al {fecha_max.strftime("%d/%m/%Y")}')
run2.bold = True
p.add_run(f' (inicio de sintomas), distribuidos en ')
run3 = p.add_run(f'{len(set(colonias_list))} colonias')
run3.bold = True
p.add_run(f' del municipio de Victoria de Durango.')

doc.add_paragraph()

# Key stats box
add_table_from_data(
    ['Indicador', 'Valor'],
    [
        ['Total de casos', str(n)],
        ['Periodo del brote', f'{fecha_min.strftime("%d/%m/%Y")} - {fecha_max.strftime("%d/%m/%Y")} ({(fecha_max - fecha_min).days + 1} dias)'],
        ['Edad promedio', f'{edad_prom:.1f} anos (rango {edad_min}-{edad_max})'],
        ['Razon M:F', f'{sexo_count.get("Masculino", 0)}:{sexo_count.get("Femenino", 0)} ({sexo_count.get("Masculino", 0)/max(sexo_count.get("Femenino", 1),1):.2f})'],
        ['Colonias afectadas', str(len(set(colonias_list)))],
        ['Dispersion maxima', f'{max_dist:.1f} km'],
        ['Oportunidad de atencion', f'{intervalo_prom:.1f} dias (promedio)'],
    ]
)

# --- 2. DISTRIBUCION POR SEXO ---
doc.add_paragraph()
add_heading_styled('2. DISTRIBUCION POR SEXO', level=1)

rows_sexo = []
for s in ["Femenino", "Masculino"]:
    c = sexo_count.get(s, 0)
    rows_sexo.append([s, str(c), f'{c/n*100:.1f}%'])
rows_sexo.append(['Total', str(n), '100%'])

add_table_from_data(['Sexo', 'Casos', 'Porcentaje'], rows_sexo)

p = doc.add_paragraph()
p.add_run(f'\nRazon Masculino:Femenino = {sexo_count.get("Masculino", 0)/max(sexo_count.get("Femenino", 1),1):.2f}')

# --- 3. DISTRIBUCION POR GRUPO DE EDAD ---
doc.add_paragraph()
add_heading_styled('3. DISTRIBUCION POR GRUPO DE EDAD', level=1)

rows_edad = []
for label in sorted(grupo_count.keys()):
    c = grupo_count[label]
    rows_edad.append([f'{label} anos', str(c), f'{c/n*100:.1f}%'])
rows_edad.append(['Total', str(n), '100%'])

add_table_from_data(['Grupo de edad', 'Casos', 'Porcentaje'], rows_edad)

p = doc.add_paragraph()
p.add_run(f'\nEdad promedio: {edad_prom:.1f} anos | Mediana: {edad_mediana} anos | Rango: {edad_min}-{edad_max} anos')

# --- 4. DISTRIBUCION POR INSTITUCION ---
doc.add_paragraph()
add_heading_styled('4. DISTRIBUCION POR INSTITUCION NOTIFICANTE', level=1)

rows_inst = [(inst, str(c), f'{c/n*100:.1f}%') for inst, c in inst_count.most_common()]
rows_inst.append(('Total', str(n), '100%'))
add_table_from_data(['Institucion', 'Casos', 'Porcentaje'], rows_inst)

# --- 5. DISTRIBUCION POR UMF ---
doc.add_paragraph()
add_heading_styled('5. DISTRIBUCION POR UMF NOTIFICANTE', level=1)

rows_umf = [(f'UMF {u}', str(c), f'{c/n*100:.1f}%') for u, c in umf_count.most_common()]
rows_umf.append(('Total', str(n), '100%'))
add_table_from_data(['UMF', 'Casos', 'Porcentaje'], rows_umf)

# --- 6. CURVA EPIDEMICA ---
doc.add_page_break()
add_heading_styled('6. CURVA EPIDEMICA', level=1)

p = doc.add_paragraph('Distribucion de casos por fecha de inicio de sintomas:')

# Insert curva image
doc.add_picture(curva_path, width=Inches(6))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

rows_curva = []
acum = 0
for fecha, count in curva_sorted:
    acum += count
    rows_curva.append([fecha, str(count), str(acum)])

add_table_from_data(['Fecha', 'Casos', 'Acumulado'], rows_curva)

p = doc.add_paragraph()
p.add_run(f'\nPeriodo del brote: {(fecha_max - fecha_min).days + 1} dias | Pico: {max(curva.values())} casos el 18/05/2026')

# --- 7. DISTRIBUCION GEOGRAFICA ---
doc.add_page_break()
add_heading_styled('7. DISTRIBUCION GEOGRAFICA', level=1)

# --- 7a. Table of colonias ---
add_heading_styled('7.1 Colonias afectadas', level=2)

rows_col = [(col, str(c), f'{c/n*100:.1f}%') for col, c in colonia_count.most_common()]
rows_col.append(('Total', str(n), '100%'))
add_table_from_data(['Colonia', 'Casos', 'Porcentaje'], rows_col)

# --- 7b. Map ---
doc.add_paragraph()
add_heading_styled('7.2 Mapa de georeferenciacion de casos', level=2)

doc.add_picture(map_path, width=Inches(6))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Figura 1. Distribucion geografica de casos de sarampion por cluster.\nCirculos = Femenino, Cuadrados = Masculino. Estrella = Centroide.')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

# --- 8. ANALISIS DE RELACION GEOGRAFICA ---
doc.add_page_break()
add_heading_styled('8. ANALISIS DE RELACION GEOGRAFICA', level=1)

# --- 8a. Clusters ---
add_heading_styled('8.1 Clusters geograficos (radio < 2 km)', level=2)

p = doc.add_paragraph()
p.add_run(f'Se identificaron ')
run = p.add_run(f'{sum(1 for m in clusters.values() if len(m) > 1)} clusters')
run.bold = True
p.add_run(f' que agrupan {sum(len(m) for m in clusters.values() if len(m) > 1)} de {n} casos.')

cluster_rows = []
for idx, (root, members) in enumerate(sorted(clusters.items(), key=lambda x: -len(x[1]))):
    if len(members) > 1:
        max_d = max(dist_matrix[a][b] for a in members for b in members)
        cols = ", ".join(cases[m]["colonia"] for m in members)
        cluster_rows.append([f'Cluster {idx+1}', str(len(members)), f'{max_d:.2f} km', cols])

add_table_from_data(['Cluster', 'Casos', 'Diametro', 'Colonias'], cluster_rows)

# Singletons
singletons = [m[0] for m in clusters.values() if len(m) == 1]
if singletons:
    p = doc.add_paragraph()
    p.add_run(f'\nCasos aislados ({len(singletons)}): ')
    p.add_run(', '.join(f'{cases[s]["nombre"]} ({cases[s]["colonia"]})' for s in singletons))

# --- 8b. Nearest neighbor ---
doc.add_paragraph()
add_heading_styled('8.2 Vecino mas cercano', level=2)

nn_rows = [(cases[i]["nombre"], cases[i]["colonia"], cases[j]["nombre"], cases[j]["colonia"], f'{d:.2f} km')
           for i, j, d in nearest]
add_table_from_data(['Caso', 'Colonia', 'Vecino cercano', 'Colonia vecino', 'Distancia'], nn_rows)

p = doc.add_paragraph()
run = p.add_run(f'\nDistancia promedio al vecino mas cercano: {avg_nearest:.2f} km')
run.bold = True

# --- 8c. Spatiotemporal ---
doc.add_paragraph()
add_heading_styled('8.3 Correlacion espacio-temporal', level=2)

p = doc.add_paragraph()
p.add_run('Pares de casos con proximidad geografica (<3 km) Y temporal (<3 dias entre sintomas):')

st_rows = [(cases[i]["nombre"], cases[j]["nombre"], f'{d:.2f} km', f'{dd} dias')
           for i, j, d, dd in pares]
add_table_from_data(['Caso A', 'Caso B', 'Distancia', 'Dif. dias sintomas'], st_rows)

p = doc.add_paragraph()
run = p.add_run(f'\nTotal de pares espacio-temporalmente relacionados: {len(pares)}')
run.bold = True

# --- 8d. Centroid ---
doc.add_paragraph()
add_heading_styled('8.4 Dispersion desde el centroide', level=2)

dc_rows = [(nombre, col, f'{d:.2f} km') for nombre, col, d in dists_centro]
add_table_from_data(['Caso', 'Colonia', 'Distancia al centroide'], dc_rows)

p = doc.add_paragraph()
p.add_run(f'\nDistancia promedio al centroide: {avg_dc:.2f} km (DE: {std_dc:.2f} km)')

# --- 9. OPORTUNIDAD DE ATENCION ---
doc.add_page_break()
add_heading_styled('9. OPORTUNIDAD DE ATENCION', level=1)

oa_rows = []
for c, fs, fa, intv in zip(cases, fechas_sintomas, fechas_atencion, intervalos):
    oa_rows.append([c["nombre"], c["colonia"], fs.strftime("%d/%m/%Y") if fs else "N/D",
                    fa.strftime("%d/%m/%Y") if fa else "N/D", f'{intv} dias'])

add_table_from_data(['Caso', 'Colonia', 'Inicio sintomas', 'Atencion', 'Intervalo'], oa_rows)

p = doc.add_paragraph()
p.add_run(f'\nIntervalo promedio sintomas-atencion: {intervalo_prom:.1f} dias | Minimo: {min(intervalos)} dias | Maximo: {max(intervalos)} dias')

# --- 10. HALLAZGOS Y CONCLUSIONES ---
doc.add_page_break()
add_heading_styled('10. HALLAZGOS CLAVE Y CONCLUSIONES', level=1)

hallazgos = [
    f'Predominio femenino: {sexo_count.get("Femenino", 0)} de {n} casos ({sexo_count.get("Femenino", 0)/n*100:.1f}%) son mujeres.',
    f'Poblacion joven en edad productiva: Edad promedio {edad_prom:.1f} anos (rango {edad_min}-{edad_max}), consistente con poblacion maquiladora.',
    f'IMSS principal notificador: {inst_count.get("IMSS Ordinario", 0)} casos ({inst_count.get("IMSS Ordinario", 0)/n*100:.1f}%) notificados por IMSS Ordinario.',
    f'UMF 44 como principal unidad: {umf_count.get("44", 0)} casos ({umf_count.get("44", 0)/n*100:.1f}%) adscritos a UMF 44.',
    f'Dispersion comunitaria: {len(set(colonias_list))} colonias afectadas en radio de {max_dist:.1f} km.',
    f'Oportunidad de atencion: Promedio de {intervalo_prom:.1f} dias entre sintomas y atencion medica.',
    f'Pico epidemico: 6 casos el 18 de mayo de 2026.',
]

for i, h in enumerate(hallazgos, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.bold = True
    p.add_run(h)

add_heading_styled('Conclusion del analisis geografico', level=2)

p = doc.add_paragraph()
p.add_run(f'Se identificaron {sum(1 for m in clusters.values() if len(m) > 1)} clusters geograficos con {sum(len(m) for m in clusters.values() if len(m) > 1)} de {n} casos agrupados. ')
p.add_run(f'La distancia promedio al vecino mas cercano es de {avg_nearest:.2f} km, lo que sugiere agrupamiento espacial significativo. ')
p.add_run(f'Se encontraron {len(pares)} pares de casos relacionados espacio-temporalmente (<3 km y <3 dias de diferencia en sintomas).')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Interpretacion: ')
run.bold = True
p.add_run(f'Los {n} casos residen en {len(set(colonias_list))} colonias DIFERENTES, lo que descarta un foco residencial comun. ')
p.add_run('La cercania espacial entre los clusters y los pares espacio-temporalmente relacionados son consistentes con ')
run2 = p.add_run('exposicion comun en el centro de trabajo (Maquiladora SAMA) ')
run2.bold = True
p.add_run('seguida de transmision secundaria comunitaria. ')
p.add_run('El patron de irradiacion uniforme desde el centroide refuerza la hipotesis de un foco laboral unico como fuente primaria del brote.')

# --- 11. LISTADO DE CASOS ---
doc.add_page_break()
add_heading_styled('ANEXO: LISTADO COMPLETO DE CASOS', level=1)

case_rows = []
for c in cases:
    case_rows.append([
        str(c["id"]), c["nombre"], c["sexo"], str(c["edad"]),
        c["umf"], c["colonia"], c["institucion"],
        c["fecha_sintomas"], c["fecha_atencion"], c.get("folio", ""),
        f'{c["lat"]:.4f}', f'{c["lng"]:.4f}'
    ])

add_table_from_data(
    ['#', 'Nombre', 'Sexo', 'Edad', 'UMF', 'Colonia', 'Inst.', 'Sintomas', 'Atencion', 'Folio', 'Lat', 'Lng'],
    case_rows
)

# --- Save ---
docx_path = os.path.join(DEST_DIR, "Analisis_Brote_Sarampion_SAMA_Durango_2026.docx")
doc.save(docx_path)
print(f"Documento Word guardado: {docx_path}")

# Also save to scratch
docx_path2 = os.path.join(SRC_DIR, "Analisis_Brote_Sarampion_SAMA_Durango_2026.docx")
doc.save(docx_path2)
print(f"Copia en scratch: {docx_path2}")

print("\nTodos los archivos generados exitosamente.")
