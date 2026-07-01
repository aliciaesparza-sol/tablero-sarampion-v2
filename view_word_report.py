import sys
from docx import Document

path = r"C:/Users/aicil/.gemini/antigravity-ide/scratch/vaccination_report_full.docx"

try:
    doc = Document(path)
except Exception as e:
    sys.exit(f'Error loading Word file: {e}')

# Print headings and paragraphs
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        print(text)

# Print tables
for table in doc.tables:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append('\t'.join(cells))
    print('\n'.join(rows))
    print('\n')
