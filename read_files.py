import docx
import pandas as pd
import os

docx_path = r"c:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\mezquital\TARJETAS\Tarjeta_Estrategia_operativa_Mezquital_30abril2026.docx"
xlsx_path = r"c:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\mezquital\TARJETAS\PLAN DE VUELOS OK..xlsx"

def read_docx(path):
    doc = docx.Document(path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Also look at tables
    tables_data = []
    for i, table in enumerate(doc.tables):
        table_text = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            table_text.append(row_text)
        tables_data.append(table_text)
        
    return full_text, tables_data

def read_xlsx(path):
    df = pd.read_excel(path)
    return df

try:
    print("--- DOCX CONTENT ---")
    text, tables = read_docx(docx_path)
    for t in text:
        if t.strip():
            print(t)
    
    print("\n--- DOCX TABLES ---")
    for i, table in enumerate(tables):
        print(f"Table {i+1}:")
        for row in table:
            print(row)

    print("\n--- XLSX CONTENT ---")
    df = read_xlsx(xlsx_path)
    print(df.to_string())

except Exception as e:
    print(f"Error: {e}")
