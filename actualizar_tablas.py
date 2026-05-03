
# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# 1. DATOS
# ============================================================
df = pd.read_csv(r'c:\Descargas_SRP\SRP-SR-2025_28-04-2026 08-14-18.csv', encoding='latin1', sep=';')
num_cols = [c for c in df.columns if c not in ['id','INSTITUCION','DELEGACION','ESTADO','JURISDICCION','MUNICIPIO','CLUES','Fecha de registro','Temporada']]
for c in num_cols:
    df[c] = pd.to_numeric(df[c], errors='coerce').fillna(0)
dur26 = df[(df['Temporada']==2026) & (df['ESTADO']=='DURANGO')].copy()

C = {
    'srp_6_11': 'SRP 6 A 11 MESES PRIMERA', 'srp_1a': 'SRP 1 ANIO  PRIMERA',
    'srp_2_5': 'SRP 2 A 5 ANIOS PRIMERA', 'srp_6a': 'SRP 6 ANIOS PRIMERA',
    'srp_7_9': 'SRP 7 A 9 ANIOS PRIMERA', 'srp_10_19': 'SRP 10 A 19 ANIOS PRIMERA',
    'srp_20_29': 'SRP 20 A 29 ANIOS PRIMERA', 'srp_30_39': 'SRP 30 A 39 ANIOS PRIMERA',
    'srp_40_49': 'SRP 40 A 49 ANIOS PRIMERA',
    'srp_salud': 'SRP PERSONAL DE SALUD PRIMERA', 'srp_educ': 'SRP PERSONAL EDUCATIVO PRIMERA',
    'srp_jorn': 'SRP JORNALEROS AGRICOLAS PRIMERA', 'srp_pt': 'SRP  PRIMERA TOTAL',
    'srp_18m': 'SRP 18 MESES SEGUNDA', 'srp_st': 'SRP SEGUNDA TOTAL',
    'sr_6_11': 'SR 6 A 11 MESES PRIMERA', 'sr_1a': 'SR 1 ANIO PRIMERA',
    'sr_pt': 'SR PRIMERA TOTAL', 'sr_st': 'SR SEGUNDA TOTAL',
}
JURS = ['DURANGO','GOMEZ PALACIO','SANTIAGO PAPASQUIARO','RODEO']
JL = {'DURANGO':'Durango','GOMEZ PALACIO':'Gómez Palacio','SANTIAGO PAPASQUIARO':'Santiago Papasquiaro','RODEO':'Rodeo'}

da   = dur26[dur26['SEMANA']==51]
dant = dur26[dur26['SEMANA']==50]
dacu = dur26[dur26['SEMANA']<=51]
daa  = dur26[dur26['SEMANA']<=50]

def v(d,k): return int(d[C[k]].sum())
def vj(d,j,k): return int(d[d['JURISDICCION']==j][C[k]].sum())
def n(x): return f"{x:,}"
def arrow(x): return f"↑ {n(x)}" if x >= 0 else f"↓ {n(abs(x))}"

# Combinar SRP+SR para dosis cero (6-11m) como en el original
def dc_jur(d, j):
    sub = d[d['JURISDICCION']==j]
    return int(sub[C['srp_6_11']].sum() + sub[C['sr_6_11']].sum())
def dc_tot(d):
    return int(d[C['srp_6_11']].sum() + d[C['sr_6_11']].sum())

# Combinar SRP+SR para 2a dosis 18m
def s18_jur(d, j):
    return int(d[d['JURISDICCION']==j][C['srp_18m']].sum() + d[d['JURISDICCION']==j][C['sr_6_11']].sum())

# Acumulados para Tabla 0 (Dosis Cero SRP+SR 6-11m)
dc_acum_ant = {}; dc_acum_act = {}
for j in JURS:
    dc_acum_ant[j] = dc_jur(daa, j)
    dc_acum_act[j] = dc_jur(dacu, j)
dc_acum_ant['TOTAL'] = dc_tot(daa)
dc_acum_act['TOTAL'] = dc_tot(dacu)

# Acumulados para Tabla 4 (2a dosis SRP 18m)
s18_acum_ant = {}; s18_acum_act = {}
for j in JURS:
    s18_acum_ant[j] = vj(daa, j, 'srp_18m')
    s18_acum_act[j] = vj(dacu, j, 'srp_18m')
s18_acum_ant['TOTAL'] = v(daa, 'srp_18m')
s18_acum_act['TOTAL'] = v(dacu, 'srp_18m')

# SRP 1a dosis por grupo edad - ACUMULADOS
srp1_grps_acum = {}
for j in JURS:
    sub = dacu[dacu['JURISDICCION']==j]
    srp1_grps_acum[j] = {
        '6_11': dc_jur(dacu, j),
        '1a':   int(sub[C['srp_1a']].sum()),
        '2_5':  int(sub[C['srp_2_5']].sum()),
        '6a':   int(sub[C['srp_6a']].sum()),
        '7_9':  int(sub[C['srp_7_9']].sum()),
        '10_19':int(sub[C['srp_10_19']].sum()),
        '20_29':int(sub[C['srp_20_29']].sum()),
    }

# Jornaleros acumulados
jorn_acum = {}
for j in JURS:
    jorn_acum[j] = vj(dacu, j, 'srp_jorn')

# ============================================================
# 2. ABRIR DOCUMENTO ACTUALIZADO
# ============================================================
doc = Document(r'C:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\CONASABI\EVIDENCIAS CONASABI_24ABRIL2026\DSP_Vacunación_CONASABIAc2_24abril2026_ACTUALIZADO.docx')

tables = doc.tables

def set_cell(table, r, c, text):
    """Set cell text preserving formatting"""
    cell = table.cell(r, c)
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ''
    if cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = str(text)
    else:
        cell.paragraphs[0].add_run(str(text))

def set_merged_row(table, r, text):
    """Set text in merged row (all cols have same text)"""
    cell = table.cell(r, 0)
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ''
    if cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = str(text)
    else:
        cell.paragraphs[0].add_run(str(text))

# ============================================================
# TABLA 0: Dosis Cero (SR/SRP 6-11m) - 8 filas x 4 cols
# F0: Título | F1: Headers | F2-F5: Jurisdicciones | F6: Total | F7: Referencia
# ============================================================
t = tables[0]
set_merged_row(t, 0, 'Tabla 1. Análisis Comparativo de Aplicación de Dosis Cero (SR/SRP 6 a 11 meses)')
set_cell(t, 1, 0, 'Jurisdicción')
set_cell(t, 1, 1, 'Corte 17/abr')
set_cell(t, 1, 2, 'Corte 24/abr')
set_cell(t, 1, 3, 'Variación')

jur_rows = {2:'DURANGO', 3:'GOMEZ PALACIO', 4:'SANTIAGO PAPASQUIARO', 5:'RODEO'}
for fi, j in jur_rows.items():
    set_cell(t, fi, 0, JL[j])
    set_cell(t, fi, 1, n(dc_acum_ant[j]))
    set_cell(t, fi, 2, n(dc_acum_act[j]))
    set_cell(t, fi, 3, arrow(dc_acum_act[j] - dc_acum_ant[j]))

set_cell(t, 6, 0, 'Total estatal')
set_cell(t, 6, 1, n(dc_acum_ant['TOTAL']))
set_cell(t, 6, 2, n(dc_acum_act['TOTAL']))
set_cell(t, 6, 3, arrow(dc_acum_act['TOTAL'] - dc_acum_ant['TOTAL']))

set_merged_row(t, 7, 'Referencia: Secretaría de Salud – CeNSIA. Informe de avance sectorial de dosis aplicadas de vacuna SR/SRP. Plataforma SIS-CeNSIA. Corte al 24 de abril de 2026.')
print("✅ Tabla 0 actualizada: Dosis Cero")

# ============================================================
# TABLA 1: Comparación intersemanal - 5 filas x 4 cols
# F0: Título | F1: Headers | F2: DC | F3: 18m | F4: Referencia
# ============================================================
t = tables[1]
set_merged_row(t, 0, 'Tabla 2. Comparación intersemanal (17/abr vs 24/abr/2026)')
set_cell(t, 1, 0, 'Rubro')
set_cell(t, 1, 1, 'Semana anterior (17/abr)')
set_cell(t, 1, 2, 'Semana actual (24/abr)')
set_cell(t, 1, 3, 'Variación')

set_cell(t, 2, 0, 'SRP 6–11 meses (Dosis cero)')
set_cell(t, 2, 1, n(dc_acum_ant['TOTAL']))
set_cell(t, 2, 2, n(dc_acum_act['TOTAL']))
set_cell(t, 2, 3, f"+{dc_acum_act['TOTAL']-dc_acum_ant['TOTAL']}")

set_cell(t, 3, 0, 'SRP 18 meses (2ª dosis)')
set_cell(t, 3, 1, n(s18_acum_ant['TOTAL']))
set_cell(t, 3, 2, n(s18_acum_act['TOTAL']))
set_cell(t, 3, 3, f"+{s18_acum_act['TOTAL']-s18_acum_ant['TOTAL']}")

set_merged_row(t, 4, 'Referencia: Secretaría de Salud – CeNSIA. Plataforma SIS-CeNSIA. Corte de información al 24 de abril de 2026.')
print("✅ Tabla 1 actualizada: Comparación intersemanal")

# ============================================================
# TABLA 2: SRP 1a dosis por jurisdicción y grupo de edad - 7 filas x 8 cols
# F0: Título | F1: Headers | F2-F5: Jurisdicciones | F6: Total
# ============================================================
t = tables[2]
set_merged_row(t, 0, 'Tabla 3. SRP primera dosis por jurisdicción y grupo de edad (acumulado al 24/abr/2026)')
grp_keys = ['6_11','1a','2_5','6a','7_9','10_19','20_29']
grp_headers = ['Jurisdicción','6–11 m','1 año','2–5 años','6 años','7–9 años','10–19 años','20–29 años']
for ci, h in enumerate(grp_headers):
    set_cell(t, 1, ci, h)

jur_order = ['DURANGO','GOMEZ PALACIO','RODEO','SANTIAGO PAPASQUIARO']
for ri, j in enumerate(jur_order):
    set_cell(t, ri+2, 0, JL[j])
    for gi, gk in enumerate(grp_keys):
        set_cell(t, ri+2, gi+1, n(srp1_grps_acum[j][gk]))

# Total row
totals_grp = {}
for gk in grp_keys:
    totals_grp[gk] = sum(srp1_grps_acum[j][gk] for j in JURS)
set_cell(t, 6, 0, 'Total')
for gi, gk in enumerate(grp_keys):
    set_cell(t, 6, gi+1, n(totals_grp[gk]))
print("✅ Tabla 2 actualizada: SRP 1a dosis por grupo edad")

# ============================================================
# TABLA 3: Población prioritaria - 7 filas x 5 cols
# F0: Título | F1: Headers | F2-F5: Jurs | F6: Referencia
# ============================================================
t = tables[3]
set_merged_row(t, 0, 'Tabla 4. Cobertura acumulada en población prioritaria (corte al 24 de abril de 2026)')

# Calcular niveles de captación
cap_levels = {}
jorn_vals = {}
for j in JURS:
    jorn_vals[j] = jorn_acum[j]
    srp_t = vj(dacu, j, 'srp_pt')
    if srp_t > 20000: cap_levels[j] = 'Alto'
    elif srp_t > 10000: cap_levels[j] = 'Moderado-alto'
    elif srp_t > 1000: cap_levels[j] = 'Bajo'
    else: cap_levels[j] = 'Muy bajo'

pers_levels = {}
for j in JURS:
    ps = vj(dacu, j, 'srp_salud')
    if ps > 500: pers_levels[j] = 'Alto'
    elif ps > 100: pers_levels[j] = 'Moderado'
    else: pers_levels[j] = 'Bajo' if ps > 10 else 'Muy bajo'

analisis = {
    'DURANGO': 'Mayor volumen operativo estatal',
    'GOMEZ PALACIO': 'Buen desempeño en captación laboral',
    'SANTIAGO PAPASQUIARO': 'Rezago operativo',
    'RODEO': 'Actividad limitada'
}

jur_order_t3 = ['DURANGO','GOMEZ PALACIO','SANTIAGO PAPASQUIARO','RODEO']
for ri, j in enumerate(jur_order_t3):
    set_cell(t, ri+2, 0, JL[j])
    set_cell(t, ri+2, 1, cap_levels[j])
    jl = 'Alto' if jorn_vals[j]>100 else ('Bajo' if jorn_vals[j]>10 else 'Mínimo')
    set_cell(t, ri+2, 2, f"{jl} ({jorn_vals[j]})")
    set_cell(t, ri+2, 3, pers_levels[j])
    set_cell(t, ri+2, 4, analisis[j])

set_merged_row(t, 6, 'Referencia: Secretaría de Salud – CeNSIA. Plataforma SIS-CeNSIA. Corte de información al 24 de abril de 2026.')
print("✅ Tabla 3 actualizada: Población prioritaria")

# ============================================================
# TABLA 4: 2a Dosis SRP 18m - 8 filas x 4 cols
# F0: Título | F1: Headers | F2-F5: Jurs | F6: Total | F7: Referencia
# ============================================================
t = tables[4]
set_merged_row(t, 0, 'Tabla 5. Cumplimiento de Esquema Básico – Segunda Dosis SRP (18 meses)')
set_cell(t, 1, 0, 'Jurisdicción')
set_cell(t, 1, 1, 'Corte 17/abr')
set_cell(t, 1, 2, 'Corte 24/abr')
set_cell(t, 1, 3, 'Variación')

jur_order_t4 = ['DURANGO','GOMEZ PALACIO','SANTIAGO PAPASQUIARO','RODEO']
for ri, j in enumerate(jur_order_t4):
    set_cell(t, ri+2, 0, JL[j])
    set_cell(t, ri+2, 1, n(s18_acum_ant[j]))
    set_cell(t, ri+2, 2, n(s18_acum_act[j]))
    dif = s18_acum_act[j] - s18_acum_ant[j]
    set_cell(t, ri+2, 3, arrow(dif))

set_cell(t, 6, 0, 'Total estatal')
set_cell(t, 6, 1, n(s18_acum_ant['TOTAL']))
set_cell(t, 6, 2, n(s18_acum_act['TOTAL']))
dif_tot = s18_acum_act['TOTAL'] - s18_acum_ant['TOTAL']
set_cell(t, 6, 3, arrow(dif_tot))

set_merged_row(t, 7, 'Referencia: Secretaría de Salud – CeNSIA. Informe de avance sectorial de dosis aplicadas de vacuna SRP. Plataforma SIS-CeNSIA. Corte al 24 de abril de 2026.')
print("✅ Tabla 4 actualizada: 2a Dosis SRP 18m")

# ============================================================
# GUARDAR
# ============================================================
OUTPUT = r'C:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\CONASABI\EVIDENCIAS CONASABI_24ABRIL2026\DSP_Vacunación_CONASABIAc2_24abril2026_ACTUALIZADO.docx'
doc.save(OUTPUT)
print(f"\n✅ Documento final guardado: {OUTPUT}")
