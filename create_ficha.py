import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

template_path = r"c:\Users\aicil\OneDrive\Escritorio\hoja MEMBRETADA GIGANTE2026_carta.docx"
output_path = r"c:\Users\aicil\OneDrive\Escritorio\PVU\SEMANA NACIONAL DE VACUNACION\SEMANA NACIONAL DE VACUNACIÓN 2026\NUEVA_FICHA_TECNICA_ARRANQUE.docx"

try:
    doc = Document(template_path)
    
    # Clear existing content
    for paragraph in doc.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    
    for table in doc.tables:
        t = table._element
        t.getparent().remove(t)
        
    def add_p(text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = 'Arial'
        return p

    # TITLE
    add_p("Ficha Técnica\nArranque de la Semana Nacional de Vacunación 2026\n“Durango se vacuna, Durango se cuida”", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p("Durango, Dgo., abril de 2026", bold=False, size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_paragraph()

    add_p("1. Nombre de la Actividad", bold=True)
    add_p("Arranque de la Semana Nacional de Vacunación 2026 y Estrategia de Recuperación de Esquemas.")
    doc.add_paragraph()

    add_p("2. Fecha y Horario", bold=True)
    add_p("Periodo: Del 25 de abril al 2 de mayo de 2026\nArranque oficial: 25 de abril de 2026\nHorario de atención: 12:00 a 19:00 horas")
    doc.add_paragraph()

    add_p("3. Sede del Arranque", bold=True)
    add_p("Ubicación: Calle Constitución, Zona Centro, Durango\nEntre: 5 de Febrero y 20 de Noviembre")
    doc.add_paragraph()

    add_p("4. Propósito General", bold=True)
    add_p("Fortalecer las acciones del Programa de Vacunación Universal en el estado, mediante la promoción de la aplicación gratuita y universal de vacunas en módulos estratégicos accesibles. Se pondrá especial énfasis en los grupos vulnerables que enfrentan barreras de acceso a los servicios de salud (tales como poblaciones periféricas urbanas, comunidades rurales, zonas fronterizas y pueblos indígenas), facilitando la identificación e intervención oportuna para la recuperación de esquemas de vacunación.")
    doc.add_paragraph()

    add_p("5. Objetivo Específico", bold=True)
    add_p("Iniciar y/o completar los esquemas de vacunación de la población conforme a la línea de vida y el Programa de Vacunación Universal, administrando todas las vacunas disponibles con especial énfasis en: BCG, hexavalente acelular, SRP, SR, DTP, Tdpa, rotavirus, neumococo, hepatitis A y B, y VSR. La estrategia busca facilitar el acceso universal instalando módulos interinstitucionales en zonas de alta afluencia para incrementar las coberturas y promover la prevención de enfermedades.")
    doc.add_paragraph()

    add_p("6. Población Objetivo", bold=True)
    add_p("• Población en general\n• Niñas y niños\n• Adolescentes\n• Personas adultas\n• Grupos con esquemas incompletos o desconocidos\n• Poblaciones vulnerables con barreras de acceso a los servicios de salud")
    doc.add_paragraph()

    add_p("7. Estrategia Operativa", bold=True)
    add_p("La jornada inicial del arranque se efectuará mediante la instalación interinstitucional de 4 Células Fijas de Vacunación (CFV) en una zona de alta concentración poblacional.")
    add_p("Acciones a realizar:")
    add_p("• Revisión de la Cartilla Nacional de Salud.\n• Aplicación de vacunas conforme a las necesidades del esquema basándose en el stock disponible.\n• Orientación preventiva y concientización social de la vacunación.\n• Registro exhaustivo de las dosis aplicadas en el sistema.")
    doc.add_paragraph()

    add_p("8. Personal y Logística de Células de Vacunación", bold=True)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Área'
    hdr_cells[1].text = 'Personal'
    hdr_cells[2].text = 'Función'
    # bold header
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'

    row_data = [
        ('Coordinación', '1 responsable', 'Supervisión del módulo'),
        ('Registro', '4 personas', 'Captura y registro de datos'),
        ('Vacunadores', '4 enfermeros', 'Aplicación del biológico'),
        ('Preparación de biológico', '2 enfermeros', 'Dilución y control de termo'),
        ('Observación', '1 personal', 'Vigilancia de ESAVI'),
        ('Logística', '1 apoyo', 'Orden y flujo de pacientes'),
        ('Total operativo por CFV', '13 personas', ''),
        ('Gran Total Operativo', '52 - 67 personas (por las 4 células de inicio y apoyo logístico general)', '')
    ]
    for area, personal, funcion in row_data:
        row_cells = table.add_row().cells
        row_cells[0].text = area
        row_cells[1].text = personal
        row_cells[2].text = funcion
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)

    doc.add_paragraph()
    add_p("9. Instituciones Participantes", bold=True)
    add_p("Estrategia interinstitucional con la participación de:\n• Secretaría de Salud del Estado de Durango (SSD)\n• Instituto Mexicano del Seguro Social (IMSS)\n• Instituto de Seguridad y Servicios Sociales de los Trabajadores del Estado (ISSSTE)")
    doc.add_paragraph()

    add_p("10. Responsables del Evento", bold=True)
    add_p("Dr. Moisés Nájera Torres - Secretario de Salud y Director General de SSD.\nDra. Citlali Solís Campos - Directora de Salud Pública.\nDra. Aleida Ana Sánchez Monreal - Subdirectora de Epidemiología y Medicina Preventiva.\nDr. Silvano Ramírez Soto - Jefe del Departamento de Enfermedades Transmisibles.\nDra. Alicia J. Esparza Aldaba - Responsable del Programa de Vacunación Universal.\nL.E. Lourdes Saraí Alcantar Vargas - Responsable de Red de Frío.\nL.E. Myriam Machado Arreola - Responsable de la campaña.")

    doc.save(output_path)
    print("Document successfully created and saved.")

except Exception as e:
    print(f"Error creating Document: {e}")
