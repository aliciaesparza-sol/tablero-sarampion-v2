
# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# 1. DATOS DEL CSV
# ============================================================
df = pd.read_csv(r'c:\Descargas_SRP\SRP-SR-2025_28-04-2026 08-14-18.csv', encoding='latin1', sep=';')
num_cols = [c for c in df.columns if c not in ['id','INSTITUCION','DELEGACION','ESTADO','JURISDICCION','MUNICIPIO','CLUES','Fecha de registro','Temporada']]
for c in num_cols:
    df[c] = pd.to_numeric(df[c], errors='coerce').fillna(0)

dur26 = df[(df['Temporada']==2026) & (df['ESTADO']=='DURANGO')].copy()

C = {
    'srp_6_11': 'SRP 6 A 11 MESES PRIMERA',
    'srp_1a':   'SRP 1 ANIO  PRIMERA',
    'srp_2_5':  'SRP 2 A 5 ANIOS PRIMERA',
    'srp_6a':   'SRP 6 ANIOS PRIMERA',
    'srp_7_9':  'SRP 7 A 9 ANIOS PRIMERA',
    'srp_10_12':'SRP 10 A 12 ANIOS PRIMERA',
    'srp_13_19':'SRP 13 A 19 ANIOS PRIMERA',
    'srp_20_29':'SRP 20 A 29 ANIOS PRIMERA',
    'srp_30_39':'SRP 30 A 39 ANIOS PRIMERA',
    'srp_40_49':'SRP 40 A 49 ANIOS PRIMERA',
    'srp_salud':'SRP PERSONAL DE SALUD PRIMERA',
    'srp_educ': 'SRP PERSONAL EDUCATIVO PRIMERA',
    'srp_jorn': 'SRP JORNALEROS AGRICOLAS PRIMERA',
    'srp_pt':   'SRP  PRIMERA TOTAL',
    'srp_18m':  'SRP 18 MESES SEGUNDA',
    'srp_2_5_2':'SRP 2 A 5 ANIOS SEGUNDA',
    'srp_st':   'SRP SEGUNDA TOTAL',
    'sr_6_11':  'SR 6 A 11 MESES PRIMERA',
    'sr_1a':    'SR 1 ANIO PRIMERA',
    'sr_2_5':   'SR 2 A 5 ANIOS PRIMERA',
    'sr_6a':    'SR 6 ANIOS PRIMERA',
    'sr_7_9':   'SR 7 A 9 ANIOS PRIMERA',
    'sr_10_12': 'SR 10 A 12 ANIOS PRIMERA',
    'sr_13_19': 'SR 13 A 19 ANIOS PRIMERA',
    'sr_20_29': 'SR 20 A 29 ANIOS PRIMERA',
    'sr_30_39': 'SR 30 A 39 ANIOS PRIMERA',
    'sr_40_49': 'SR 40 A 49 ANIOS PRIMERA',
    'sr_salud': 'SR PERSONAL DE SALUD PRIMERA',
    'sr_educ':  'SR PERSONAL EDUCATIVO PRIMERA',
    'sr_jorn':  'SR JORNALEROS AGRICOLAS PRIMERA',
    'sr_pt':    'SR PRIMERA TOTAL',
    'sr_18m':   'SR 18 MESES SEGUNDA',
    'sr_st':    'SR SEGUNDA TOTAL',
}

JURS = ['DURANGO','GOMEZ PALACIO','SANTIAGO PAPASQUIARO','RODEO']
JL   = {'DURANGO':'Durango','GOMEZ PALACIO':'Gómez Palacio','SANTIAGO PAPASQUIARO':'Santiago Papasquiaro','RODEO':'Rodeo'}

da  = dur26[dur26['SEMANA']==51]
dant= dur26[dur26['SEMANA']==50]
dacu= dur26[dur26['SEMANA']<=51]
daa = dur26[dur26['SEMANA']<=50]

def v(d,k): return int(d[C[k]].sum())
def vj(d,j,k): return int(d[d['JURISDICCION']==j][C[k]].sum())
def n(x): return f"{x:,}"

# ============================================================
# 2. ABRIR DOCUMENTO Y ELIMINAR IMÁGENES
# ============================================================
doc = Document(r'C:\Users\aicil\.gemini\antigravity\scratch\informe_original.docx')

# Eliminar imágenes inline
for para in doc.paragraphs:
    for run in para.runs:
        for tag in [qn('w:drawing'), qn('w:pict')]:
            for el in run._element.findall('.//' + tag):
                el.getparent().remove(el)

# Eliminar párrafos vacíos tras eliminar imágenes (limpiar)
for para in doc.paragraphs:
    if para.text.strip() == '' and len(para.runs) == 0:
        pass  # los dejamos, no rompemos el layout

# ============================================================
# 3. HELPERS FORMATO
# ============================================================
def shd(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), hex_color)
    tcPr.append(s)

def ct(cell, text, bold=False, sz=9, col=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(sz)
    if col:
        r.font.color.rgb = RGBColor(*col)

HDR = '1F497D'; SUB='2E75B6'; TOT='BDD7EE'; ALT='DEEAF1'; WHT='FFFFFF'

def make_table(doc, headers, data_rows, totals=None, col_widths=None):
    n_rows = 1 + len(data_rows) + (1 if totals else 0)
    table = doc.add_table(rows=n_rows, cols=len(headers))
    table.style = 'Table Grid'
    # Header
    for i, h in enumerate(headers):
        shd(table.cell(0, i), HDR)
        ct(table.cell(0, i), h, bold=True, col=(255,255,255), sz=8)
    # Data
    for ri, row in enumerate(data_rows):
        bg = ALT if ri % 2 == 0 else WHT
        for ci, val in enumerate(row):
            shd(table.cell(ri+1, ci), bg)
            al = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            ct(table.cell(ri+1, ci), val, align=al, sz=8)
    # Totals
    if totals:
        ri = len(data_rows) + 1
        for ci, val in enumerate(totals):
            shd(table.cell(ri, ci), TOT)
            al = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            ct(table.cell(ri, ci), val, bold=True, align=al, sz=8)
    return table

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False, italic=False, sz=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(sz)
    return p

# ============================================================
# 4. ACTUALIZAR TABLAS EXISTENTES
# ============================================================
tables = doc.tables

# TABLA 0 (Tabla 1 del doc): Dosis Cero comparativo 4 jurisdicciones
# Encabezados actuales: Jurisdicción | Corte 10/abr | Corte 17/abr | Variación
# Nuevos: Jurisdicción | Corte 17/abr | Corte 24/abr | Variación
t0 = tables[0]
# Fila 1 = headers
try:
    t0.cell(1,1).paragraphs[0].runs[0].text = 'Corte 17/abr'
    t0.cell(1,2).paragraphs[0].runs[0].text = 'Corte 24/abr'
except: pass

# Filas de datos (F2=Durango, F3=GomPal, F4=StgPap, F5=Rodeo, F6=Estatal, F7=Acumulado?)
jur_map = {2:'DURANGO', 3:'GOMEZ PALACIO', 4:'SANTIAGO PAPASQUIARO', 5:'RODEO'}
for fi, jur in jur_map.items():
    ant_v = vj(dant, jur, 'srp_6_11')
    act_v = vj(da,   jur, 'srp_6_11')
    dif_v = act_v - ant_v
    try:
        for ci, val in enumerate([JL[jur], n(ant_v), n(act_v), f"{'+' if dif_v>=0 else ''}{n(dif_v)}"]):
            cell = t0.cell(fi, ci)
            cell.text = val
    except: pass

# Fila total estatal
try:
    ant_e = v(dant,'srp_6_11'); act_e = v(da,'srp_6_11'); dif_e = act_e-ant_e
    for ci, val in enumerate(['Total Estatal', n(ant_e), n(act_e), f"{'+' if dif_e>=0 else ''}{n(dif_e)}"]):
        t0.cell(6, ci).text = val
except: pass

# Fila acumulado
try:
    acum_ant_v = v(daa,'srp_6_11'); acum_act_v = v(dacu,'srp_6_11')
    for ci, val in enumerate(['Acumulado 2026', n(acum_ant_v), n(acum_act_v), f"+{n(acum_act_v-acum_ant_v)}"]):
        t0.cell(7, ci).text = val
except: pass

# ============================================================
# 5. ACTUALIZAR TEXTO DE PÁRRAFOS CLAVE
# ============================================================
REEMPLAZOS = [
    ('Corte: 24 de abril de 2026', 'Corte: 24 de abril de 2026'),
    ('corte al 17 de abril', 'corte al 24 de abril'),
    ('17 de abril de 2026', '24 de abril de 2026'),
    ('10 de abril', '17 de abril'),
    ('10,534', n(v(dacu,'srp_6_11'))),
    ('10,343', n(v(daa,'srp_6_11'))),
    ('19,507', n(v(dacu,'srp_18m'))),
    ('19,267', n(v(daa,'srp_18m'))),
    ('semana epidemiológica 16', 'semana epidemiológica 17'),
    ('semana 16', 'semana 17'),
    ('Sem. 15', 'Sem. 16'),
    ('Sem. 16', 'Sem. 17'),
]

for p in doc.paragraphs:
    for old, new in REEMPLAZOS:
        if old in p.text:
            for run in p.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)

# ============================================================
# 6. REGENERAR TABLAS RESTANTES CON DATOS NUEVOS
# ============================================================
# Borramos contenido de tablas 1-4 y las reemplazamos con datos actualizados
# Esto es más seguro que intentar editar celda por celda con estructura desconocida

def clear_and_fill_table(table, headers, data_rows, totals=None):
    """Rellena una tabla existente con nuevos datos"""
    # Limpiar todas las celdas
    all_rows = list(table.rows)
    n_existing = len(all_rows)
    n_needed = 1 + len(data_rows) + (1 if totals else 0)

    # Rellenar encabezados
    if n_existing > 0:
        for ci, h in enumerate(headers):
            if ci < len(all_rows[0].cells):
                shd(all_rows[0].cells[ci], HDR)
                ct(all_rows[0].cells[ci], h, bold=True, col=(255,255,255), sz=8)

    # Rellenar filas de datos
    for ri, row_data in enumerate(data_rows):
        if ri+1 < n_existing:
            row = all_rows[ri+1]
            bg = ALT if ri % 2 == 0 else WHT
            for ci, val in enumerate(row_data):
                if ci < len(row.cells):
                    shd(row.cells[ci], bg)
                    al = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
                    ct(row.cells[ci], val, align=al, sz=8)

    # Totals
    if totals and n_existing > len(data_rows)+1:
        tr = all_rows[len(data_rows)+1]
        for ci, val in enumerate(totals):
            if ci < len(tr.cells):
                shd(tr.cells[ci], TOT)
                al = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
                ct(tr.cells[ci], val, bold=True, align=al, sz=8)

# Tabla 1 (idx 1): Primera Dosis SRP por Jurisdicción - comparativo intersemanal
if len(tables) > 1:
    t1_headers = ['Jurisdicción', 'Sem. Ant.\n(11-17 abr)', 'Sem. Act.\n(18-24 abr)', 'Diferencia', 'Acumulado 2026']
    t1_data = []
    for jur in JURS:
        ant_v = vj(dant, jur, 'srp_pt')
        act_v = vj(da,   jur, 'srp_pt')
        acu_v = vj(dacu, jur, 'srp_pt')
        dif_v = act_v - ant_v
        t1_data.append([JL[jur], n(ant_v), n(act_v), f"{'+' if dif_v>=0 else ''}{n(dif_v)}", n(acu_v)])
    ant_e=v(dant,'srp_pt'); act_e=v(da,'srp_pt'); acu_e=v(dacu,'srp_pt')
    t1_tot = ['Total Estatal', n(ant_e), n(act_e), f"{'+' if act_e-ant_e>=0 else ''}{n(act_e-ant_e)}", n(acu_e)]
    clear_and_fill_table(tables[1], t1_headers, t1_data, t1_tot)

# Tabla 2 (idx 2): Primera Dosis SRP por Grupo de Edad
if len(tables) > 2:
    grp_headers = ['Jurisdicción','6-11m','1 año','2-5a','6a','7-9a','10-12a','13-19a','20-29a','30-39a','40-49a','P.Salud','P.Educ','Jorn.','Total']
    grp_keys = ['srp_6_11','srp_1a','srp_2_5','srp_6a','srp_7_9','srp_10_12','srp_13_19','srp_20_29','srp_30_39','srp_40_49','srp_salud','srp_educ','srp_jorn','srp_pt']
    t2_data = []
    for jur in JURS:
        row = [JL[jur]] + [n(vj(da, jur, k)) for k in grp_keys]
        t2_data.append(row)
    t2_tot = ['Total Estatal'] + [n(v(da, k)) for k in grp_keys]
    clear_and_fill_table(tables[2], grp_headers, t2_data, t2_tot)

# Tabla 3 (idx 3): Segunda Dosis SRP 18m comparativo
if len(tables) > 3:
    t3_headers = ['Jurisdicción', 'Sem. Ant.\n(11-17 abr)', 'Sem. Act.\n(18-24 abr)', 'Diferencia', 'Acumulado 2026']
    t3_data = []
    for jur in JURS:
        ant_v = vj(dant, jur, 'srp_18m')
        act_v = vj(da,   jur, 'srp_18m')
        acu_v = vj(dacu, jur, 'srp_18m')
        dif_v = act_v - ant_v
        t3_data.append([JL[jur], n(ant_v), n(act_v), f"{'+' if dif_v>=0 else ''}{n(dif_v)}", n(acu_v)])
    ant_e=v(dant,'srp_18m'); act_e=v(da,'srp_18m'); acu_e=v(dacu,'srp_18m')
    t3_tot = ['Total Estatal', n(ant_e), n(act_e), f"{'+' if act_e-ant_e>=0 else ''}{n(act_e-ant_e)}", n(acu_e)]
    clear_and_fill_table(tables[3], t3_headers, t3_data, t3_tot)

# Tabla 4 (idx 4): SR comparativo
if len(tables) > 4:
    t4_headers = ['Jurisdicción', 'SR 1a Ant.', 'SR 1a Act.', 'Dif. 1a', 'SR 2a Ant.', 'SR 2a Act.', 'Dif. 2a', 'Acum. SR Total']
    t4_data = []
    for jur in JURS:
        sr1a = vj(dant,jur,'sr_pt'); sr1b = vj(da,jur,'sr_pt')
        sr2a = vj(dant,jur,'sr_st'); sr2b = vj(da,jur,'sr_st')
        acu  = vj(dacu,jur,'sr_pt') + vj(dacu,jur,'sr_st')
        t4_data.append([JL[jur], n(sr1a), n(sr1b), f"{'+' if sr1b-sr1a>=0 else ''}{n(sr1b-sr1a)}",
                         n(sr2a), n(sr2b), f"{'+' if sr2b-sr2a>=0 else ''}{n(sr2b-sr2a)}", n(acu)])
    sr1_tot_a=v(dant,'sr_pt'); sr1_tot_b=v(da,'sr_pt')
    sr2_tot_a=v(dant,'sr_st'); sr2_tot_b=v(da,'sr_st')
    acu_tot = v(dacu,'sr_pt') + v(dacu,'sr_st')
    t4_tot = ['Total Estatal', n(sr1_tot_a), n(sr1_tot_b),
              f"{'+' if sr1_tot_b-sr1_tot_a>=0 else ''}{n(sr1_tot_b-sr1_tot_a)}",
              n(sr2_tot_a), n(sr2_tot_b),
              f"{'+' if sr2_tot_b-sr2_tot_a>=0 else ''}{n(sr2_tot_b-sr2_tot_a)}", n(acu_tot)]
    clear_and_fill_table(tables[4], t4_headers, t4_data, t4_tot)

# ============================================================
# 7. GUARDAR DOCUMENTO FINAL
# ============================================================
OUTPUT = r'C:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\CONASABI\EVIDENCIAS CONASABI_24ABRIL2026\DSP_Vacunación_CONASABIAc2_24abril2026_ACTUALIZADO.docx'
doc.save(OUTPUT)
print(f"✅ Documento guardado en:\n{OUTPUT}")
print(f"\nRESUMEN DE DATOS CLAVE - CORTE 24 ABRIL 2026")
print(f"  Dosis cero SRP (6-11m) semana actual: {n(v(da,'srp_6_11'))}")
print(f"  Dosis cero SRP (6-11m) acumulado:     {n(v(dacu,'srp_6_11'))}")
print(f"  SRP 1a dosis semana actual:            {n(v(da,'srp_pt'))}")
print(f"  SRP 1a dosis acumulado:                {n(v(dacu,'srp_pt'))}")
print(f"  SRP 2a dosis (18m) semana actual:      {n(v(da,'srp_18m'))}")
print(f"  SRP 2a dosis (18m) acumulado:          {n(v(dacu,'srp_18m'))}")
print(f"  SR 1a dosis semana actual:             {n(v(da,'sr_pt'))}")
print(f"  SR 1a dosis acumulado:                 {n(v(dacu,'sr_pt'))}")
print(f"  SR 2a dosis semana actual:             {n(v(da,'sr_st'))}")
print(f"  SR 2a dosis acumulado:                 {n(v(dacu,'sr_st'))}")
print(f"  GRAN TOTAL ACUMULADO:                  {n(v(dacu,'srp_pt')+v(dacu,'srp_st')+v(dacu,'sr_pt')+v(dacu,'sr_st'))}")
