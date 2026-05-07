import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def update_report():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('INFORME DE VACUNACIÓN - MUNICIPIO DE EL MEZQUITAL', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('CORTE AL 04 DE MAYO DE 2026')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    source_text = 'Fuente: SIS/CeNSIA, consultado el 04 de mayo de 2026; 19:00 hs.'
    
    # Section 1: 2025
    doc.add_heading('1. DOSIS APLICADAS DURANTE 2025', level=1)
    doc.add_paragraph(source_text)
    
    data_2025 = [
        ['Valores', 'IMSS B', 'SSA', 'Suma total'],
        ['SRP PRIMERA TOTAL', '587', '1,043', '1,630'],
        ['SRP SEGUNDA TOTAL', '1,526', '1,255', '2,781'],
        ['SR PRIMERA TOTAL', '495', '645', '1,140'],
        ['SR SEGUNDA TOTAL', '704', '254', '958'],
        ['TOTAL', '3,312', '3,197', '6,509']
    ]
    t1 = doc.add_table(rows=6, cols=4)
    t1.style = 'Table Grid'
    for i, row in enumerate(data_2025):
        for j, text in enumerate(row):
            t1.rows[i].cells[j].text = text

    # Section 2: 2026
    doc.add_heading('2. DOSIS APLICADAS DURANTE 2026', level=1)
    doc.add_paragraph(source_text)
    
    data_2026 = [
        ['Valores', 'IMSS B', 'SSA', 'Suma total'],
        ['SRP PRIMERA TOTAL', '121', '2,063', '2,184'],
        ['SRP SEGUNDA TOTAL', '229', '6,448', '6,677'],
        ['SR PRIMERA TOTAL', '2,630', '1,201', '3,831'],
        ['SR SEGUNDA TOTAL', '207', '1,569', '1,776'],
        ['TOTAL', '3,187', '11,281', '14,468']
    ]
    t2 = doc.add_table(rows=6, cols=4)
    t2.style = 'Table Grid'
    for i, row in enumerate(data_2026):
        for j, text in enumerate(row):
            t2.rows[i].cells[j].text = text

    # Section 3: 2025-2026 COMBINED
    doc.add_heading('3. DOSIS APLICADAS ACUMULADAS 2025 - 2026', level=1)
    doc.add_paragraph(source_text)
    
    # Calculate Sums
    # SRP1: 1630 + 2184 = 3814
    # SRP2: 2781 + 6677 = 9458
    # SR1: 1140 + 3831 = 4971
    # SR2: 958 + 1776 = 2734
    # TOTAL: 6509 + 14468 = 20977
    
    data_combined = [
        ['Valores', 'IMSS B', 'SSA', 'Suma total'],
        ['SRP PRIMERA TOTAL', '708', '3,106', '3,814'],
        ['SRP SEGUNDA TOTAL', '1,755', '7,703', '9,458'],
        ['SR PRIMERA TOTAL', '3,125', '1,846', '4,971'],
        ['SR SEGUNDA TOTAL', '911', '1,823', '2,734'],
        ['TOTAL', '6,499', '14,478', '20,977']
    ]
    t3 = doc.add_table(rows=6, cols=4)
    t3.style = 'Table Grid'
    for i, row in enumerate(data_combined):
        for j, text in enumerate(row):
            t3.rows[i].cells[j].text = text

    # Section 4: Heatmap
    doc.add_heading('4. MAPA DE CALOR: ALCANCE POR LOCALIDAD', level=1)
    doc.add_paragraph('El siguiente gráfico representa el alcance (%) de vacunación en las principales localidades del municipio.')
    
    img_path = r"C:\Users\aicil\.gemini\antigravity\scratch\heatmap_vacunacion.png"
    doc.add_picture(img_path, width=Inches(6.0))
    
    # Section 5: Details
    doc.add_heading('5. DETALLE DE COBERTURA POR LOCALIDAD', level=1)
    doc.add_paragraph('Basado en censos INEGI 2020 y dosis aplicadas durante el periodo.')
    
    # Add a summary of top localities
    doc.add_paragraph('Localidades con mayor alcance documentado:', style='List Bullet')
    doc.add_paragraph('Las Joyas (Alcance > 100%)', style='List Bullet')
    doc.add_paragraph('Aguacates (Angostura) (Alcance > 100%)', style='List Bullet')
    doc.add_paragraph('Cerro Bolillo (Alcance > 100%)', style='List Bullet')

    output_path = r"C:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\mezquital\INFORMES\Informe_Vacunacion_Mezquital_2026_Final_Heatmap.docx"
    doc.save(output_path)
    print(f"Final report saved to {output_path}")

update_report()
