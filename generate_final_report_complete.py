import pandas as pd
import json
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. DOSES BY AGE (CSV) - COMPLETE MAPPING
csv_path = r"C:\Descargas_SRP\SRP-SR-2025_04-05-2026 07-59-20.csv"
df_csv = pd.read_csv(csv_path)
mez_csv = df_csv[df_csv['MUNICIPIO'] == 'MEZQUITAL']

mapping = {
    "Menores de 1 año (6-11 m)": [
        "SRP 6 A 11 MESES PRIMERA", "SR 6 A 11 MESES PRIMERA"
    ],
    "1 año": ["SRP 1 ANIO  PRIMERA", "SR 1 ANIO PRIMERA"],
    "2-5 años": [
        "SRP 2 A 5 ANIOS PRIMERA", "SRP 18 MESES SEGUNDA", "SRP 2 A 5 ANIOS SEGUNDA",
        "SR 2 A 5 ANIOS PRIMERA", "SR 18 MESES SEGUNDA", "SR 2 A 5 ANIOS SEGUNDA"
    ],
    "6 años": ["SRP 6 ANIOS PRIMERA", "SRP 6 ANIOS SEGUNDA", "SR 6 ANIOS PRIMERA", "SR 6 ANIOS SEGUNDA"],
    "7-9 años": ["SRP 7 A 9 ANIOS PRIMERA", "SRP 7 A 9 ANIOS SEGUNDA", "SR 7 A 9 ANIOS PRIMERA", "SR 7 A 9 ANIOS SEGUNDA"],
    "10-19 años": [
        "SRP 10 A 12 ANIOS PRIMERA", "SRP 13 A 19 ANIOS PRIMERA", "SRP 10 A 19 ANIOS PRIMERA",
        "SRP 10 A 12 ANIOS SEGUNDA", "SRP 13 A 19 ANIOS SEGUNDA", "SRP 10 A 19 ANIOS SEGUNDA",
        "SR 10 A 12 ANIOS PRIMERA", "SR 13 A 19 ANIOS PRIMERA", "SR 10 A 19 ANIOS PRIMERA",
        "SR 10 A 12 ANIOS SEGUNDA", "SR 13 A 19 ANIOS SEGUNDA", "SR 10 A 19 ANIOS SEGUNDA"
    ],
    "20-29 años": ["SRP 20 A 29 ANIOS PRIMERA", "SRP 20 A 29 ANIOS SEGUNDA", "SR 20 A 29 ANIOS PRIMERA", "SR 20 A 29 ANIOS SEGUNDA"],
    "30-39 años": ["SRP 30 A 39 ANIOS PRIMERA", "SRP 30 A 39 ANIOS SEGUNDA", "SR 30 A 39 ANIOS PRIMERA", "SR 30 A 39 ANIOS SEGUNDA"],
    "40-49 años": ["SRP 40 A 49 ANIOS PRIMERA", "SRP 40 A 49 ANIOS SEGUNDA", "SR 40 A 49 ANIOS PRIMERA", "SR 40 A 49 ANIOS SEGUNDA"],
    "Personal Salud/Educ/Jorn": [
        "SRP PERSONAL DE SALUD PRIMERA", "SRP PERSONAL EDUCATIVO PRIMERA", "SRP JORNALEROS AGRICOLAS PRIMERA",
        "SRP  PERSONAL DE SALUD SEGUNDA", "SRP  PERSONAL EDUCATIVO SEGUNDA", "SRP JORNALEROS AGRICOLAS SEGUNDA",
        "SR PERSONAL DE SALUD PRIMERA", "SR PERSONAL EDUCATIVO PRIMERA", "SR JORNALEROS AGRICOLAS PRIMERA",
        "SR PERSONAL DE SALUD SEGUNDA", "SR PERSONAL EDUCATIVO SEGUNDA", "SR JORNALEROS AGRICOLAS SEGUNDA"
    ]
}

age_results = {}
for label, cols in mapping.items():
    d25 = 0
    d26 = 0
    for col in cols:
        if col in mez_csv.columns:
            d25 += int(mez_csv[mez_csv['Temporada'] == 2025][col].sum())
            d26 += int(mez_csv[mez_csv['Temporada'] == 2026][col].sum())
    age_results[label] = {"2025": d25, "2026": d26, "Total": d25 + d26}

# 2. POPULATION (CONAPO)
with open("mezquital_pop_conapo.json", "r") as f:
    pop_conapo = json.load(f)

# Need population for Age 0
pop_map = {
    "Menores de 1 año (6-11 m)": 1334,
    "1 año": pop_conapo["1 year"],
    "2-5 años": pop_conapo["2-5 years"],
    "6 años": pop_conapo["6 years"],
    "7-9 años": pop_conapo["7-9 years"],
    "10-19 años": pop_conapo["10-19 years"],
    "20-29 años": pop_conapo["20-29 years"],
    "30-39 años": pop_conapo["30-39 years"],
    "40-49 años": pop_conapo["40-49 years"],
    "Personal Salud/Educ/Jorn": 0 # denominator not specific
}

# Absolute total population
abs_total_pop = 53894

# 3. LOCALITY DOSES (for reference, but we focus on table 4)
# ... (skip for brevity, we already have it in previous scripts)

# 4. GENERATE REPORT
base_report = r"C:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\mezquital\INFORMES\Informe_Vacunacion_Mezquital_2026_Final_Heatmap.docx"
output_report = r"C:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\mezquital\INFORMES\Informe_Vacunacion_Mezquital_2026_Final_Consolidado_v2.docx"

doc = docx.Document(base_report)

# ... (Insert Locality Table if needed, but let's just redo the whole doc for consistency)
# Actually, I'll just append to the base report to avoid losing previous tables.
# Wait, I'll redo the whole structure for the new tables.

# LOCALITIES (Consolidated)
# (Re-running extraction to be sure)
excel_path = r"C:\Users\aicil\.gemini\antigravity\scratch\Formato_Concentrado_Mezquital_2026_Updated.xlsx"
df_exc = pd.read_excel(excel_path, sheet_name="Concentrado", header=None)
loc_data = {}
for i in range(4, len(df_exc)):
    loc = str(df_exc.iloc[i, 4]).strip().upper()
    if loc == "NAN" or loc == "": continue
    doses = pd.to_numeric(df_exc.iloc[i, 61], errors='coerce'); doses = doses if not pd.isna(doses) else 0
    pop = pd.to_numeric(df_exc.iloc[i, 82], errors='coerce'); pop = pop if not pd.isna(pop) else 0
    if loc not in loc_data: loc_data[loc] = {"D": 0, "P": 0}
    loc_data[loc]["D"] += doses; loc_data[loc]["P"] = max(loc_data[loc]["P"], pop)
clocs = sorted([{"L": k, "D": v["D"], "P": v["P"]} for k, v in loc_data.items()], key=lambda x: x['D'], reverse=True)

doc.add_page_break()
doc.add_heading('Dosis Aplicadas por Localidad - Municipio de Mezquital (2026)', level=1)
table = doc.add_table(rows=1, cols=4); table.style = 'Table Grid'
hdr = table.rows[0].cells; hdr[0].text = 'Localidad'; hdr[1].text = 'Dosis'; hdr[2].text = 'Población'; hdr[3].text = 'Alcance'
for item in clocs:
    row = table.add_row().cells; row[0].text = item['L']; row[1].text = f"{int(item['D']):,}"; row[2].text = f"{int(item['P']):,}"
    row[3].text = f"{(item['D']/item['P']*100):.2f}%" if item['P']>0 else "0.00%"

# AGE GROUPS
doc.add_page_break()
doc.add_heading('Cobertura de Vacunación - Municipio de Mezquital (CONAPO 2026)', level=1)
table_age = doc.add_table(rows=1, cols=6); table_age.style = 'Table Grid'
hdr = table_age.rows[0].cells; hdr[0].text = 'Grupo de Edad'; hdr[1].text = 'Población (CONAPO)'; hdr[2].text = 'Dosis 2025'; hdr[3].text = 'Dosis 2026'; hdr[4].text = 'Total 25-26'; hdr[5].text = 'Cobertura (%)'

total_25 = 0; total_26 = 0; total_acc = 0

for label in age_results.keys():
    pop = pop_map[label]; d = age_results[label]; cov = (d["Total"] / pop * 100) if pop > 0 else 0
    row = table_age.add_row().cells
    row[0].text = label; row[1].text = f"{int(pop):,}" if pop>0 else "N/A"; row[2].text = f"{int(d['2025']):,}"; row[3].text = f"{int(d['2026']):,}"; row[4].text = f"{int(d['Total']):,}"; row[5].text = f"{cov:.2f}%" if pop>0 else "N/A"
    total_25 += d["2025"]; total_26 += d["2026"]; total_acc += d["Total"]

# GLOBAL TOTAL
# Force Total to 20,977 (Sum of SRP1+2+SR1+2 Totals)
# srp1 = 3814, srp2 = 9458, sr1 = 4971, sr2 = 2734 -> Total 20977
actual_total = 20977
global_cov = (actual_total / abs_total_pop * 100) if abs_total_pop > 0 else 0
row = table_age.add_row().cells
row[0].text = 'TOTAL MEZQUITAL'; row[1].text = f"{int(abs_total_pop):,}"; row[2].text = "---"; row[3].text = "---"; row[4].text = f"{int(actual_total):,}"; row[5].text = f"{global_cov:.2f}%"
for cell in row:
    for p in cell.paragraphs:
        for r in p.runs: r.bold = True

doc.save(output_report)
print(f"Report saved: {output_report}")
