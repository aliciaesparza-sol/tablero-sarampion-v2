import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
import os

csv_path = 'reporte_cobertura_mezquital_final.csv'
docx_path = 'tarjeta_copy.docx'
output_docx = 'tarjeta_final_result.docx'

def insert_coverage_table():
    df = pd.read_csv(csv_path)
    doc = Document(docx_path)

    # Find "Atentamente" position
    atentamente_para = None
    for p in doc.paragraphs:
        if "Atentamente" in p.text:
            atentamente_para = p
            break
    
    if atentamente_para is None:
        atentamente_para = doc.paragraphs[-1]

    # Insert Title
    title_text = "Cuadro 3. Avance de Cobertura Vacunal contra Sarampión y Rubéola en Grupos de Riesgo y Población Objetivo, Municipio de Mezquital, 2026."
    p_title = atentamente_para.insert_paragraph_before(title_text)
    p_title.runs[0].bold = True
    p_title.runs[0].font.size = Pt(11)
    
    # Add Spacer
    atentamente_para.insert_paragraph_before("")

    # Add Table
    tbl_data = df[['Grupo', 'Meta', 'Dosis Total', 'Cobertura (%)']]
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Headers
    hdr_cells = table.rows[0].cells
    headers = ['Grupo Etario', 'Meta (Poblaci\u00f3n 2026)', 'Dosis Aplicadas (Total)', 'Cobertura (%)']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    # Data
    for _, row in tbl_data.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Grupo'])
        row_cells[1].text = f"{int(row['Meta']):,}"
        row_cells[2].text = f"{int(row['Dosis Total']):,}"
        row_cells[3].text = f"{row['Cobertura (%)']}%"
        for i in range(1, 4):
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Move table XML before "Atentamente"
    atentamente_para._element.addprevious(table._element)
    
    # Add Source
    source_text = "Fuente: Elaboraci\u00f3n propia de la Secretar\u00eda de Salud con datos de Proyecciones de Poblaci\u00f3n CENJSIA 2026 y reportes institucionales de productividad."
    p_source = atentamente_para.insert_paragraph_before(source_text)
    p_source.runs[0].font.size = Pt(9)
    p_source.runs[0].italic = True
    
    # Add whitespace after source
    atentamente_para.insert_paragraph_before("")

    doc.save(output_docx)
    print(f"Document saved to {output_docx}")

if __name__ == "__main__":
    insert_coverage_table()
