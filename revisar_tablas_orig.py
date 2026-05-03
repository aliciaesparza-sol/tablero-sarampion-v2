
# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

doc = Document(r'C:\Users\aicil\.gemini\antigravity\scratch\informe_original.docx')

print("=== TABLAS DEL DOCUMENTO ORIGINAL ===\n")
for ti, t in enumerate(doc.tables):
    print(f"\n--- TABLA {ti} ({len(t.rows)} filas x {len(t.columns)} cols) ---")
    for ri, row in enumerate(t.rows):
        cells = [c.text.strip()[:80] for c in row.cells]
        print(f"  F{ri}: {cells}")

print("\n\n=== PÁRRAFOS CERCANOS A TABLAS (títulos y fuentes) ===")
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt and ('Tabla' in txt or 'Fuente' in txt or 'tabla' in txt or 'fuente' in txt or 'Elaboración' in txt or 'elaboración' in txt or 'Nota' in txt):
        print(f"  [{i}] ESTILO={p.style.name!r} | {txt[:300]}")
