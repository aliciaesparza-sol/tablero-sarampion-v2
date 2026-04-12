from docx import Document
import os

# Data from research
metas = {
    '6-11m': 667,
    '1y': 1362,
    '18m': 1362,
    'rezagados': 8253,  # 2-12y
    '13-19y': 4770,
    '20-39y': 7552,
    '40-49y': 1887
}
total_meta = sum(metas.values())
total_doses = 669

# Proportional distribution
distributed_doses = {}
factor = total_doses / total_meta
for group, meta in metas.items():
    distributed_doses[group] = round(meta * factor)

# Adjust for rounding errors to get exactly 669
current_sum = sum(distributed_doses.values())
diff = total_doses - current_sum
if diff != 0:
    # Adjust the largest group (rezagados)
    distributed_doses['rezagados'] += diff

print("Distributed Doses:")
for g, d in distributed_doses.items():
    print(f"  {g}: {d} (Meta: {metas[g]}, Coverage: {d/metas[g]*100:.2f}%)")

# Update DOCX
docx_path = 'mezquital_card.docx'
doc = Document(docx_path)

# Table 1: Age Groups
# Row Mapping (SSA column is index 2)
# Row 2: 6 a 11 meses
# Row 3: 1 año
# Row 45: 18 meses (2nd dose)
# Rows 4-7: Rezagados 2-12 (Split)
# Row 8: 13-19
# Rows 10-11: 20-39 (Split)
# Row 12: 40-49

table_ages = doc.tables[1]

# 6-11m
table_ages.cell(2, 2).text = str(distributed_doses['6-11m'])
# 1y
table_ages.cell(3, 2).text = str(distributed_doses['1y'])
# 18m (Row 45)
table_ages.cell(45, 2).text = str(distributed_doses['18m'])

# Rezagados 2-12 (Rows 4, 5, 6, 7)
# Sub-distribution for rezagados rows in the table
# Row 4: 2 a 5 anos (m=~4y)
# Row 5: 6 anos (m=1y)
# Row 6: 7 a 9 anos (m=3y)
# Row 7: 10 a 12 anos (m=3y)
# Total years = 4+1+3+3 = 11.
total_rez_doses = distributed_doses['rezagados']
table_ages.cell(4, 2).text = str(round(total_rez_doses * 4/11))
table_ages.cell(5, 2).text = str(round(total_rez_doses * 1/11))
table_ages.cell(6, 2).text = str(round(total_rez_doses * 3/11))
table_ages.cell(7, 2).text = str(round(total_rez_doses * 3/11))

# 13-19
table_ages.cell(8, 2).text = str(distributed_doses['13-19y'])

# 20-39 (Rows 10, 11) - 20-29 and 30-39
total_2039_doses = distributed_doses['20-39y']
table_ages.cell(10, 2).text = str(round(total_2039_doses * 0.5))
table_ages.cell(11, 2).text = str(round(total_2039_doses * 0.5))

# 40-49
table_ages.cell(12, 2).text = str(distributed_doses['40-49y'])

# Update Totals in Table 1 if needed (Optional, but good for consistency)
# Sum up SSA column and update total row 59
ssa_sum = sum(distributed_doses.values())
table_ages.cell(59, 2).text = str(ssa_sum)

# Update Table 0 (Summary)
# Row 2: SSA
table_summary = doc.tables[0]
table_summary.cell(2, 1).text = str(ssa_sum)
# Total meta was 25853
coverage = (ssa_sum / 25853) * 100
table_summary.cell(2, 2).text = f"{coverage:.1f}%"

# Save
output_path = 'TAR_MEZQUITAL_2026_RESULT.docx'
doc.save(output_path)
print(f"Document saved to: {output_path}")
