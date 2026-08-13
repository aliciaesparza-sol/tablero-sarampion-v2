# -*- coding: utf-8 -*-
"""
INFORME DIARIO AUTOMATIZADO - VACUNACION SARAMPION DURANGO
==========================================================
9 rubros exactos del PDF de referencia | Hoja membretada | Gráfica
Firma: CSC/KLAR/AJEA
Envío automático diario a las 7:00 AM
"""

import sys, io, os, json, glob, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

import pandas as pd
from io import StringIO
from datetime import datetime, date, timedelta
import locale
import time

# ═══ CONFIGURACIÓN ═══════════════════════════════════════════════
BASE_DIR        = os.path.dirname(os.path.abspath(__file__))
CONFIG_PATH     = os.path.join(BASE_DIR, "config_correo.json")

# En GitHub Actions usamos la carpeta 'salida/' relativa al repo.
# En Windows local seguimos usando OneDrive si existe, sino 'salida/' local.
_CARPETA_LOCAL  = r"C:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\INFORME AUTOMATIZADO SARAMPION"
_CARPETA_CI     = os.path.join(BASE_DIR, "salida")
CARPETA_INFORME = _CARPETA_LOCAL if os.path.isdir(os.path.dirname(_CARPETA_LOCAL)) else _CARPETA_CI

# Plantilla membretada: busca en el repositorio, luego en OneDrive local.
_PLANTILLA_LOCAL = r"C:\Users\aicil\OneDrive\Escritorio\hoja MEMBRETADA GIGANTE2026_carta.docx"
_PLANTILLA_REPO  = os.path.join(BASE_DIR, "assets", "plantilla_membretada.docx")
PLANTILLA_DOCX  = _PLANTILLA_REPO if os.path.exists(_PLANTILLA_REPO) else (
                   _PLANTILLA_LOCAL if os.path.exists(_PLANTILLA_LOCAL) else None)

LOG_PATH        = os.path.join(BASE_DIR, "log_informe_diario.txt")
CENSIA_URL      = "https://siscensia.salud.gob.mx/sarampion_2025/"
# Credenciales: primero variables de entorno (GitHub Secrets), luego valores locales
CENSIA_USER     = os.environ.get("CENSIA_USER", "E_DGO_ADMIN")
CENSIA_PASS     = os.environ.get("CENSIA_PASS", "QWERTY")
FIRMA           = "CSC/KLAR/AJEA"

JURS = ['DURANGO', 'GOMEZ PALACIO', 'SANTIAGO PAPASQUIARO', 'RODEO']
JL   = {
    'DURANGO':              'Durango',
    'GOMEZ PALACIO':        'Gómez Palacio',
    'SANTIAGO PAPASQUIARO': 'Santiago\nPapasquiaro',
    'RODEO':                'Rodeo',
}
INSTITUCIONES = ['SSA', 'IMSS', 'IMSS-BIENESTAR', 'ISSSTE', 'SEDENA']
INST_LABEL    = {
    'SSA':            'SSD',
    'IMSS':           'IMSS',
    'IMSS-BIENESTAR': 'IMSS\nBienestar',
    'ISSSTE':         'ISSSTE',
    'SEDENA':         'SEDENA',
}
INST_LABEL_CHART = {
    'SSA':            'SSD',
    'IMSS':           'IMSS',
    'IMSS-BIENESTAR': 'IMSS Bienestar',
    'ISSSTE':         'ISSSTE',
    'SEDENA':         'SEDENA',
}
GRUPOS_EDAD = [
    ('6 a 11 meses',        'SRP 6 A 11 MESES',         'SR 6 A 11 MESES'),
    ('1 año',               'SRP 1 ANIO ',               'SR 1 ANIO'),
    ('18 meses',            'SRP 18 MESES',              'SR 18 MESES'),
    ('2 a 5 años',          'SRP 2 A 5 ANIOS',           'SR 2 A 5 ANIOS'),
    ('6 años',              'SRP 6 ANIOS',               'SR 6 ANIOS'),
    ('7 a 9 años',          'SRP 7 A 9 ANIOS',           'SR 7 A 9 ANIOS'),
    ('10 a 12 años',        'SRP 10 A 12 ANIOS',         'SR 10 A 12 ANIOS'),
    ('13 a 19 años',        'SRP 13 A 19 ANIOS',         'SR 13 A 19 ANIOS'),
    ('10 a 19 años',        'SRP 10 A 19 ANIOS',         'SR 10 A 19 ANIOS'),
    ('20 a 29 años',        'SRP 20 A 29 ANIOS',         'SR 20 A 29 ANIOS'),
    ('30 a 39 años',        'SRP 30 A 39 ANIOS',         'SR 30 A 39 ANIOS'),
    ('40 a 49 años',        'SRP 40 A 49 ANIOS',         'SR 40 A 49 ANIOS'),
    ('Personal de salud',   'SRP PERSONAL DE SALUD',     'SR PERSONAL DE SALUD'),
    ('Personal educativo',  'SRP PERSONAL EDUCATIVO',    'SR PERSONAL EDUCATIVO'),
    ('Jornaleros agrícolas','SRP JORNALEROS AGRICOLAS',  'SR JORNALEROS AGRICOLAS'),
]
SIN_2A_DOSIS_SRP = ('SRP 6 A 11 MESES', 'SRP 1 ANIO ')
SIN_1A_DOSIS_SRP = ('SRP 18 MESES',)
SIN_2A_DOSIS_SR  = ('SR 6 A 11 MESES', 'SR 1 ANIO')
SIN_1A_DOSIS_SR  = ('SR 18 MESES',)

# ═══ PALETA DE COLORES (AZUL MARINO Y BLANCO) ═══════════════════
C_PRIMARY   = '1B3A6B'   # Encabezados de tabla, títulos (Azul marino)
C_SECONDARY = '2E6DA4'   # Sub-títulos, subtablas (Azul medio)
C_ALT1      = 'D6E4F0'   # Fila par (Azul muy claro)
C_ALT2      = 'FFFFFF'   # Fila impar (Blanco)
C_TOTAL     = 'A9CCE3'   # Fila de totales (Azul claro)
C_SUBTITLE  = 'EBF5FB'   # Fondo sub-encabezados de vacuna (Azul pálido)

from docx.shared import RGBColor
HDR_COLOR = RGBColor( 27,  58, 107)   # #1B3A6B (Azul marino principal)
SUB_COLOR = RGBColor( 46, 109, 164)   # #2E6DA4 (Azul secundario)
BLK_COLOR = RGBColor(  0,   0,   0)
WHT_COLOR = RGBColor(255, 255, 255)
GRY_COLOR = RGBColor( 90,  90,  90)


# ══════════════════════════════════════════════════════════════════
# LOGGER
# ══════════════════════════════════════════════════════════════════
class Logger:
    def __init__(self, path):
        self.path = path

    def _write(self, msg):
        print(msg)
        try:
            with open(self.path, 'a', encoding='utf-8') as f:
                f.write(msg + '\n')
        except Exception:
            pass

    def section(self, icon, title):
        sep = '─' * 60
        self._write(sep)
        self._write(f'{icon} {title}')
        self._write(sep)

    def ok(self, msg):   self._write('   ✔ ' + msg)
    def info(self, msg): self._write('   ℹ  ' + msg)
    def warn(self, msg): self._write('   ⚠  ' + msg)
    def err(self, msg):  self._write('   ✗ ' + msg)
    def raw(self, msg):  self._write(msg)

log = Logger(LOG_PATH)


# ══════════════════════════════════════════════════════════════════
# PASO 1: DESCARGAR CSV
# ══════════════════════════════════════════════════════════════════
def descargar_csv_censia():
    log.section('📥', 'PASO 1: Descargando CSV de CeNSIA...')
    output_csv = os.path.join(BASE_DIR, "censia_descarga_hoy.csv")
    
    from playwright.sync_api import sync_playwright
    
    def attempt_once():
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            context = browser.new_context(accept_downloads=True)
            page    = context.new_page()
            
            page.goto(CENSIA_URL, wait_until="domcontentloaded", timeout=45000)
            try:
                page.fill("input[name='user']", CENSIA_USER, timeout=15000)
                page.fill("input[name='pass']", CENSIA_PASS)
                page.click("button[type='submit']")
                page.wait_for_load_state("domcontentloaded", timeout=15000)
            except Exception:
                pass
                
            page.goto(
                "https://siscensia.salud.gob.mx/sarampion_2025/ssa/reporte.php",
                wait_until="domcontentloaded", timeout=45000)
                
            page.wait_for_selector("button#descarga_todos, a#descarga_todos", timeout=15000)
            
            with page.expect_download(timeout=180000) as dl:
                clicked = False
                for sel in ["button#descarga_todos", "a#descarga_todos",
                            "button:has-text('Descargar')", "a:has-text('CSV')"]:
                    try:
                        el = page.query_selector(sel)
                        if el:
                            el.click()
                            clicked = True
                            break
                    except Exception:
                        continue
                if not clicked:
                    raise Exception("Could not find or click download button")
                    
            dl.value.save_as(output_csv)
            browser.close()
            return True

    max_retries = 3
    success = False
    last_err = None
    for attempt in range(1, max_retries + 1):
        try:
            log.raw(f'   Navegando e iniciando sesión (Intento {attempt}/{max_retries})...')
            if attempt_once():
                success = True
                break
        except Exception as e:
            last_err = e
            log.warn(f"Intento {attempt} falló: {e}")
            if attempt < max_retries:
                time.sleep(5)
                
    if success:
        size = os.path.getsize(output_csv)
        log.ok(f"CSV descargado: {size // 1024} KB → {output_csv}")
        return output_csv
    else:
        log.warn(f"No se pudo descargar de CeNSIA después de {max_retries} intentos: {last_err}")
        log.raw('   Buscando CSV más reciente disponible...')
        candidates = (
            glob.glob(r"C:\SRP\SRP-SR-*.csv") +
            glob.glob(r"C:\Users\aicil\Downloads\SRP-SR-*.csv") +
            glob.glob(os.path.join(BASE_DIR, "SRP-SR-*.csv")) +
            glob.glob(os.path.join(BASE_DIR, "censia_descarga_hoy.csv"))
        )
        if candidates:
            best = max(candidates, key=os.path.getmtime)
            log.ok(f"Usando: {os.path.basename(best)}")
            return best
        log.err("No se encontró ningún CSV disponible. Abortando.")
        sys.exit(1)


# ══════════════════════════════════════════════════════════════════
# PASO 2: PROCESAR DATOS
# ══════════════════════════════════════════════════════════════════
def cargar_y_procesar_csv(csv_path):
    log.section('⚙ ', 'PASO 2: Procesando datos...')
    with open(csv_path, encoding='latin1') as f:
        content = f.read()
    lines = content.split('\n')
    fixed = []
    for line in lines:
        line = line.strip()
        if line.startswith('"') and line.endswith('"'):
            line = line[1:-1].replace('""', '"')
        fixed.append(line)
    df = pd.read_csv(StringIO('\n'.join(fixed)), encoding='latin1', low_memory=False)
    df.columns = df.columns.str.strip()
    skip = {'id', 'INSTITUCION', 'DELEGACION', 'ESTADO', 'JURISDICCION',
            'MUNICIPIO', 'CLUES', 'Fecha de registro', 'Temporada', 'SEMANA'}
    for col in [c for c in df.columns if c not in skip]:
        df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)
    df['Fecha de registro'] = pd.to_datetime(df['Fecha de registro'], errors='coerce')
    dur = df[(df['ESTADO'] == 'DURANGO') & (df['Temporada'].isin([2025, 2026]))].copy()
    fecha_max = dur['Fecha de registro'].max()
    fecha_inicio = fecha_max - timedelta(days=1) if pd.notna(fecha_max) else None
    semana = dur['SEMANA'].max() if 'SEMANA' in dur.columns else '?'
    log.ok(f"{len(dur)} registros de Durango 2026")
    try:
        log.ok(f"Semana actual: SE-{int(semana)}")
    except Exception:
        log.ok(f"Semana actual: {semana}")
    return {'dur': dur, 'fecha_max': fecha_max, 'fecha_inicio_diario': fecha_inicio}


# ══════════════════════════════════════════════════════════════════
# HELPERS DE CÁLCULO
# ══════════════════════════════════════════════════════════════════
def safe_col(d, col_name):
    col_name = col_name.strip()
    if col_name in d.columns:
        return int(d[col_name].sum())
    for c in d.columns:
        if col_name in c:
            return int(d[c].sum())
    return 0

def totales_biologico(d):
    srp1 = safe_col(d, 'SRP  PRIMERA TOTAL') or safe_col(d, 'SRP PRIMERA TOTAL')
    srp2 = safe_col(d, 'SRP SEGUNDA TOTAL')
    sr1  = safe_col(d, 'SR PRIMERA TOTAL')
    sr2  = safe_col(d, 'SR SEGUNDA TOTAL')
    return srp1, srp2, sr1, sr2

def totales_por_inst(d, inst):
    return totales_biologico(d[d['INSTITUCION'] == inst])

def totales_por_jur(d, jur):
    return totales_biologico(d[(d['JURISDICCION'] == jur) & (d['INSTITUCION'] == 'SSA')])

def get_dosis_edad(d, srp_prefix, sr_prefix):
    srp1 = srp2 = sr1 = sr2 = 0
    for c in d.columns:
        cn = c.strip()
        if srp_prefix.strip() in cn:
            if 'PRIMERA' in cn.upper():  srp1 += int(d[c].sum())
            elif 'SEGUNDA' in cn.upper(): srp2 += int(d[c].sum())
        if sr_prefix.strip() in cn and not cn.startswith('SRP'):
            if 'PRIMERA' in cn.upper():  sr1 += int(d[c].sum())
            elif 'SEGUNDA' in cn.upper(): sr2 += int(d[c].sum())
    return srp1, srp2, sr1, sr2

def n(x):
    return "{:,}".format(int(x))


# ══════════════════════════════════════════════════════════════════
# GRÁFICA (RUBRO 5)
# ══════════════════════════════════════════════════════════════════
def generar_grafica(d_diario, periodo_label):
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import numpy as np
    from matplotlib.patches import FancyBboxPatch

    labels  = [INST_LABEL_CHART[i] for i in INSTITUCIONES]
    srp_tot = []
    sr_tot  = []
    for inst in INSTITUCIONES:
        s1, s2, r1, r2 = totales_por_inst(d_diario, inst)
        srp_tot.append(s1 + s2)
        sr_tot.append(r1 + r2)

    x     = np.arange(len(labels))
    width = 0.38

    fig, ax = plt.subplots(figsize=(9, 4.2), dpi=150)
    fig.patch.set_facecolor('#FFFFFF')
    ax.set_facecolor('#FFFFFF')

    bars1 = ax.bar(x - width/2, srp_tot, width, label='SRP',
                   color='#1B3A6B', zorder=3, edgecolor='white', linewidth=0.5)
    bars2 = ax.bar(x + width/2, sr_tot,  width, label='SR',
                   color='#2E6DA4', zorder=3, edgecolor='white', linewidth=0.5)

    ax.set_title(
        f'5. Dosis aplicadas por institución — Últimas 24 horas',
        fontsize=10, fontweight='bold', color='#1B3A6B', pad=10
    )
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=9, color='#1B3A6B')
    ax.set_ylabel('Dosis aplicadas', fontsize=9, color='#1B3A6B')
    ax.yaxis.grid(True, linestyle='--', alpha=0.5, color='#D6E4F0', zorder=0)
    ax.set_axisbelow(True)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#D6E4F0')
    ax.spines['bottom'].set_color('#D6E4F0')
    ax.tick_params(colors='#1B3A6B')
    legend = ax.legend(fontsize=9, framealpha=0.8, edgecolor='#D6E4F0')

    for bar in bars1:
        h = bar.get_height()
        if h > 0:
            ax.text(bar.get_x() + bar.get_width() / 2, h + max(srp_tot + sr_tot) * 0.02,
                    str(int(h)), ha='center', va='bottom', fontsize=8,
                    color='#1B3A6B', fontweight='bold')
    for bar in bars2:
        h = bar.get_height()
        if h > 0:
            ax.text(bar.get_x() + bar.get_width() / 2, h + max(srp_tot + sr_tot) * 0.02,
                    str(int(h)), ha='center', va='bottom', fontsize=8,
                    color='#2E6DA4', fontweight='bold')

    # Fuente pequeña debajo
    fig.text(0.5, -0.01, periodo_label, ha='center', fontsize=7.5,
             color='#7F8C8D', style='italic')

    plt.tight_layout(pad=1.2)
    ruta = os.path.join(BASE_DIR, "_grafica_diaria_tmp.png")
    plt.savefig(ruta, dpi=150, bbox_inches='tight', facecolor='#FFFFFF')
    plt.close()
    return ruta


# ══════════════════════════════════════════════════════════════════
# PASO 3: GENERAR INFORME WORD
# ══════════════════════════════════════════════════════════════════
def generar_informe_word(datos):
    log.section('📄', 'PASO 3: Generando informe Word...')

    from docx import Document
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    dur          = datos['dur']
    fecha_max    = datos['fecha_max']
    fecha_inicio = datos['fecha_inicio_diario']

    # Filtros temporales
    d_diario = dur[dur['Fecha de registro'] > fecha_inicio] if (
        pd.notna(fecha_max) and fecha_inicio is not None
    ) else dur.iloc[0:0]

    # Cadenas de fecha
    try:
        locale.setlocale(locale.LC_TIME, 'Spanish_Mexico.1252')
    except Exception:
        try:
            locale.setlocale(locale.LC_TIME, 'es_MX.UTF-8')
        except Exception:
            pass

    # Usar la fecha actual como corte (hoy)
    fecha_corte_str = date.today().strftime("%d DE %B DE %Y").upper()
    fecha_hora_consulta = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    fecha_hora_reporte = fecha_hora_consulta
    fecha_hora_ultimo = ""
    periodo_diario = ""
    periodo_encabezado = ""
    if fecha_inicio is not None:
        periodo_diario = f"{fecha_inicio.strftime('%d/%m/%Y %H:%M:%S')}-{date.today().strftime('%d/%m/%Y %H:%M:%S')}"
        periodo_encabezado = f"{fecha_inicio.strftime('%d de %B de %Y %H:%M')} a {date.today().strftime('%d de %B de %Y %H:%M')}"

    # ── Abrir plantilla (preserva header membretado) o crear documento nuevo ──
    from docx.shared import Cm
    if PLANTILLA_DOCX and os.path.exists(PLANTILLA_DOCX):
        doc = Document(PLANTILLA_DOCX)
        for p in doc.paragraphs:
            p.clear()
        log.ok(f"Usando plantilla membretada: {os.path.basename(PLANTILLA_DOCX)}")
    else:
        doc = Document()
        log.warn("Plantilla membretada no encontrada. Generando documento sin membrete.")

    # ── Ajustar márgenes para que el contenido no invada el logo del membrete ──
    for section in doc.sections:
        section.top_margin    = Cm(4.5)   # espacio suficiente bajo el logo
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)

    # ══ HELPERS XML ══
    def set_keep_with_next(paragraph):
        """Evita que el título quede separado de su tabla en salto de página."""
        pPr = paragraph._p.get_or_add_pPr()
        kn  = OxmlElement('w:keepNext')
        pPr.append(kn)

    def set_cant_split(table):
        """Evita que las filas de la tabla se corten entre páginas."""
        for row in table.rows:
            trPr = row._tr.get_or_add_trPr()
            cs   = OxmlElement('w:cantSplit')
            trPr.append(cs)

    def add_page_break():
        p  = doc.add_paragraph()
        r  = p.add_run()
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        r._r.append(br)
        return p

    def shd_cell(cell, hex_str):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        s    = OxmlElement('w:shd')
        s.set(qn('w:val'),   'clear')
        s.set(qn('w:color'), 'auto')
        s.set(qn('w:fill'),  hex_str)
        tcPr.append(s)

    def ct(cell, text, bold=False, sz=8, color=None,
           align=WD_ALIGN_PARAGRAPH.CENTER):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(1.5)
        p.paragraph_format.space_after  = Pt(1.5)
        r = p.add_run(str(text))
        r.bold      = bold
        r.font.size = Pt(sz)
        r.font.name = 'Calibri'
        if color:
            r.font.color.rgb = color

    def add_title(text, sz=12, bold=True, color=HDR_COLOR,
                  align=WD_ALIGN_PARAGRAPH.CENTER):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(text)
        r.bold = bold; r.font.size = Pt(sz)
        r.font.name = 'Calibri'; r.font.color.rgb = color
        return p

    def add_section_title(num, text, page_break_before=False):
        if page_break_before:
            add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(f"{num}. {text}")
        r.bold = True; r.font.size = Pt(10)
        r.font.name = 'Calibri'; r.font.color.rgb = HDR_COLOR
        set_keep_with_next(p)
        return p

    def add_subtitle(text, page_break_before=False):
        if page_break_before:
            add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(text)
        r.bold = True; r.font.size = Pt(9)
        r.font.name = 'Calibri'; r.font.color.rgb = SUB_COLOR
        set_keep_with_next(p)
        return p

    def build_table(headers, rows,
                    hdr_hex=C_PRIMARY, alt_hex=C_ALT1,
                    alt2_hex=C_ALT2,  tot_hex=C_TOTAL):
        tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.style  = 'Table Grid'
        tbl.autofit = True
        # Encabezado
        for i, h in enumerate(headers):
            shd_cell(tbl.cell(0, i), hdr_hex)
            ct(tbl.cell(0, i), h, bold=True, color=WHT_COLOR, sz=8)
        # Filas
        for ri, row in enumerate(rows):
            is_total = (ri == len(rows) - 1)
            bg = tot_hex if is_total else (alt_hex if ri % 2 == 0 else alt2_hex)
            for ci, val in enumerate(row):
                shd_cell(tbl.cell(ri + 1, ci), bg)
                al = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
                ct(tbl.cell(ri + 1, ci), val,
                   bold=is_total, align=al, sz=8,
                   color=HDR_COLOR if is_total else BLK_COLOR)
        set_cant_split(tbl)
        return tbl

    def spacer(pt=3):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(pt)

    def add_image(path_img, width_cm=13.0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        p.add_run().add_picture(path_img, width=Cm(width_cm))
        return p

    # ══════════════════════════════════════════════
    # ENCABEZADOS DEL DOCUMENTO
    # ══════════════════════════════════════════════
    add_title("SECRETARÍA DE SALUD", sz=13, bold=True, color=HDR_COLOR)
    add_title("DIRECCIÓN DE SALUD PÚBLICA", sz=11, bold=True, color=SUB_COLOR)
    add_title("REPORTE DE SEGUIMIENTO DE VACUNACIÓN SARAMPIÓN",
              sz=12, bold=True, color=HDR_COLOR)
    add_title("CORTE AL " + fecha_corte_str, sz=11, bold=True, color=HDR_COLOR)

    # ── RUBRO 1: Biológico Global ──
    srp1_g, srp2_g, sr1_g, sr2_g = totales_biologico(dur)
    add_section_title(1, "Dosis aplicadas por biológico y esquema (Global)")
    build_table(
        ['Biológico', '1a Dosis', '2a Dosis', 'Total'],
        [
            ['SRP',        n(srp1_g), n(srp2_g), n(srp1_g+srp2_g)],
            ['SR',         n(sr1_g),  n(sr2_g),  n(sr1_g+sr2_g)],
            ['Gran Total', n(srp1_g+sr1_g), n(srp2_g+sr2_g),
                           n(srp1_g+srp2_g+sr1_g+sr2_g)],
        ]
    )

    # ── RUBRO 2: Biológico 24h ──
    srp1_d, srp2_d, sr1_d, sr2_d = totales_biologico(d_diario)
    add_section_title(2, f"Dosis aplicadas por biológico y esquema (Del {periodo_encabezado})")
    build_table(
        ['Biológico', '1a Dosis', '2a Dosis', 'Total'],
        [
            ['SRP',        n(srp1_d), n(srp2_d), n(srp1_d+srp2_d)],
            ['SR',         n(sr1_d),  n(sr2_d),  n(sr1_d+sr2_d)],
            ['Gran Total', n(srp1_d+sr1_d), n(srp2_d+sr2_d),
                           n(srp1_d+srp2_d+sr1_d+sr2_d)],
        ]
    )

    # ── RUBRO 3: Institución Global ──
    add_section_title(3, "Dosis aplicadas por institución (Global)")
    rows3 = []; t1=t2=t3=t4=0
    for inst in INSTITUCIONES:
        s1,s2,r1,r2 = totales_por_inst(dur, inst)
        t1+=s1;t2+=s2;t3+=r1;t4+=r2
        rows3.append([INST_LABEL[inst], n(s1),n(s2),n(s1+s2),n(r1),n(r2),n(r1+r2),n(s1+s2+r1+r2)])
    rows3.append(['Total', n(t1),n(t2),n(t1+t2),n(t3),n(t4),n(t3+t4),n(t1+t2+t3+t4)])
    build_table(
        ['Institución','SRP 1a','SRP 2a','Total\nSRP','SR 1a','SR 2a','Total\nSR','Total'],
        rows3)

    # ── RUBRO 4: Institución 24h (nueva página para que no tape logo) ──
    add_section_title(4, f"Dosis aplicadas por institución (Del {periodo_encabezado})",
                      page_break_before=True)
    rows4 = []; t1=t2=t3=t4=0
    for inst in INSTITUCIONES:
        s1,s2,r1,r2 = totales_por_inst(d_diario, inst)
        t1+=s1;t2+=s2;t3+=r1;t4+=r2
        rows4.append([INST_LABEL[inst], n(s1),n(s2),n(s1+s2),n(r1),n(r2),n(r1+r2),n(s1+s2+r1+r2)])
    rows4.append(['Total', n(t1),n(t2),n(t1+t2),n(t3),n(t4),n(t3+t4),n(t1+t2+t3+t4)])
    build_table(
        ['Institución','SRP 1a','SRP 2a','Total SRP','SR 1a','SR 2a','Total SR','Total'],
        rows4)

    # ── RUBRO 5: Gráfica 24h ──
    add_section_title(5, f"Gráfica de dosis aplicadas por institución (Del {periodo_encabezado})")
    try:
        ruta_grafica = generar_grafica(d_diario, f"Del {periodo_encabezado}")
        add_image(ruta_grafica, width_cm=13.0)
    except Exception as e:
        log.warn(f"No se pudo generar la gráfica: {e}")
        spacer()

    # ── RUBRO 6: Jurisdicción Global (nueva página) ──
    add_section_title(6, "Desglose SSD por Jurisdicción Sanitaria (Global)",
                      page_break_before=True)
    rows6 = []; t1=t2=t3=t4=0
    for jur in JURS:
        s1,s2,r1,r2 = totales_por_jur(dur, jur)
        t1+=s1;t2+=s2;t3+=r1;t4+=r2
        rows6.append([JL[jur], n(s1),n(s2),n(s1+s2),n(r1),n(r2),n(r1+r2),n(s1+s2+r1+r2)])
    rows6.append(['Total SSD', n(t1),n(t2),n(t1+t2),n(t3),n(t4),n(t3+t4),n(t1+t2+t3+t4)])
    build_table(
        ['Jurisdicción','SRP\n1a','SRP\n2a','T. SRP','SR 1a','SR 2a','T. SR','Total'],
        rows6)

    # ── RUBRO 7: Jurisdicción 24h ──
    add_section_title(7, f"Desglose SSD por Jurisdicción Sanitaria (Del {periodo_encabezado})")
    rows7 = []; t1=t2=t3=t4=0
    for jur in JURS:
        s1,s2,r1,r2 = totales_por_jur(d_diario, jur)
        t1+=s1;t2+=s2;t3+=r1;t4+=r2
        rows7.append([JL[jur], n(s1),n(s2),n(s1+s2),n(r1),n(r2),n(r1+r2),n(s1+s2+r1+r2)])
    rows7.append(['Total SSD', n(t1),n(t2),n(t1+t2),n(t3),n(t4),n(t3+t4),n(t1+t2+t3+t4)])
    build_table(
        ['Jurisdicción','SRP 1a','SRP 2a','T. SRP','SR 1a','SR 2a','T. SR','Total'],
        rows7)

    # ── RUBRO 8: Grupo de edad Global (nueva página) ──
    add_section_title(8, "Desglose por grupo de edad (Global)",
                      page_break_before=True)

    add_subtitle("Vacuna SRP")
    rows8a = []; t1=t2=0
    for label, sp, rp in GRUPOS_EDAD:
        s1,s2,_,_ = get_dosis_edad(dur, sp, 'XXXIGNORE')
        t1+=s1; t2+=s2
        d1 = '—' if any(p.strip() in sp for p in SIN_1A_DOSIS_SRP) else n(s1)
        d2 = '—' if any(p.strip() in sp for p in SIN_2A_DOSIS_SRP) else n(s2)
        rows8a.append([label, d1, d2])
    rows8a.append(['Total', n(t1), n(t2)])
    build_table(['Grupo de Edad','1a Dosis','2a Dosis'], rows8a)

    spacer(4)
    add_subtitle("Vacuna SR")
    rows8b = []; t1=t2=0
    for label, sp, rp in GRUPOS_EDAD:
        _,_,r1,r2 = get_dosis_edad(dur, 'XXXIGNORE', rp)
        t1+=r1; t2+=r2
        d1 = '—' if any(p.strip() in rp for p in SIN_1A_DOSIS_SR) else n(r1)
        d2 = '—' if any(p.strip() in rp for p in SIN_2A_DOSIS_SR) else n(r2)
        rows8b.append([label, d1, d2])
    rows8b.append(['Total', n(t1), n(t2)])
    build_table(['Grupo de Edad','1a Dosis','2a Dosis'], rows8b)

    # ── RUBRO 9: Grupo de edad 24h (nueva página) ──
    add_section_title(9, f"Desglose por grupo de edad (Del {periodo_encabezado})",
                      page_break_before=True)

    add_subtitle("Vacuna SRP")
    rows9a = []; t1=t2=0
    for label, sp, rp in GRUPOS_EDAD:
        s1,s2,_,_ = get_dosis_edad(d_diario, sp, 'XXXIGNORE')
        t1+=s1; t2+=s2
        d1 = '—' if any(p.strip() in sp for p in SIN_1A_DOSIS_SRP) else n(s1)
        d2 = '—' if any(p.strip() in sp for p in SIN_2A_DOSIS_SRP) else n(s2)
        rows9a.append([label, d1, d2])
    rows9a.append(['Total', n(t1), n(t2)])
    build_table(['Grupo de Edad','1a Dosis','2a Dosis'], rows9a)

    spacer(4)
    add_subtitle("Vacuna SR")
    rows9b = []; t1=t2=0
    for label, sp, rp in GRUPOS_EDAD:
        _,_,r1,r2 = get_dosis_edad(d_diario, 'XXXIGNORE', rp)
        t1+=r1; t2+=r2
        d1 = '—' if any(p.strip() in rp for p in SIN_1A_DOSIS_SR) else n(r1)
        d2 = '—' if any(p.strip() in rp for p in SIN_2A_DOSIS_SR) else n(r2)
        rows9b.append([label, d1, d2])
    rows9b.append(['Total', n(t1), n(t2)])
    build_table(['Grupo de Edad','1a Dosis','2a Dosis'], rows9b)

    # ══ PIE DE PÁGINA ══
    spacer(8)
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r  = pie.add_run(
        "*Fuente: SISTEMAS DE INFORMACIÓN CENTRO NACIONAL DE SALUD PARA LA INFANCIA Y LA\n"
        "ADOLESCENCIA (CENSIA). https://siscensia.salud.gob.mx/*\n"
        f"Fecha y hora de consulta: {fecha_hora_consulta}\n"
        f"Fecha y hora del reporte: {fecha_hora_reporte}\n"
        f"Fecha y hora del último registro: {fecha_hora_ultimo}\n"
        f"Periodo de reporte diario: {periodo_diario}"
    )
    r.font.size = Pt(7.5)
    r.font.name = 'Calibri'
    r.font.color.rgb = GRY_COLOR
    r.italic = True

    spacer(10)
    firma_p = doc.add_paragraph()
    firma_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = firma_p.add_run(FIRMA)
    r2.font.size = Pt(9)
    r2.font.name = 'Calibri'
    r2.bold = True
    r2.font.color.rgb = HDR_COLOR

    # ══ GUARDAR ══
    os.makedirs(CARPETA_INFORME, exist_ok=True)
    nombre   = "Vacunacion_Sarampion_" + datetime.now().strftime("%d.%m.%y") + ".docx"
    ruta_out = os.path.join(CARPETA_INFORME, nombre)
    doc.save(ruta_out)
    log.ok(f"Informe guardado: {ruta_out}")
    return ruta_out, fecha_corte_str


# ══════════════════════════════════════════════════════════════════
# PASO 4: ENVIAR POR CORREO
# ══════════════════════════════════════════════════════════════════
def enviar_informe(ruta_informe, fecha_corte_str, csv_path=None, test_email=None):
    log.section('📧', 'PASO 4: Enviando informe por correo...')
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text      import MIMEText
    from email.mime.base      import MIMEBase
    from email                import encoders

    if not os.path.exists(CONFIG_PATH):
        log.warn("No se encontró config_correo.json. No se enviará correo.")
        return

    # Credenciales: variables de entorno tienen prioridad (GitHub Actions)
    # Si no hay config_correo.json (CI), se usan solo env vars
    cfg = {}
    if os.path.exists(CONFIG_PATH):
        with open(CONFIG_PATH, "r", encoding="utf-8") as f:
            cfg = json.load(f)

    smtp_server = cfg.get("smtp_server", "smtp.gmail.com")
    smtp_port   = int(cfg.get("smtp_port", 465))
    use_tls     = cfg.get("use_tls", False)  # False = SSL directo (puerto 465)
    sender      = os.environ.get("GMAIL_USER") or cfg.get("sender_email", "")
    password    = os.environ.get("GMAIL_APP_PASSWORD") or cfg.get("sender_password", "")
    
    if test_email:
        recipients = [test_email]
        log.info(f"Modo de prueba activado: enviando ÚNICAMENTE a {test_email}")
    else:
        # MAIL_TO env var: lista separada por comas (para GitHub Actions)
        mail_to_env = os.environ.get("MAIL_TO", "")
        if mail_to_env:
            recipients = [m.strip() for m in mail_to_env.split(",") if m.strip()]
        else:
            recipients = cfg.get("recipient_emails", [])

    subject = f"Reporte Vacunación Sarampión — Corte al {fecha_corte_str} | Durango"
    
    body_text = (
        "Buen día, estimados:\n\n"
        "Por este medio se remite el informe diario de sarampión para su conocimiento y seguimiento.\n\n"
        "Quedo atenta a cualquier duda, comentario o aclaración que consideren pertinente.\n\n"
        "Reciban un cordial saludo.\n\n"
        "Atentamente,\n\n"
        "Alicia Esparza Aldaba\n"
        "Jefa del Departamento de Enfermedades Transmisibles"
    )
    
    body_html = (
        "<html><body style='font-family: Calibri, Arial, sans-serif; font-size: 11pt; color: #333333;'>\n"
        "<p>Buen día, estimados:</p>\n"
        "<p>Por este medio se remite el informe diario de sarampión para su conocimiento y seguimiento.</p>\n"
        "<p>Quedo atenta a cualquier duda, comentario o aclaración que consideren pertinente.</p>\n"
        "<p>Reciban un cordial saludo.</p>\n"
        "<p>Atentamente,</p>\n"
        "<p><strong>Alicia Esparza Aldaba</strong><br>\n"
        "Jefa del Departamento de Enfermedades Transmisibles</p>\n"
        "</body></html>"
    )

    if not sender or not recipients:
        log.warn("Credenciales o destinatarios incompletos en config_correo.json.")
        return

    msg = MIMEMultipart()
    msg['From']    = sender
    msg['To']      = ", ".join(recipients)
    msg['Subject'] = subject
    
    # Agregar cuerpo en formato texto y html alternativo
    alt_part = MIMEMultipart('alternative')
    alt_part.attach(MIMEText(body_text, 'plain', 'utf-8'))
    alt_part.attach(MIMEText(body_html, 'html', 'utf-8'))
    msg.attach(alt_part)

    # Adjuntar reporte Word
    if os.path.exists(ruta_informe):
        with open(ruta_informe, "rb") as f:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(f.read())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition",
                        f"attachment; filename={os.path.basename(ruta_informe)}")
        msg.attach(part)
        log.raw(f"   📎 Adjunto: {os.path.basename(ruta_informe)}")

    # Adjuntar CSV de CeNSIA
    if csv_path and os.path.exists(csv_path):
        with open(csv_path, "rb") as f:
            part_csv = MIMEBase("application", "octet-stream")
            part_csv.set_payload(f.read())
        encoders.encode_base64(part_csv)
        part_csv.add_header("Content-Disposition",
                            f"attachment; filename={os.path.basename(csv_path)}")
        msg.attach(part_csv)
        log.raw(f"   📎 Adjunto: {os.path.basename(csv_path)}")

    max_retries = 3
    for attempt in range(1, max_retries + 1):
        try:
            if use_tls:
                server = smtplib.SMTP(smtp_server, smtp_port, timeout=60)
                server.ehlo()
                server.starttls()
                server.ehlo()
            else:
                server = smtplib.SMTP_SSL(smtp_server, smtp_port, timeout=60)
            server.login(sender, password)
            server.sendmail(sender, recipients, msg.as_string())
            server.quit()
            log.ok(f"Correo enviado a: {', '.join(recipients)} (intento {attempt})")
            break
        except Exception as e:
            log.err(f"Error al enviar correo (intento {attempt}): {e}")
            if attempt < max_retries:
                log.info("Esperando 5 segundos antes de reintentar...")
                time.sleep(5)
            else:
                log.err("Todos los intentos fallaron. Abortando.")
                raise


# ══════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Genera y envía el reporte diario de sarampión.")
    parser.add_argument("--test-email", type=str, help="Envía una prueba únicamente a este correo electrónico.")
    args = parser.parse_args()

    sep = '=' * 60
    for line in [sep,
                 "  INFORME DIARIO AUTOMATIZADO — SARAMPIÓN DURANGO 2026",
                 "  " + datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
                 sep]:
        log.raw(line)

    csv_path              = descargar_csv_censia()
    datos                 = cargar_y_procesar_csv(csv_path)
    ruta_doc, fecha_corte = generar_informe_word(datos)
    
    # Convertir a PDF
    log.section('📄', 'PASO 3.5: Convirtiendo informe a PDF...')
    ruta_pdf = ruta_doc.replace(".docx", ".pdf")
    convertido = False

    # Intentar primero con LibreOffice (funciona en Linux/GitHub Actions y Windows)
    import subprocess, shutil
    libreoffice_cmds = ["libreoffice", "soffice",
                        r"C:\Program Files\LibreOffice\program\soffice.exe"]
    for cmd in libreoffice_cmds:
        if shutil.which(cmd) or os.path.exists(cmd):
            try:
                out_dir = os.path.dirname(ruta_doc)
                result = subprocess.run(
                    [cmd, "--headless", "--convert-to", "pdf",
                     "--outdir", out_dir, ruta_doc],
                    capture_output=True, text=True, timeout=120
                )
                if result.returncode == 0 and os.path.exists(ruta_pdf):
                    log.ok(f"Informe PDF generado (LibreOffice): {ruta_pdf}")
                    ruta_adjunto = ruta_pdf
                    convertido = True
                    break
                else:
                    log.warn(f"LibreOffice retornó código {result.returncode}: {result.stderr[:200]}")
            except Exception as e:
                log.warn(f"LibreOffice falló: {e}")
            break

    # Fallback: docx2pdf (solo Windows con Word instalado)
    if not convertido:
        try:
            from docx2pdf import convert
            convert(ruta_doc, ruta_pdf)
            log.ok(f"Informe PDF generado (docx2pdf): {ruta_pdf}")
            ruta_adjunto = ruta_pdf
            convertido = True
        except Exception as e:
            log.err(f"No se pudo convertir a PDF: {e}")
            log.info("Se enviará el archivo DOCX en su lugar.")
            ruta_adjunto = ruta_doc

    # Actualizar report_path en config_correo.json
    try:
        if os.path.exists(CONFIG_PATH):
            with open(CONFIG_PATH, "r", encoding="utf-8") as f:
                cfg_data = json.load(f)
            cfg_data["report_path"] = ruta_adjunto
            with open(CONFIG_PATH, "w", encoding="utf-8") as f:
                json.dump(cfg_data, f, indent=2, ensure_ascii=False)
            log.ok("Ruta del informe actualizada en config_correo.json")
    except Exception as ex:
        log.warn(f"No se pudo actualizar report_path en config_correo.json: {ex}")

    enviar_informe(ruta_adjunto, fecha_corte, csv_path, test_email=args.test_email)

    for line in ["", sep,
                 f"  ✅ PROCESO COMPLETADO — {datetime.now().strftime('%H:%M:%S')}",
                 f"  📄 Informe: {ruta_adjunto}",
                 sep]:
        log.raw(line)
