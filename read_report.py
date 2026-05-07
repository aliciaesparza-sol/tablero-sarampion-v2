import docx

file_path = r"C:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\mezquital\INFORMES\Informe_Vacunacion_Mezquital_2026.15.04.2026.docx"

try:
    doc = docx.Document(file_path)
    print(f"Report: {file_path}")
    print("--- CONTENT ---")
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
    
    print("\n--- TABLES ---")
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i}:")
        for row in table.rows:
            print(" | ".join(cell.text.strip() for cell in row.cells))

except Exception as e:
    print(f"Error: {e}")
